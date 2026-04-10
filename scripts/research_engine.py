#!/usr/bin/env python3
"""Zenith Research Engine — AI-powered research assistant actions.
Provides: research_chat, search_papers, web_search_action, extract_pdf_text,
          check_novelty, verify_citations, run_experiment_action, export_chat,
          generate_section

Inspired by AutoResearchClaw (https://github.com/aiming-lab/AutoResearchClaw)
by Aiming Lab. Full credits in README.md.
"""
import os, sys, json, tempfile, time, re, random, urllib.request, urllib.error, urllib.parse

try:
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
except ImportError:
    pass

try:
    import requests
    _HAS_REQUESTS = True
except ImportError:
    _HAS_REQUESTS = False

try:
    from scihub import SciHub
    _HAS_SCIHUB = True
except Exception:
    _HAS_SCIHUB = False
# ── Retry / rate-limit constants (mirrors AutoResearchClaw patterns) ──
_MAX_RETRIES = 3
_RETRY_BASE_DELAY = 2.0
_MAX_BACKOFF_SEC = 60
_RETRYABLE_CODES = frozenset({429, 500, 502, 503, 504, 529})
_USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36"
_NO_TEMPERATURE_MODELS = frozenset({"o3", "o3-mini", "o4-mini", "deepseek-reasoner"})

TEMP_DIR = os.path.join(tempfile.gettempdir(), "Zenith")
RESEARCH_DIR = os.path.join(TEMP_DIR, "Research")
EXPORTS_DIR = os.path.join(RESEARCH_DIR, "exports")
EXPERIMENTS_DIR = os.path.join(RESEARCH_DIR, "experiments")
PAPERS_DIR = os.path.join(RESEARCH_DIR, "papers")

VECTORDB_DIR = os.path.join(RESEARCH_DIR, "vector_db")
PROMPT_LOGS_DIR = os.path.join(RESEARCH_DIR, "prompt_logs")
for _d in [RESEARCH_DIR, EXPORTS_DIR, EXPERIMENTS_DIR, PAPERS_DIR, VECTORDB_DIR, PROMPT_LOGS_DIR]:
    os.makedirs(_d, exist_ok=True)


# ── Prompt Logger — disk-based, survives subprocess boundaries ──
# Each invoke("process_file") runs a NEW Python process, so in-memory dicts are
# wiped between phases.  We persist to a single JSON file on disk instead.

_PROMPT_LOG_FILE = os.path.join(PROMPT_LOGS_DIR, "_active_pipeline_prompts.json")


def _log_prompt(agent_name, system_prompt, user_prompt, variables, output_text, tokens=None):
    """Append a prompt exchange to the on-disk log file."""
    entry = {
        "agent": agent_name,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "system_prompt": system_prompt,
        "user_prompt": user_prompt,
        "variables": variables,
        "output": output_text if isinstance(output_text, str) else str(output_text),
        "tokens": tokens or {},
    }
    # Append atomically — read → append → write
    try:
        existing = []
        if os.path.isfile(_PROMPT_LOG_FILE):
            with open(_PROMPT_LOG_FILE, "r", encoding="utf-8") as f:
                existing = json.load(f)
        existing.append(entry)
        with open(_PROMPT_LOG_FILE, "w", encoding="utf-8") as f:
            json.dump(existing, f, ensure_ascii=False)
    except Exception:
        pass  # never crash the pipeline over logging


def _clear_prompt_logs():
    """Clear all prompt logs (call at pipeline start)."""
    try:
        if os.path.isfile(_PROMPT_LOG_FILE):
            os.remove(_PROMPT_LOG_FILE)
    except Exception:
        pass


def _get_prompt_logs():
    """Read all prompt logs from disk, grouped by agent name."""
    result = {}
    try:
        if os.path.isfile(_PROMPT_LOG_FILE):
            with open(_PROMPT_LOG_FILE, "r", encoding="utf-8") as f:
                entries = json.load(f)
            for entry in entries:
                agent = entry.pop("agent", "unknown")
                if agent not in result:
                    result[agent] = []
                result[agent].append(entry)
    except Exception:
        pass
    return result


# ── ChromaDB Vector Store (lightweight local vector DB for GraphRAG) ──
try:
    import chromadb
    _HAS_CHROMADB = True
except ImportError:
    _HAS_CHROMADB = False


def _init_vector_collection(project_id: str):
    """Initialize or get a ChromaDB collection for a research project."""
    if not _HAS_CHROMADB:
        return None
    db_path = os.path.join(VECTORDB_DIR, project_id)
    os.makedirs(db_path, exist_ok=True)
    client = chromadb.PersistentClient(path=db_path)
    collection = client.get_or_create_collection(
        name="research_papers",
        metadata={"hnsw:space": "cosine"},
    )
    return collection


def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
    """Split text into overlapping chunks for embedding."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


def ingest_into_vectordb(args):
    """Phase 2.2 — Ingest extracted texts into ChromaDB vector store.
    Args: {project_id, papers: [{title, doi, text}], query}
    Returns: {ok, chunks_stored, collection_size}"""
    project_id = args.get("project_id", "default")
    papers = args.get("papers", [])
    query = args.get("query", "")

    if not _HAS_CHROMADB:
        return {"ok": True, "chunks_stored": 0, "collection_size": 0,
                "warning": "chromadb not installed. Install via: pip install chromadb. Proceeding without vector search."}

    collection = _init_vector_collection(project_id)
    if not collection:
        return {"ok": True, "chunks_stored": 0, "collection_size": 0,
                "warning": "Could not initialize vector DB."}

    chunks_stored = 0
    for pi, paper in enumerate(papers):
        text = paper.get("text", "")
        title = paper.get("title", f"Paper {pi}")
        doi = paper.get("doi", "")
        if not text or len(text) < 50:
            continue
        chunks = _chunk_text(text, chunk_size=1200, overlap=200)
        ids = [f"p{pi}_c{ci}" for ci in range(len(chunks))]
        metadatas = [{"paper_idx": pi, "title": title, "doi": doi, "chunk_idx": ci, "query": query} for ci in range(len(chunks))]
        # Batch add (ChromaDB handles embeddings internally with default model)
        try:
            collection.add(documents=chunks, ids=ids, metadatas=metadatas)
            chunks_stored += len(chunks)
        except Exception as e:
            # Duplicates or other issues — skip silently
            if "already exists" not in str(e).lower():
                pass

    return {"ok": True, "chunks_stored": chunks_stored,
            "collection_size": collection.count()}


def query_vectordb(args):
    """Query the vector DB for relevant chunks.
    Args: {project_id, query, n_results, section_type}
    Returns: {ok, results: [{text, title, doi, score}]}"""
    project_id = args.get("project_id", "default")
    query_text = args.get("query", "")
    n_results = args.get("n_results", 10)
    section_type = args.get("section_type", "")

    if not _HAS_CHROMADB:
        return {"ok": True, "results": [], "warning": "chromadb not installed"}

    collection = _init_vector_collection(project_id)
    if not collection or collection.count() == 0:
        return {"ok": True, "results": []}

    search_query = f"{section_type}: {query_text}" if section_type else query_text
    try:
        results = collection.query(query_texts=[search_query], n_results=min(n_results, collection.count()))
        formatted = []
        if results and results.get("documents"):
            for i, doc in enumerate(results["documents"][0]):
                meta = results["metadatas"][0][i] if results.get("metadatas") else {}
                dist = results["distances"][0][i] if results.get("distances") else 0
                formatted.append({
                    "text": doc, "title": meta.get("title", ""),
                    "doi": meta.get("doi", ""), "score": round(1 - dist, 4),
                })
        return {"ok": True, "results": formatted}
    except Exception as e:
        return {"ok": True, "results": [], "error": str(e)}


# ══════════════════════════════════════════════════════════════════════════════
# ██  LLM CALL HELPERS (mirrors _call_llm from process_files.py)
# ══════════════════════════════════════════════════════════════════════════════

## ── Model Tier Resolution ──

TIER_MODELS = {
    "google": {"strong": "gemini-3.1-pro-preview", "fast": "gemini-3.1-flash-lite-preview"},
    "openai": {"strong": "gpt-4.1", "fast": "gpt-4.1-mini"},
    "anthropic": {"strong": "claude-sonnet-4-5-20260115", "fast": "claude-haiku-4-5-20250514"},
    "deepseek": {"strong": "deepseek-reasoner", "fast": "deepseek-chat"},
    "groq": {"strong": "llama-3.3-70b-versatile", "fast": "llama-3.1-8b-instant"},
}

def _resolve_model(provider, tier, fallback_model=""):
    """Resolve a model tier ('strong'/'fast') to a concrete model ID."""
    if not tier or tier == "none":
        return fallback_model
    return TIER_MODELS.get(provider, {}).get(tier, fallback_model)


def _llm_chat(provider, api_key, model, messages, temperature=0.7, max_tokens=16384,
              response_schema=None, thinking_config=None, code_execution=False, google_search=False):
    """Multi-turn LLM chat with retry logic (follows AutoResearchClaw patterns).
    messages = [{role, content}, ...].

    Enhanced parameters (Gemini-specific, graceful degradation for others):
      response_schema: JSON Schema dict for structured output (Gemini: native, others: prompt injection)
      thinking_config: {"budget": int} to enable Gemini thinking mode
      code_execution: bool to enable Gemini inline code execution
      google_search: bool to enable native Gemini Google Search grounding

    Returns {text, usage: {input_tokens, output_tokens}, structured: dict|None}."""

    # For non-Google providers, inject structured output schema into system prompt
    effective_messages = list(messages)
    if response_schema and provider != "google":
        schema_instruction = f"\n\nYou MUST respond with ONLY valid JSON matching this exact schema (no markdown fences, no commentary):\n{json.dumps(response_schema, indent=2)}"
        # Find system message and append, or prepend as new system message
        found_sys = False
        for i, m in enumerate(effective_messages):
            if m["role"] == "system":
                effective_messages[i] = {**m, "content": m["content"] + schema_instruction}
                found_sys = True
                break
        if not found_sys:
            effective_messages.insert(0, {"role": "system", "content": f"Respond with valid JSON matching this schema:{schema_instruction}"})

    def _build_request():
        """Build the urllib request for the given provider."""
        if provider == "openai":
            url = "https://api.openai.com/v1/chat/completions"
            mdl = model or "gpt-5.4-nano"
            body = {"model": mdl, "messages": effective_messages, "max_tokens": max_tokens}
            # Reasoning models reject temperature param (ARC pattern)
            if not any(mdl.startswith(p) for p in _NO_TEMPERATURE_MODELS):
                body["temperature"] = temperature
            # OpenAI supports json_schema response format
            if response_schema:
                body["response_format"] = {"type": "json_object"}
            payload = json.dumps(body).encode()
            return urllib.request.Request(url, data=payload, headers={
                "Authorization": f"Bearer {api_key}", "Content-Type": "application/json",
                "User-Agent": _USER_AGENT}), mdl

        elif provider == "anthropic":
            url = "https://api.anthropic.com/v1/messages"
            mdl = model or "claude-sonnet-4-5-20260115"
            sys_msgs = [m for m in effective_messages if m["role"] == "system"]
            non_sys = [m for m in effective_messages if m["role"] != "system"]
            body = {"model": mdl, "max_tokens": max_tokens, "messages": non_sys,
                    "temperature": temperature}
            if sys_msgs:
                body["system"] = sys_msgs[0]["content"]
            payload = json.dumps(body).encode()
            return urllib.request.Request(url, data=payload, headers={
                "x-api-key": api_key, "Content-Type": "application/json",
                "anthropic-version": "2023-06-01", "User-Agent": _USER_AGENT}), mdl

        elif provider == "google":
            mdl = model or "gemini-3.1-flash-lite-preview"
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{mdl}:generateContent?key={api_key}"

            # Separate system instruction from conversation
            system_text = ""
            contents = []
            for m in effective_messages:
                if m["role"] == "system":
                    system_text += m["content"] + "\n"
                elif m["role"] == "assistant":
                    contents.append({"role": "model", "parts": [{"text": m["content"]}]})
                else:
                    contents.append({"role": "user", "parts": [{"text": m["content"]}]})

            # Build request body
            body = {"contents": contents}

            # System instruction (proper Gemini API field instead of faking as user/model turns)
            if system_text.strip():
                body["systemInstruction"] = {"parts": [{"text": system_text.strip()}]}

            # Generation config
            gen_config = {"maxOutputTokens": max_tokens}
            if temperature is not None:
                gen_config["temperature"] = temperature

            # Structured output (native Gemini response_schema)
            if response_schema:
                gen_config["responseMimeType"] = "application/json"
                gen_config["responseSchema"] = response_schema

            body["generationConfig"] = gen_config

            # Thinking mode (Gemini 2.5+ feature)
            if thinking_config:
                budget = thinking_config.get("budget", 8192)
                body["generationConfig"]["thinkingConfig"] = {"thinkingBudget": budget}

            tools_list = []
            # Code execution tool
            if code_execution:
                tools_list.append({"codeExecution": {}})
            # Google Search tool
            if google_search:
                tools_list.append({"googleSearch": {}})
                
            if tools_list:
                body["tools"] = tools_list

            payload = json.dumps(body).encode()
            return urllib.request.Request(url, data=payload, headers={
                "Content-Type": "application/json", "User-Agent": _USER_AGENT}), mdl

        elif provider == "deepseek":
            url = "https://api.deepseek.com/chat/completions"
            mdl = model or "deepseek-chat"
            payload = json.dumps({"model": mdl, "messages": effective_messages,
                                  "max_tokens": max_tokens, "temperature": temperature}).encode()
            return urllib.request.Request(url, data=payload, headers={
                "Authorization": f"Bearer {api_key}", "Content-Type": "application/json",
                "User-Agent": _USER_AGENT}), mdl

        elif provider == "groq":
            url = "https://api.groq.com/openai/v1/chat/completions"
            mdl = model or "llama-3.3-70b-versatile"
            payload = json.dumps({"model": mdl, "messages": effective_messages,
                                  "max_tokens": max_tokens, "temperature": temperature}).encode()
            return urllib.request.Request(url, data=payload, headers={
                "Authorization": f"Bearer {api_key}", "Content-Type": "application/json",
                "User-Agent": _USER_AGENT}), mdl

        else:
            return None, model or ""

    req, mdl = _build_request()
    if req is None:
        return {"error": f"Unknown provider: {provider}"}

    # ── Retry loop with exponential backoff + jitter (ARC pattern) ──
    last_error = None
    for attempt in range(_MAX_RETRIES):
        try:
            with urllib.request.urlopen(req, timeout=300) as resp:
                data = json.loads(resp.read().decode())

            text = ""
            usage = {"input_tokens": 0, "output_tokens": 0}
            structured = None

            if provider in ("openai", "groq", "deepseek"):
                text = data["choices"][0]["message"]["content"]
                u = data.get("usage", {})
                usage = {"input_tokens": u.get("prompt_tokens", 0),
                         "output_tokens": u.get("completion_tokens", 0)}
            elif provider == "anthropic":
                text = data["content"][0]["text"]
                u = data.get("usage", {})
                usage = {"input_tokens": u.get("input_tokens", 0),
                         "output_tokens": u.get("output_tokens", 0)}
            elif provider == "google":
                cands = data.get("candidates", [])
                if cands and cands[0].get("content", {}).get("parts"):
                    parts = cands[0]["content"]["parts"]
                    # Collect text from all parts (thinking mode may have multiple)
                    text_parts = []
                    code_results = []
                    for part in parts:
                        if "text" in part:
                            text_parts.append(part["text"])
                        elif "executableCode" in part:
                            code_results.append({"code": part["executableCode"].get("code", ""), "language": part["executableCode"].get("language", "")})
                        elif "codeExecutionResult" in part:
                            code_results.append({"output": part["codeExecutionResult"].get("output", ""), "outcome": part["codeExecutionResult"].get("outcome", "")})
                    text = "\n".join(text_parts)
                    if code_results:
                        structured = structured or {}
                        structured["code_execution_results"] = code_results
                else:
                    feedback = data.get("promptFeedback", {})
                    block_reason = feedback.get("blockReason", "")
                    if block_reason:
                        return {"error": f"Google blocked request: {block_reason}"}
                    return {"error": "Google returned empty response. Try a different query."}
                u = data.get("usageMetadata", {})
                usage = {"input_tokens": u.get("promptTokenCount", 0),
                         "output_tokens": u.get("candidatesTokenCount", 0)}

            # Parse structured output if we used response_schema
            if response_schema and text:
                try:
                    structured = json.loads(text)
                except json.JSONDecodeError:
                    # Try to extract JSON from the response
                    json_match = re.search(r'[\[{][\s\S]*[\]}]', text)
                    if json_match:
                        try:
                            structured = json.loads(json_match.group())
                        except json.JSONDecodeError:
                            pass

            result = {"text": text, "usage": usage}
            if structured is not None:
                result["structured"] = structured

            # Auto-log prompt for export
            _caller = ""
            import traceback as _tb
            _stack = _tb.extract_stack(limit=4)
            for _frame in reversed(_stack[:-1]):
                if _frame.name not in ("_llm_chat", "_build_request", "<module>"):
                    _caller = _frame.name
                    break
            if _caller:
                _sys_p = next((m["content"] for m in messages if m.get("role") == "system"), "")
                _usr_p = next((m["content"] for m in messages if m.get("role") == "user"), "")
                _log_prompt(_caller, _sys_p, _usr_p,
                            {"provider": provider, "model": mdl, "temperature": temperature,
                             "max_tokens": max_tokens, "code_execution": code_execution},
                            text, usage)

            return result

        except urllib.error.HTTPError as e:
            status = e.code
            body = ""
            try:
                body = e.read().decode()[:500]
            except Exception:
                pass
            last_error = f"API error {status}: {body[:300]}"

            # Non-retryable errors — fail immediately
            if status in (400, 401, 403, 404):
                if status == 400 and any(kw in body.lower() for kw in ("rate limit", "overloaded", "temporarily", "capacity", "throttl", "retry")):
                    pass  # fall through to retry
                else:
                    return {"error": last_error}

            # Retryable: 429, 500, 502, 503, 504, 529
            if status in _RETRYABLE_CODES or (status == 400 and "rate" in body.lower()):
                delay = min(_RETRY_BASE_DELAY * (2 ** attempt), _MAX_BACKOFF_SEC)
                delay += random.uniform(0, delay * 0.3)  # jitter
                time.sleep(delay)
                req, mdl = _build_request()
                if req is None:
                    return {"error": f"Unknown provider: {provider}"}
                continue

            return {"error": last_error}

        except (urllib.error.URLError, OSError) as e:
            last_error = f"Connection error: {e}"
            if attempt < _MAX_RETRIES - 1:
                delay = min(_RETRY_BASE_DELAY * (2 ** attempt), _MAX_BACKOFF_SEC)
                time.sleep(delay)
                req, mdl = _build_request()
                if req is None:
                    return {"error": f"Unknown provider: {provider}"}
                continue
            return {"error": last_error}

        except Exception as e:
            return {"error": str(e)}

    return {"error": f"All {_MAX_RETRIES} retries failed. Last: {last_error}"}


# ══════════════════════════════════════════════════════════════════════════════
# ██  TOOL IMPLEMENTATIONS
# ══════════════════════════════════════════════════════════════════════════════

def _search_arxiv(query, max_results=50):
    """Search arXiv via its Atom API."""
    try:
        q = urllib.parse.quote(query)
        url = f"http://export.arxiv.org/api/query?search_query=all:{q}&start=0&max_results={max_results}&sortBy=relevance"
        req = urllib.request.Request(url, headers={"User-Agent": "Zenith/1.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            raw = resp.read().decode()
        import xml.etree.ElementTree as ET
        root = ET.fromstring(raw)
        ns = {"a": "http://www.w3.org/2005/Atom"}
        papers = []
        for entry in root.findall("a:entry", ns):
            title = (entry.findtext("a:title", "", ns) or "").strip().replace("\n", " ")
            summary = (entry.findtext("a:summary", "", ns) or "").strip().replace("\n", " ")
            authors = [a.findtext("a:name", "", ns) for a in entry.findall("a:author", ns)]
            published = entry.findtext("a:published", "", ns)[:4] if entry.findtext("a:published", "", ns) else ""
            arxiv_id = (entry.findtext("a:id", "", ns) or "").split("/abs/")[-1]
            link = entry.findtext("a:id", "", ns) or ""
            papers.append({
                "title": title, "authors": authors, "year": published,
                "abstract": summary[:500], "doi": "", "citations": 0,
                "url": link, "source": "arXiv", "arxiv_id": arxiv_id,
            })
        return papers
    except Exception as e:
        return []


def _search_semantic_scholar(query, max_results=50, year_min=None):
    """Search Semantic Scholar API with retry (ARC pattern)."""
    try:
        q = urllib.parse.quote(query)
        url = f"https://api.semanticscholar.org/graph/v1/paper/search?query={q}&limit={max_results}&fields=title,authors,year,abstract,citationCount,externalIds,url"
        if year_min:
            url += f"&year={year_min}-"

        data = None
        for attempt in range(_MAX_RETRIES):
            try:
                req = urllib.request.Request(url, headers={
                    "User-Agent": _USER_AGENT, "Accept": "application/json"})
                with urllib.request.urlopen(req, timeout=30) as resp:
                    data = json.loads(resp.read().decode())
                break
            except urllib.error.HTTPError as e:
                if e.code == 429:
                    delay = min(2 ** (attempt + 1), 30)
                    time.sleep(delay + random.uniform(0, delay * 0.3))
                    continue
                return []
            except (urllib.error.URLError, OSError):
                if attempt < _MAX_RETRIES - 1:
                    time.sleep(2 ** attempt)
                    continue
                return []

        if data is None:
            return []

        papers = []
        for p in data.get("data", []):
            authors = [a.get("name", "") for a in (p.get("authors") or [])]
            doi = (p.get("externalIds") or {}).get("DOI", "")
            papers.append({
                "title": p.get("title", ""), "authors": authors,
                "year": p.get("year", ""), "abstract": (p.get("abstract") or "")[:500],
                "doi": doi, "citations": p.get("citationCount", 0),
                "url": p.get("url", ""), "source": "Semantic Scholar",
            })
        return papers
    except Exception:
        return []


def _search_openalex(query, max_results=50):
    """Search OpenAlex API."""
    try:
        import urllib.parse
        q = urllib.parse.quote(query)
        url = f"https://api.openalex.org/works?search={q}&per_page={max_results}&sort=relevance_score:desc&select=title,authorships,publication_year,doi,cited_by_count,id"
        req = urllib.request.Request(url, headers={"User-Agent": "mailto:zenith@example.com"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode())
        papers = []
        for w in data.get("results", []):
            authors = [a.get("author", {}).get("display_name", "") for a in (w.get("authorships") or [])]
            doi = (w.get("doi") or "").replace("https://doi.org/", "")
            papers.append({
                "title": w.get("title", ""), "authors": authors,
                "year": w.get("publication_year", ""), "abstract": "",
                "doi": doi, "citations": w.get("cited_by_count", 0),
                "url": w.get("id", ""), "source": "OpenAlex",
            })
        return papers
    except Exception:
        return []


def _web_search_brave(query, api_key, max_results=100):
    """Web search via Brave Search API (https://brave.com/search/api/)."""
    try:
        encoded = urllib.parse.quote_plus(query)
        url = f"https://api.search.brave.com/res/v1/web/search?q={encoded}&count={max_results}"
        req = urllib.request.Request(url, headers={
            "Accept": "application/json",
            "Accept-Encoding": "gzip",
            "X-Subscription-Token": api_key,
        })
        with urllib.request.urlopen(req, timeout=20) as resp:
            raw = resp.read()
            # Handle gzip
            if resp.headers.get("Content-Encoding") == "gzip":
                import gzip
                raw = gzip.decompress(raw)
            data = json.loads(raw.decode("utf-8"))
        results = []
        for r in data.get("web", {}).get("results", []):
            results.append({
                "title": r.get("title", ""), "url": r.get("url", ""),
                "snippet": r.get("description", "")[:400], "score": 0,
                "source": "brave",
            })
        # Brave may include an infobox / summary
        answer = ""
        if data.get("summarizer", {}).get("key"):
            answer = data["summarizer"].get("summary", "")
        return {"results": results, "answer": answer}
    except Exception as e:
        return {"results": [], "answer": "", "error": str(e)}


def _firecrawl_scrape(url_to_scrape, api_key):
    """Scrape a URL via Firecrawl API (https://firecrawl.dev) → returns markdown text."""
    try:
        payload = json.dumps({
            "url": url_to_scrape,
            "formats": ["markdown"],
            "onlyMainContent": True,
            "timeout": 30000,
        }).encode()
        req = urllib.request.Request(
            "https://api.firecrawl.dev/v1/scrape",
            data=payload,
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        )
        with urllib.request.urlopen(req, timeout=40) as resp:
            data = json.loads(resp.read().decode())
        if data.get("success"):
            md = data.get("data", {}).get("markdown", "")
            title = data.get("data", {}).get("metadata", {}).get("title", "")
            return {"ok": True, "markdown": md[:8000], "title": title, "url": url_to_scrape}
        return {"ok": False, "error": data.get("error", "Unknown error")}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _firecrawl_search(query, api_key, max_results=50):
    """Search via Firecrawl Search API → returns URLs + scraped content."""
    try:
        payload = json.dumps({
            "query": query,
            "limit": max_results,
            "scrapeOptions": {"formats": ["markdown"], "onlyMainContent": True},
        }).encode()
        req = urllib.request.Request(
            "https://api.firecrawl.dev/v1/search",
            data=payload,
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        )
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = json.loads(resp.read().decode())
        results = []
        if data.get("success"):
            for r in data.get("data", []):
                results.append({
                    "title": r.get("metadata", {}).get("title", r.get("url", "")),
                    "url": r.get("url", ""),
                    "snippet": r.get("markdown", "")[:500],
                    "score": 0,
                    "source": "firecrawl",
                })
        return {"results": results}
    except Exception as e:
        return {"results": [], "error": str(e)}


def _web_search_tavily(query, api_key, max_results=50):
    """Web search via Tavily AI."""
    try:
        url = "https://api.tavily.com/search"
        payload = json.dumps({
            "api_key": api_key, "query": query, "max_results": max_results,
            "search_depth": "advanced", "include_answer": True,
        }).encode()
        req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode())
        results = []
        for r in data.get("results", []):
            results.append({
                "title": r.get("title", ""), "url": r.get("url", ""),
                "snippet": r.get("content", "")[:300], "score": r.get("score", 0),
                "source": "tavily",
            })
        return {"results": results, "answer": data.get("answer", "")}
    except Exception as e:
        return {"results": [], "answer": "", "error": str(e)}


def _web_search_duckduckgo(query, max_results=50):
    """Fallback web search via DuckDuckGo HTML (no API key needed).
    Uses ARC's pattern: separate regex for links/snippets + URL unwrapping."""
    try:
        encoded = urllib.parse.quote_plus(query)
        url = f"https://html.duckduckgo.com/html/?q={encoded}"
        req = urllib.request.Request(url, headers={"User-Agent": _USER_AGENT})
        with urllib.request.urlopen(req, timeout=15) as resp:
            html = resp.read().decode("utf-8", errors="replace")

        # Separate regex patterns (more robust, matches ARC)
        link_pattern = re.compile(r'<a[^>]*class="result__a"[^>]*href="([^"]*)"[^>]*>(.*?)</a>', re.DOTALL)
        snippet_pattern = re.compile(r'<a[^>]*class="result__snippet"[^>]*>(.*?)</a>', re.DOTALL)
        links = link_pattern.findall(html)
        snippets = snippet_pattern.findall(html)

        results = []
        for i, (raw_url, title_html) in enumerate(links[:max_results]):
            title = re.sub(r'<[^>]+>', '', title_html).strip()
            snippet = re.sub(r'<[^>]+>', '', snippets[i]).strip() if i < len(snippets) else ""
            # Unwrap DDG redirect URLs (ARC pattern)
            result_url = raw_url
            if "duckduckgo.com" in raw_url:
                parsed = urllib.parse.urlparse(raw_url)
                uddg = urllib.parse.parse_qs(parsed.query).get("uddg")
                if uddg:
                    result_url = urllib.parse.unquote(uddg[0])
                else:
                    continue
            results.append({"title": title, "url": result_url, "snippet": snippet[:300]})
        return {"results": results}
    except Exception:
        return {"results": []}


# ══════════════════════════════════════════════════════════════════════════════
# ██  V5.6 PIPELINE — PubMed, CrossRef, Sci-Hub, Unpaywall
# ══════════════════════════════════════════════════════════════════════════════

def _search_pubmed(query, max_results=200):
    """Search PubMed via NCBI E-utilities API.
    Returns list of paper dicts with title, authors, year, abstract, doi, pmid."""
    try:
        # Step 1: ESearch → get PMIDs
        encoded = urllib.parse.quote_plus(query)
        search_url = (
            f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
            f"?db=pubmed&term={encoded}&retmax={max_results}&retmode=json&sort=relevance"
        )
        req = urllib.request.Request(search_url, headers={"User-Agent": "Zenith/2.0 (mailto:zenith@example.com)"})
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode())

        pmids = data.get("esearchresult", {}).get("idlist", [])
        if not pmids:
            return []

        time.sleep(0.4)  # NCBI rate limit: 3 req/sec without key

        # Step 2: EFetch → get full metadata as XML
        ids_str = ",".join(pmids)
        fetch_url = (
            f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
            f"?db=pubmed&id={ids_str}&retmode=xml"
        )
        req2 = urllib.request.Request(fetch_url, headers={"User-Agent": "Zenith/2.0"})
        with urllib.request.urlopen(req2, timeout=30) as resp2:
            xml_data = resp2.read().decode("utf-8", errors="replace")

        import xml.etree.ElementTree as ET
        root = ET.fromstring(xml_data)

        papers = []
        for article in root.findall(".//PubmedArticle"):
            medline = article.find(".//MedlineCitation")
            if medline is None:
                continue
            art = medline.find("Article")
            if art is None:
                continue

            # Title
            title_el = art.find("ArticleTitle")
            title = "".join(title_el.itertext()).strip() if title_el is not None else ""

            # Abstract
            abstract_parts = []
            abstract_el = art.find("Abstract")
            if abstract_el is not None:
                for at in abstract_el.findall("AbstractText"):
                    label = at.get("Label", "")
                    text = "".join(at.itertext()).strip()
                    if label:
                        abstract_parts.append(f"{label}: {text}")
                    else:
                        abstract_parts.append(text)
            abstract = " ".join(abstract_parts)[:600]

            # Authors
            authors = []
            author_list = art.find("AuthorList")
            if author_list is not None:
                for au in author_list.findall("Author"):
                    last = au.findtext("LastName", "")
                    fore = au.findtext("ForeName", "")
                    if last:
                        authors.append(f"{last} {fore}".strip())

            # Year
            pub_date = art.find(".//PubDate")
            year = ""
            if pub_date is not None:
                year = pub_date.findtext("Year", "")
                if not year:
                    medline_date = pub_date.findtext("MedlineDate", "")
                    if medline_date:
                        ym = re.match(r'(\d{4})', medline_date)
                        year = ym.group(1) if ym else ""

            # DOI
            doi = ""
            for eid in art.findall(".//ELocationID"):
                if eid.get("EIdType") == "doi":
                    doi = (eid.text or "").strip()
                    break
            if not doi:
                article_ids = article.find(".//PubmedData/ArticleIdList")
                if article_ids is not None:
                    for aid in article_ids.findall("ArticleId"):
                        if aid.get("IdType") == "doi":
                            doi = (aid.text or "").strip()
                            break

            # PMID
            pmid = medline.findtext("PMID", "")

            # MeSH terms
            mesh_terms = []
            mesh_list = medline.find("MeshHeadingList")
            if mesh_list is not None:
                for mh in mesh_list.findall("MeshHeading"):
                    desc = mh.find("DescriptorName")
                    if desc is not None:
                        mesh_terms.append(desc.text or "")

            papers.append({
                "title": title,
                "authors": authors,
                "year": year,
                "abstract": abstract,
                "doi": doi,
                "citations": 0,
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else "",
                "source": "PubMed",
                "pmid": pmid,
                "mesh_terms": mesh_terms[:10],
            })

        return papers
    except Exception as e:
        return []


def _enrich_crossref(doi):
    """Enrich a paper record with CrossRef metadata (citation count, full metadata).
    Returns dict with enriched fields or empty dict on failure."""
    if not doi:
        return {}
    try:
        encoded = urllib.parse.quote(doi, safe='')
        url = f"https://api.crossref.org/works/{encoded}"
        req = urllib.request.Request(url, headers={
            "User-Agent": "Zenith/2.0 (mailto:zenith@example.com)",
            "Accept": "application/json",
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode())

        msg = data.get("message", {})
        citations = msg.get("is-referenced-by-count", 0)
        titles = msg.get("title", [])
        title = titles[0] if titles else ""

        authors = []
        for a in msg.get("author", []):
            name = f"{a.get('family', '')} {a.get('given', '')}".strip()
            if name:
                authors.append(name)

        year = ""
        pub_date = msg.get("published-print", msg.get("published-online", {}))
        date_parts = pub_date.get("date-parts", [[]])
        if date_parts and date_parts[0]:
            year = str(date_parts[0][0])

        journal = ""
        containers = msg.get("container-title", [])
        if containers:
            journal = containers[0]

        return {
            "citations": citations,
            "title": title,
            "authors": authors,
            "year": year,
            "journal": journal,
            "type": msg.get("type", ""),
            "subject": msg.get("subject", []),
        }
    except Exception:
        return {}


_SCIHUB_MIRRORS = [
    "https://sci-hub.ru",
    "https://sci-hub.st",
    "https://sci-hub.se",
    "https://sci-hub.su",
    "https://sci-hub.box",
    "https://sci-hub.red",
    "https://sci-hub.al",
    "https://sci-hub.mk",
    "https://sci-hub.ee",
    "https://sci-hub.in",
    "https://sci-hub.shop",
]


def _scihub_extract_pdf_url(html, mirror=""):
    """Extract PDF URL from Sci-Hub HTML using BeautifulSoup.
    Returns (pdf_url, None) or (None, captcha_info).
    Works with modern Sci-Hub (2024+) which uses <object> tags."""
    try:
        from scihub import SciHub as _SH
        return _SH.extract_pdf_url(html, mirror)
    except ImportError:
        pass

    # Fallback: inline BS4 extraction if scihub.py somehow unavailable
    from bs4 import BeautifulSoup as _BS
    soup = _BS(html, 'html.parser')

    # Check for CAPTCHA
    captcha_img = soup.find('img', id='captcha') or soup.find('img', src=re.compile(r'captcha', re.I))
    if captcha_img:
        captcha_form = captcha_img.find_parent('form')
        if captcha_form:
            action = captcha_form.get('action', '')
            img_src = captcha_img.get('src', '')
            return None, {"captcha_img_url": _scihub_fix_url(img_src, mirror),
                          "form_action": _scihub_fix_url(action, mirror) if action else ""}

    # 1. <object type="application/pdf" data="...">
    obj_tag = soup.find('object', attrs={'type': 'application/pdf'})
    if obj_tag and obj_tag.get('data'):
        return _scihub_fix_url(obj_tag['data'].split('#')[0], mirror), None

    # 2. <embed src="...">
    embed_tag = soup.find('embed', src=True)
    if embed_tag:
        url = embed_tag['src'].split('#')[0]
        if '.pdf' in url or '/storage/' in url:
            return _scihub_fix_url(url, mirror), None

    # 3. <iframe src="...">
    iframe_tag = soup.find('iframe', src=True)
    if iframe_tag:
        return _scihub_fix_url(iframe_tag['src'].split('#')[0], mirror), None

    # 4. JS inline: url: '/storage/...'
    for script in soup.find_all('script'):
        text = script.string or ''
        m = re.search(r'''['"]?url['"]?\s*:\s*['"]([^'"]+\.pdf[^'"]*)['"]''', text)
        if m:
            return _scihub_fix_url(m.group(1).split('#')[0], mirror), None

    # 5. Direct <a> link to .pdf
    for a_tag in soup.find_all('a', href=True):
        href = a_tag['href']
        if '.pdf' in href and ('storage' in href or 'download' in href):
            return _scihub_fix_url(href, mirror), None

    # 6. Regex fallback
    m = re.search(r'(https?://[^\s"\'<>]+\.pdf)', html)
    if m:
        return m.group(1), None

    return None, None


def _scihub_fix_url(url, mirror):
    """Fix relative URLs to absolute."""
    if url.startswith("//"):
        return "https:" + url
    elif url.startswith("/"):
        return mirror + url
    elif not url.startswith("http"):
        return "https://" + url
    return url


def _scihub_download_pdf(pdf_url, doi, output_dir):
    """Download PDF from URL, validate, and save. Returns result dict."""
    headers = {"User-Agent": _USER_AGENT}

    if _HAS_REQUESTS:
        resp = requests.get(pdf_url, headers=headers, timeout=45, verify=False)
        pdf_bytes = resp.content
    else:
        req = urllib.request.Request(pdf_url, headers=headers)
        import ssl
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        with urllib.request.urlopen(req, timeout=45, context=ctx) as resp:
            pdf_bytes = resp.read()

    if len(pdf_bytes) < 1024 or pdf_bytes[:5] != b'%PDF-':
        return {"ok": False, "error": "Downloaded content is not a valid PDF"}

    safe_doi = re.sub(r'[^a-zA-Z0-9_.-]', '_', doi)
    out_path = os.path.join(output_dir, f"scihub_{safe_doi}.pdf")
    with open(out_path, "wb") as f:
        f.write(pdf_bytes)
    return {"ok": True, "path": out_path, "url": pdf_url, "size": len(pdf_bytes)}


def _fetch_scihub(doi, output_dir=None):
    """Fetch a paper PDF from Sci-Hub with CAPTCHA detection.
    Returns:
      {ok: True, path, url, mirror, size} — success
      {ok: False, captcha_required: True, captcha_img_b64, mirror, doi, cookies_b64} — needs CAPTCHA
      {ok: False, error} — failure
    """
    if not doi:
        return {"ok": False, "error": "No DOI provided"}
    if output_dir is None:
        output_dir = PAPERS_DIR
    os.makedirs(output_dir, exist_ok=True)

    # Strategy 1: Use our rewritten scihub.py (BS4-based, handles <object> tags)
    if _HAS_SCIHUB and _HAS_REQUESTS:
        try:
            sh = SciHub(mirrors=_SCIHUB_MIRRORS, timeout=30)
            result = sh.fetch(doi)
            if result and result.get('pdf'):
                pdf_bytes = result['pdf']
                valid, _ = sh.validate_pdf(pdf_bytes)
                if valid:
                    safe_doi = re.sub(r'[^a-zA-Z0-9_.-]', '_', doi)
                    out_path = os.path.join(output_dir, f"scihub_{safe_doi}.pdf")
                    with open(out_path, "wb") as f:
                        f.write(pdf_bytes)
                    return {"ok": True, "path": out_path, "url": result.get('url', ''),
                            "mirror": result.get('mirror', ''), "size": len(pdf_bytes)}
            elif result and result.get('err'):
                # SciHub class tried all mirrors — fall through to manual approach
                pass
        except Exception:
            pass

    # Strategy 2: Manual mirror crawl with CAPTCHA handling
    import base64
    session = requests.Session() if _HAS_REQUESTS else None

    for mirror in _SCIHUB_MIRRORS:
        try:
            page_url = f"{mirror}/{doi}"
            headers = {
                "User-Agent": _USER_AGENT,
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            }

            if session:
                resp = session.get(page_url, headers=headers, timeout=20, verify=False)
                if resp.status_code == 403 or resp.status_code >= 500:
                    continue
                html = resp.text
            else:
                req = urllib.request.Request(page_url, headers=headers)
                with urllib.request.urlopen(req, timeout=20) as resp:
                    html = resp.read().decode("utf-8", errors="replace")

            if len(html) < 200 or 'article not found' in html.lower():
                continue

            pdf_url, captcha_info = _scihub_extract_pdf_url(html, mirror)

            # CAPTCHA detected — return image for user to solve
            if captcha_info:
                captcha_img_url = captcha_info.get("captcha_img_url", "")
                try:
                    if session:
                        img_resp = session.get(captcha_img_url, headers=headers, timeout=10, verify=False)
                        img_bytes = img_resp.content
                    else:
                        img_req = urllib.request.Request(captcha_img_url, headers=headers)
                        with urllib.request.urlopen(img_req, timeout=10) as img_resp:
                            img_bytes = img_resp.read()
                    captcha_b64 = base64.b64encode(img_bytes).decode("ascii")
                except Exception:
                    captcha_b64 = ""

                cookies_dict = {}
                if session and session.cookies:
                    cookies_dict = dict(session.cookies)

                return {
                    "ok": False,
                    "captcha_required": True,
                    "captcha_img_b64": captcha_b64,
                    "mirror": mirror,
                    "doi": doi,
                    "form_action": captcha_info.get("form_action", ""),
                    "cookies": cookies_dict,
                }

            if not pdf_url:
                continue

            dl = _scihub_download_pdf(pdf_url, doi, output_dir)
            if dl.get("ok"):
                dl["mirror"] = mirror
                return dl

        except Exception:
            continue

    return {"ok": False, "error": f"Could not fetch DOI {doi} from any Sci-Hub mirror"}


def solve_scihub_captcha(args):
    """Submit CAPTCHA solution to Sci-Hub and download the paper.
    Args: {solution, mirror, doi, form_action, cookies, output_dir}
    Returns: {ok, path, url, mirror, size} or {ok: False, error}"""
    solution = args.get("solution", "")
    mirror = args.get("mirror", "")
    doi = args.get("doi", "")
    form_action = args.get("form_action", "")
    cookies = args.get("cookies", {})
    output_dir = args.get("output_dir", PAPERS_DIR)

    if not solution or not mirror or not doi:
        return {"error": "Missing captcha solution, mirror, or DOI"}
    if not _HAS_REQUESTS:
        return {"error": "requests package required for CAPTCHA solving"}

    os.makedirs(output_dir, exist_ok=True)

    try:
        session = requests.Session()
        if cookies:
            session.cookies.update(cookies)

        # Submit CAPTCHA solution
        action_url = form_action if form_action.startswith("http") else f"{mirror}{form_action or '/'}"
        resp = session.post(action_url, data={
            "answer": solution,
            "id": doi,
        }, headers={"User-Agent": _USER_AGENT}, timeout=20, verify=False, allow_redirects=True)

        html = resp.text

        # Check if CAPTCHA was solved — look for PDF URL
        pdf_url, captcha_info = _scihub_extract_pdf_url(html)

        if captcha_info:
            return {"ok": False, "error": "CAPTCHA solution was incorrect, please try again",
                    "captcha_required": True}

        if not pdf_url:
            # Maybe we got redirected to the PDF directly
            if resp.headers.get("Content-Type", "").startswith("application/pdf"):
                pdf_bytes = resp.content
                if len(pdf_bytes) >= 1024 and pdf_bytes[:5] == b'%PDF-':
                    safe_doi = re.sub(r'[^a-zA-Z0-9_.-]', '_', doi)
                    out_path = os.path.join(output_dir, f"scihub_{safe_doi}.pdf")
                    with open(out_path, "wb") as f:
                        f.write(pdf_bytes)
                    return {"ok": True, "path": out_path, "url": resp.url,
                            "mirror": mirror, "size": len(pdf_bytes)}
            return {"ok": False, "error": "Could not find PDF after CAPTCHA submission"}

        pdf_url = _scihub_fix_url(pdf_url, mirror)
        dl = _scihub_download_pdf(pdf_url, doi, output_dir)
        if dl.get("ok"):
            dl["mirror"] = mirror
        return dl

    except Exception as e:
        return {"ok": False, "error": f"CAPTCHA solve error: {str(e)}"}

def _fetch_unpaywall(doi, output_dir=None):
    """Attempt to fetch open-access PDF via Unpaywall API.
    Returns {ok, path, url} or {ok: False, error}."""
    if not doi:
        return {"ok": False, "error": "No DOI provided"}

    if output_dir is None:
        output_dir = PAPERS_DIR
    os.makedirs(output_dir, exist_ok=True)

    try:
        encoded = urllib.parse.quote(doi, safe='/')
        url = f"https://api.unpaywall.org/v2/{encoded}?email=zenith@example.com"
        req = urllib.request.Request(url, headers={"User-Agent": "Zenith/2.0"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode())

        # Find best OA location with PDF
        best_loc = data.get("best_oa_location") or {}
        pdf_url = best_loc.get("url_for_pdf", "")

        if not pdf_url:
            # Try all OA locations
            for loc in data.get("oa_locations", []):
                if loc.get("url_for_pdf"):
                    pdf_url = loc["url_for_pdf"]
                    break

        if not pdf_url:
            return {"ok": False, "error": "No open-access PDF available via Unpaywall"}

        # Download PDF
        pdf_req = urllib.request.Request(pdf_url, headers={"User-Agent": _USER_AGENT})
        with urllib.request.urlopen(pdf_req, timeout=45) as pdf_resp:
            pdf_bytes = pdf_resp.read()

        if len(pdf_bytes) < 1024:
            return {"ok": False, "error": "Downloaded file too small, likely not a valid PDF"}

        safe_doi = re.sub(r'[^a-zA-Z0-9_.-]', '_', doi)
        filename = f"oa_{safe_doi}.pdf"
        out_path = os.path.join(output_dir, filename)
        with open(out_path, "wb") as f:
            f.write(pdf_bytes)

        return {"ok": True, "path": out_path, "url": pdf_url, "source": "unpaywall", "size": len(pdf_bytes)}

    except Exception as e:
        return {"ok": False, "error": str(e)}


# ══════════════════════════════════════════════════════════════════════════════
# ██  V5.6 PIPELINE — Gatekeeper, Query Architect, Triage, Acquire
# ══════════════════════════════════════════════════════════════════════════════

## ── Structured Output Schemas (Gemini native, prompt-injected for others) ──

_SCHEMA_GATEKEEPER = {
    "type": "object",
    "properties": {
        "is_valid": {"type": "boolean", "description": "Whether the research question is answerable via peer-reviewed literature"},
        "reason": {"type": "string", "description": "Explanation of validity assessment"},
        "domain": {"type": "string", "description": "Research domain/field (e.g. medicine, computer science)"},
        "keywords": {"type": "array", "items": {"type": "string"}, "description": "Key search terms extracted from the question"},
        "suggested_refinement": {"type": "string", "description": "Improved version of the question, or empty string if already good"},
    },
    "required": ["is_valid", "reason", "domain"],
}

_SCHEMA_QUERIES = {
    "type": "array",
    "items": {
        "type": "object",
        "properties": {
            "db": {"type": "string", "description": "Target database: pubmed, semantic_scholar, openalex, arxiv, web"},
            "query_string": {"type": "string", "description": "The actual search query/Boolean string"},
            "description": {"type": "string", "description": "Brief description of what this query targets"},
        },
        "required": ["db", "query_string", "description"],
    },
}

_SCHEMA_TRIAGE = {
    "type": "array",
    "items": {
        "type": "object",
        "properties": {
            "paper_index": {"type": "integer", "description": "1-based index of the paper in the batch"},
            "is_relevant": {"type": "boolean", "description": "Whether paper is relevant to the research question"},
            "relevance_score": {"type": "number", "description": "Relevance score from 0.0 to 1.0"},
            "justification": {"type": "string", "description": "1-2 sentence justification for the relevance decision"},
        },
        "required": ["paper_index", "is_relevant", "relevance_score"],
    },
}

_SCHEMA_BLUEPRINT = {
    "type": "array",
    "items": {
        "type": "object",
        "properties": {
            "section": {"type": "string", "description": "Section name (e.g. Introduction, Methods, Results, Discussion)"},
            "requirements": {"type": "array", "items": {"type": "string"}, "description": "What must be covered in this section"},
            "subsections": {"type": "array", "items": {"type": "string"}, "description": "Subsection headings within this section"},
            "needs_table": {"type": "boolean", "description": "Whether this section needs a comparison/summary table"},
            "needs_figure": {"type": "boolean", "description": "Whether this section needs a figure or chart"},
            "word_target": {"type": "integer", "description": "Target word count for this section"},
        },
        "required": ["section", "requirements"],
    },
}

_SCHEMA_CITATION_VERIFY = {
    "type": "object",
    "properties": {
        "verified": {"type": "array", "items": {"type": "object", "properties": {"citation": {"type": "string"}, "paper_index": {"type": "integer"}, "accurate": {"type": "boolean"}}, "required": ["citation", "accurate"]}, "description": "List of verified citations"},
        "hallucinated": {"type": "array", "items": {"type": "string"}, "description": "Citations that don't match any paper in the reference list"},
        "issues": {"type": "array", "items": {"type": "string"}, "description": "Specific issues found with citations"},
        "pass": {"type": "boolean", "description": "Whether the section passes citation verification"},
    },
    "required": ["verified", "hallucinated", "pass"],
}

_SCHEMA_GUIDELINES = {
    "type": "object",
    "properties": {
        "compliant": {"type": "array", "items": {"type": "string"}, "description": "Checklist items that are satisfied"},
        "violations": {"type": "array", "items": {"type": "object", "properties": {"item": {"type": "string"}, "severity": {"type": "string"}, "suggestion": {"type": "string"}}, "required": ["item", "severity"]}, "description": "Checklist items that are violated"},
        "pass": {"type": "boolean", "description": "Whether the section passes guidelines compliance"},
    },
    "required": ["compliant", "violations", "pass"],
}


def _get_step_config(args, defaults=None):
    """Extract step_config from args with fallback defaults."""
    sc = args.get("step_config", {})
    d = defaults or {}
    return {
        "system_prompt": sc.get("system_prompt", d.get("system_prompt", "")),
        "model_tier": sc.get("model_tier", d.get("model_tier", "strong")),
        "max_tokens": sc.get("max_tokens", d.get("max_tokens", 4096)),
        "temperature": sc.get("temperature", d.get("temperature", 0.3)),
        "use_structured_output": sc.get("use_structured_output", d.get("use_structured_output", False)),
        "use_thinking": sc.get("use_thinking", d.get("use_thinking", False)),
        "thinking_budget": sc.get("thinking_budget", d.get("thinking_budget", 8192)),
        "enabled_tools": sc.get("enabled_tools", d.get("enabled_tools", [])),
    }

def _build_system_prompt(args, sc, fallback, **kwargs):
    """Combine the overall base system prompt with the step-specific prompt, replacing placeholders."""
    base_prompt = args.get("system_prompt", "")
    step_prompt = sc["system_prompt"] if sc.get("system_prompt") else fallback

    study_design = args.get("study_design", "systematic_review")
    guidelines_map = {
        "systematic_review": "PRISMA 2020",
        "meta_analysis": "PRISMA-MA + MOOSE",
        "narrative_review": "SANRA",
        "scoping_review": "PRISMA-ScR",
        "subject_review": "Standard academic review",
        "educational": "Pedagogical resource",
        "case_study": "CARE",
        "comparative": "Comparative analysis",
        "exploratory": "Exploratory research",
    }
    guidelines = args.get("guidelines", guidelines_map.get(study_design, "academic"))

    fmt_args = {
        "study_design": study_design.replace("_", " "),
        "query": args.get("query", ""),
        "guidelines": guidelines,
    }
    for k, v in kwargs.items():
        fmt_args[k] = str(v)

    def _replace(text):
        if not text: return text
        for k, v in fmt_args.items():
            text = text.replace(f"{{{k}}}", str(v))
        return text

    base_prompt = _replace(base_prompt)
    step_prompt = _replace(step_prompt)

    if base_prompt and base_prompt.strip() and base_prompt != step_prompt:
        return f"{base_prompt.strip()}\n\n[Task-Specific Instructions]\n{step_prompt.strip()}"
    return step_prompt

def validate_research_query(args):
    """Phase 1.1 — The Gatekeeper: Validate if a research question is answerable.
    Args: {query, api_key, provider, model, step_config}
    Returns: {ok, is_valid, reason, domain, suggested_refinement}"""
    query = args.get("query", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "strong", "max_tokens": 2048, "temperature": 0.1, "use_structured_output": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}
    if not query.strip():
        return {"error": "Research question is required."}

    system = _build_system_prompt(args, sc, 
        "You are a research methodology expert. Evaluate whether the given research question "
        "is suitable for systematic academic literature review. Consider: specificity, scope, "
        "feasibility, and whether peer-reviewed literature likely exists on this topic."
    )

    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": system},
                        {"role": "user", "content": f"Research question: {query}"}],
                       temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                       response_schema=_SCHEMA_GATEKEEPER if sc["use_structured_output"] else None)
    if "error" in result:
        return result

    # Use structured output if available, else parse from text
    parsed = result.get("structured")
    if not parsed:
        try:
            parsed = json.loads(re.search(r'\{.*\}', result["text"], re.DOTALL).group())
        except Exception:
            parsed = {"is_valid": True, "reason": result["text"][:500], "domain": "general"}

    return {"ok": True, **parsed, "tokens": result.get("usage")}


def generate_search_queries(args):
    """Phase 1.2 — The Query Architect: Generate MeSH/Boolean search strings.
    Args: {query, domain, api_key, provider, model, step_config}
    Returns: {ok, queries: [{db, query_string, description}]}"""
    query = args.get("query", "")
    domain = args.get("domain", "biomedical")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "strong", "max_tokens": 4096, "temperature": 0.3, "use_structured_output": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}

    system = _build_system_prompt(args, sc, 
        "You are a medical librarian expert in MeSH terminology and Boolean search strategy. "
        "Given a research question, generate optimized search queries for:\n"
        "1. PubMed (using MeSH terms and Boolean operators AND/OR/NOT)\n"
        "2. Semantic Scholar / OpenAlex (natural language, key concepts)\n"
        "3. arXiv (for computational/AI-related aspects)\n"
        "4. Web search (for grey literature, preprints)\n\n"
        "Generate at least 3 complementary queries targeting different facets of the question."
    )

    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": system},
                        {"role": "user", "content": f"Research question: {query}\nDomain: {domain}"}],
                       temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                       response_schema=_SCHEMA_QUERIES if sc["use_structured_output"] else None)
    if "error" in result:
        return result

    parsed = result.get("structured")
    if not parsed:
        try:
            parsed = json.loads(re.search(r'\[.*\]', result["text"], re.DOTALL).group())
        except Exception:
            parsed = [{"db": "pubmed", "query_string": query, "description": "Direct search"}]

    return {"ok": True, "queries": parsed if isinstance(parsed, list) else [parsed],
            "tokens": result.get("usage")}


def triage_papers(args):
    """Phase 1.4 — The Triage Agent: Screen papers for relevance.
    Args: {papers: [{title, abstract, doi}], query, api_key, provider, model, step_config}
    Returns: {ok, results: [{doi, is_relevant, justification, relevance_score}]}"""
    papers = args.get("papers", [])
    query = args.get("query", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "fast", "max_tokens": 8192, "temperature": 0.1, "use_structured_output": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}
    if not papers:
        return {"ok": True, "results": [], "relevant_count": 0}

    all_results = []
    batch_size = 10

    for batch_start in range(0, len(papers), batch_size):
        batch = papers[batch_start:batch_start + batch_size]
        papers_text = ""
        for idx, p in enumerate(batch):
            papers_text += (
                f"\n--- Paper {idx + 1} ---\n"
                f"Title: {p.get('title', 'N/A')}\n"
                f"Abstract: {p.get('abstract', 'N/A')[:400]}\n"
                f"DOI: {p.get('doi', 'N/A')}\n"
            )

        system = _build_system_prompt(args, sc, 
            "You are an expert research screener. Evaluate each paper's relevance to the "
            "research question. Apply strict inclusion/exclusion criteria. "
            "For each paper, assess: direct relevance, study design appropriateness, "
            "publication quality, and methodological rigor."
        )

        result = _llm_chat(provider, api_key, model,
                           [{"role": "system", "content": system},
                            {"role": "user", "content": f"Research question: {query}\n\nPapers to screen:{papers_text}"}],
                           temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                           response_schema=_SCHEMA_TRIAGE if sc["use_structured_output"] else None)
        if "error" in result:
            for p in batch:
                all_results.append({"doi": p.get("doi", ""), "title": p.get("title", ""),
                                    "is_relevant": True, "relevance_score": 0.5,
                                    "justification": "Screening failed, included by default"})
            continue

        try:
            arr = result.get("structured") or json.loads(re.search(r'\[.*\]', result["text"], re.DOTALL).group())
            if not isinstance(arr, list):
                arr = [arr]
            for item in arr:
                pi = item.get("paper_index", 1) - 1
                if 0 <= pi < len(batch):
                    all_results.append({
                        "doi": batch[pi].get("doi", ""),
                        "title": batch[pi].get("title", ""),
                        "is_relevant": item.get("is_relevant", True),
                        "relevance_score": item.get("relevance_score", 0.5),
                        "justification": item.get("justification", ""),
                    })
        except Exception:
            for p in batch:
                all_results.append({"doi": p.get("doi", ""), "title": p.get("title", ""),
                                    "is_relevant": True, "relevance_score": 0.5,
                                    "justification": "Could not parse screening result"})

        if batch_start + batch_size < len(papers):
            time.sleep(0.5)

    relevant_count = sum(1 for r in all_results if r.get("is_relevant"))
    return {"ok": True, "results": all_results, "relevant_count": relevant_count,
            "total_screened": len(all_results)}


def acquire_papers(args):
    """Phase 1.5 — The Acquisition Engine: Download papers via Sci-Hub + Unpaywall.
    Args: {papers: [{doi, title}], output_dir (optional), skip_unpaywall (bool)}
    Returns: {ok, acquired, failed, captcha_needed}"""
    papers = args.get("papers", [])
    output_dir = args.get("output_dir", PAPERS_DIR)
    skip_unpaywall = args.get("skip_unpaywall", False)
    os.makedirs(output_dir, exist_ok=True)

    acquired = []
    failed = []
    captcha_needed = []

    for p in papers[:50]:
        doi = p.get("doi", "")
        pmid = p.get("pmid", "")
        title = p.get("title", "unknown")

        # Check if already downloaded (skip re-downloading)
        if doi:
            safe_doi = re.sub(r'[^a-zA-Z0-9_.-]', '_', doi)
            existing = [f for f in os.listdir(output_dir) if safe_doi in f and f.endswith(".pdf")] if os.path.isdir(output_dir) else []
            if existing:
                existing_path = os.path.join(output_dir, existing[0])
                if os.path.getsize(existing_path) > 1024:
                    acquired.append({"doi": doi, "title": title, "path": existing_path,
                                     "source": "cached", "size": os.path.getsize(existing_path)})
                    continue

        # Build list of identifiers to try: DOI first, then PMID, then title
        identifiers = []
        if doi:
            identifiers.append(("doi", doi))
        if pmid:
            identifiers.append(("pmid", str(pmid)))
        if title and len(title) > 10:
            identifiers.append(("title", title))

        if not identifiers:
            failed.append({"doi": "", "title": title, "error": "No DOI, PMID, or title available"})
            continue

        paper_acquired = False
        last_result = None

        for id_type, identifier in identifiers:
            # Try Unpaywall first (legal, open access) — only works with DOI
            if id_type == "doi" and not skip_unpaywall:
                result = _fetch_unpaywall(identifier, output_dir)
                if result.get("ok"):
                    acquired.append({"doi": doi, "title": title, "path": result["path"],
                                     "source": "unpaywall", "size": result.get("size", 0)})
                    paper_acquired = True
                    time.sleep(0.3)
                    break

            # Sci-Hub (accepts DOI, PMID, or title search)
            time.sleep(0.5)
            result = _fetch_scihub(identifier, output_dir)
            last_result = result
            if result.get("ok"):
                acquired.append({"doi": doi, "title": title, "path": result["path"],
                                 "source": f"scihub_{id_type}", "size": result.get("size", 0)})
                paper_acquired = True
                time.sleep(1.0)
                break

            # CAPTCHA detected — collect for interactive solving (only try once)
            if result.get("captcha_required"):
                captcha_needed.append({
                    "doi": doi, "title": title,
                    "captcha_img_b64": result.get("captcha_img_b64", ""),
                    "mirror": result.get("mirror", ""),
                    "form_action": result.get("form_action", ""),
                    "cookies": result.get("cookies", {}),
                })
                paper_acquired = True  # not failed, just pending CAPTCHA
                break

        if not paper_acquired:
            failed.append({"doi": doi, "title": title, "error": (last_result or {}).get("error", "Download failed with all identifiers")})
            time.sleep(0.5)

    return {"ok": True, "acquired": acquired, "failed": failed,
            "captcha_needed": captcha_needed,
            "acquired_count": len(acquired), "failed_count": len(failed)}


def draft_research_section(args):
    """Phase 3.2 — Lead Author Agent: Draft a research section with citations.
    Args: {section_type, query, papers_context, api_key, provider, model, guidelines, step_config, blueprint_requirements}
    Returns: {ok, text, citations_used, tokens, chart_requests, table_requests}"""
    section_type = args.get("section_type", "introduction")
    query = args.get("query", "")
    papers_context = args.get("papers_context", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    guidelines = args.get("guidelines", "PRISMA")
    blueprint_reqs = args.get("blueprint_requirements", [])
    extracted_texts = args.get("extracted_texts", "")
    project_id = args.get("project_id", "default")
    sc = _get_step_config(args, {"model_tier": "strong", "max_tokens": 32768, "temperature": 0.5, "use_thinking": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}

    # Phase 3.2 Retrieval Agent — query vector DB for section-specific context
    vector_context = ""
    if _HAS_CHROMADB:
        vdb_result = query_vectordb({"project_id": project_id, "query": query,
                                     "section_type": section_type, "n_results": 8})
        if vdb_result.get("results"):
            vector_context = "\n\n## Retrieved full-text excerpts (from vector DB):\n"
            for vi, vr in enumerate(vdb_result["results"], 1):
                vector_context += f"\n[VDB-{vi}] (Source: {vr.get('title', 'Unknown')}, score: {vr.get('score', 0)})\n{vr['text'][:800]}\n"

    system = _build_system_prompt(args, sc, 
        "You are an expert academic researcher writing a {guidelines}-compliant "
        "research paper. You are drafting the '{section_type}' section.\n\n"
        "Rules:\n"
        "1. Use formal academic prose with objective, precise language.\n"
        "2. Synthesize evidence from the provided papers. Do not hallucinate claims.\n"
        "3. Every factual claim MUST be followed by an inline numbered citation [N].\n"
        "4. If blueprint_requirements are provided, ensure they are strictly covered.\n"
        "5. Recommend locations for tables/figures if data allows (e.g. '[Insert Table comparing outcomes]').",
        section_type=section_type, guidelines=guidelines
    )

    reqs_text = ""
    if blueprint_reqs:
        reqs_text = f"\n\nSection requirements from blueprint:\n- " + "\n- ".join(blueprint_reqs)

    extra_context = ""
    if extracted_texts:
        extra_context += f"\n\nFull-text excerpts from acquired PDFs:\n{extracted_texts[:6000]}"
    if vector_context:
        extra_context += vector_context

    prompt = (
        f"Research question: {query}\n\n"
        f"Available literature:\n{papers_context[:12000]}{extra_context}\n\n"
        f"Write the {section_type} section. Be comprehensive and cite all sources.{reqs_text}"
    )

    enabled_tools = sc.get("enabled_tools", [])
    tool_desc = ""
    code_execution = False
    if enabled_tools:
        tools = []
        if "experiment" in enabled_tools:
            tools.append("- EXPERIMENT: Run Python code in a sandboxed environment")
            code_execution = True
        if "generate_chart" in enabled_tools:
            tools.append("- GENERATE_CHART: Generate a chart (bar, line, pie, scatter, heatmap) from data")
            code_execution = True
        if "generate_table" in enabled_tools:
            tools.append("- GENERATE_TABLE: Generate a formatted data table")
        if tools:
            tool_desc = (
                "\n\nYou have access to these research tools:\n" +
                "\n".join(tools) +
                "\n\nWhen you need to use a tool, include a tool call tag in your response like: "
                "[TOOL:TOOL_NAME]{\"param\": \"value\"}[/TOOL]\n"
                "You can use multiple tools in a single response. Always explain what you found after using a tool."
            )

    full_system = system + tool_desc

    thinking = {"budget": sc.get("thinking_budget", 8192)} if sc["use_thinking"] and provider == "google" else None
    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": full_system},
                        {"role": "user", "content": prompt}],
                       temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                       thinking_config=thinking, code_execution=code_execution)
    if "error" in result:
        return result

    text = result["text"]
    citations = re.findall(r'\[([^\]]+)\]', text)

    # Extract chart/figure requests from the text
    figure_requests = re.findall(r'\[FIGURE:\s*(.+?)\]', text)
    # Extract any markdown tables inline in the text
    table_matches = re.findall(r'(\|.+\|(?:\n\|.+\|)+)', text)

    return {"ok": True, "text": text, "section_type": section_type,
            "citations_used": len(citations), "tokens": result.get("usage"),
            "figures": figure_requests, "tables": [t.strip() for t in table_matches]}


def smooth_manuscript(args):
    """Phase 4.1 — The Smoothing Pass: Polish and harmonize the full manuscript.
    Args: {sections: [{type, text}], query, api_key, provider, model, papers, step_config,
           generated_figures: [{description, caption, path, index}],
           generated_tables: [{description, caption, markdown, path, index}]}
    Returns: {ok, manuscript, abstract, references, tokens}"""
    sections = args.get("sections", [])
    query = args.get("query", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    papers = args.get("papers", [])
    generated_figures = args.get("generated_figures", [])
    generated_tables = args.get("generated_tables", [])
    project_id = args.get("project_id", "default")
    sc = _get_step_config(args, {"model_tier": "strong", "max_tokens": 65536, "temperature": 0.4, "use_thinking": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}
    if not sections:
        return {"error": "No sections provided."}

    # Retrieve vector DB context for richer evidence during smoothing
    vector_context = ""
    if _HAS_CHROMADB:
        vdb_result = query_vectordb({"project_id": project_id, "query": query, "n_results": 12})
        if vdb_result.get("results"):
            vector_context = "\n\n## Full-text evidence from vector DB (use to verify claims):\n"
            for vi, vr in enumerate(vdb_result["results"], 1):
                vector_context += f"\n[VDB-{vi}] ({vr.get('title', 'Unknown')})\n{vr['text'][:600]}\n"

    # Build figure + table lookup for post-LLM injection
    # We do NOT replace figures before sending to the LLM — the LLM would remove/rewrite them.
    # Instead, we strip all figure placeholders to clean tags, send to LLM, then replace after.
    figure_list = list(generated_figures)
    table_list = list(generated_tables)

    # Normalize all placeholders to a clean [FIGURE_N] format before sending to LLM
    combined = ""
    fig_counter = 0
    fig_tag_map = {}  # tag -> figure data
    for s in sections:
        text = s.get("text", "")
        # Match multiple placeholder formats the LLM might use
        placeholder_patterns = [
            r'\[FIGURE:\s*(.+?)\]',                                          # [FIGURE: desc]
            r'\((?:Suggest|Insert|Place)\s+(?:placing\s+)?Figure\s*\d*\s*(?:here)?[:\s]*(.+?)\)',  # (Suggest placing Figure 1 here: ...)
            r'\[(?:Insert|Place)\s+Figure\s*\d*\s*(?:here)?[:\s]*(.+?)\]',   # [Insert Figure 1 here: ...]
            r'\(Placement Suggestion:\s*(.+?)\)',                             # (Placement Suggestion: ...)
        ]
        for pattern in placeholder_patterns:
            def _replace_to_tag(m, _pat=pattern):
                nonlocal fig_counter
                desc = m.group(1).strip().rstrip('.')
                fig_counter += 1
                tag = f"[FIGURE_{fig_counter}]"
                # Try to match this description to a generated figure
                matched_fig = _fuzzy_match_figure(desc, figure_list, fig_tag_map)
                fig_tag_map[tag] = {"desc": desc, "figure": matched_fig, "index": fig_counter}
                return f"\n\n{tag}\n\n"
            text = re.sub(pattern, _replace_to_tag, text, flags=re.IGNORECASE)
        combined += f"\n\n## {s.get('type', 'Section').title()}\n\n{text}"

    # Assign remaining unmatched figures to their own tags at the end
    used_figs = {id(entry["figure"]) for entry in fig_tag_map.values() if entry["figure"]}
    unmatched_figs = [f for f in figure_list if id(f) not in used_figs]
    for uf in unmatched_figs:
        fig_counter += 1
        tag = f"[FIGURE_{fig_counter}]"
        fig_tag_map[tag] = {"desc": uf.get("description", ""), "figure": uf, "index": fig_counter}
        combined += f"\n\n{tag}\n\n"

    # Append generated tables inline with clean tags
    tbl_tag_map = {}
    if table_list:
        for ti, tbl in enumerate(table_list):
            tag = f"[TABLE_{ti + 1}]"
            tbl_tag_map[tag] = tbl
            combined += f"\n\n{tag}\n\n"

    refs_context = ""
    if papers:
        refs_context = "\n\n## Available References\n"
        for i, p in enumerate(papers, 1):
            auth = ", ".join(p.get("authors", [])[:3])
            if len(p.get("authors", [])) > 3:
                auth += " et al."
            refs_context += f"[{i}] {auth}. \"{p.get('title', '')}\" ({p.get('year', 'n.d.')}). {p.get('journal', '')}. DOI: {p.get('doi', 'N/A')}\n"

    system = _build_system_prompt(args, sc, 
        "You are a senior research editor performing a final smoothing pass on an academic manuscript. "
        "Your tasks:\n"
        "1. Fix tonal inconsistencies between sections\n"
        "2. Add logical transition sentences between sections\n"
        "3. Eliminate redundancy and tighten prose\n"
        "4. Write the Abstract (structured: Background, Methods, Results, Conclusions)\n"
        "5. Write the Conclusion section summarizing main findings and implications\n"
        "6. Preserve ALL existing numbered citations [N] exactly as they are. DO NOT REMOVE OR RENUMBER citations.\n"
        "7. Append a formal References section at the very end using the provided Reference List."
    )

    # Incorporate quality swarm feedback
    quality_context = ""
    citation_issues = args.get("citation_issues", [])
    guidelines_issues = args.get("guidelines_issues", [])
    if citation_issues:
        quality_context += "\n\n## Citation Issues to Fix\n"
        for ci in citation_issues[:20]:
            quality_context += f"- [{ci.get('severity', 'info')}] {ci.get('section', '?')}: {ci.get('issue', '')}\n"
    if guidelines_issues:
        quality_context += "\n\n## Guidelines Compliance Issues to Address\n"
        for gi in guidelines_issues[:20]:
            quality_context += f"- [{gi.get('status', '?')}] {gi.get('item', '')}: {gi.get('fix', '')}\n"

    thinking = {"budget": sc.get("thinking_budget", 16384)} if sc["use_thinking"] and provider == "google" else None
    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": system},
                        {"role": "user", "content": f"Research question: {query}\n\nDraft manuscript:{combined}{refs_context}{vector_context}{quality_context}"}],
                       temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                       thinking_config=thinking)
    if "error" in result:
        return result

    manuscript = result["text"]

    # ── Post-LLM: Replace [FIGURE_N] and [TABLE_N] tags with actual markdown images ──
    real_fig_num = 0
    for tag, info in fig_tag_map.items():
        fig = info.get("figure")
        if fig and fig.get("path") and os.path.isfile(fig["path"]):
            real_fig_num += 1
            caption = fig.get("caption", info["desc"])
            filename = os.path.basename(fig["path"])
            replacement = f"\n\n**Figure {real_fig_num}.** {caption}\n\n![Figure {real_fig_num}: {caption}](assets/{filename})\n\n"
        else:
            real_fig_num += 1
            replacement = f"\n\n**Figure {real_fig_num}.** {info['desc']}\n\n"
        manuscript = manuscript.replace(tag, replacement)

    real_tbl_num = 0
    for tag, tbl in tbl_tag_map.items():
        real_tbl_num += 1
        md = tbl.get("markdown", "")
        caption = tbl.get("caption", tbl.get("description", ""))
        tbl_path = tbl.get("path", "")
        tbl_filename = os.path.basename(tbl_path) if tbl_path else ""
        replacement = f"\n\n**Table {real_tbl_num}.** {caption}\n\n"
        if md:
            replacement += f"{md}\n\n"
        if tbl_filename and tbl_path and os.path.isfile(tbl_path):
            replacement += f"![Table {real_tbl_num}: {caption}](assets/{tbl_filename})\n\n"
        manuscript = manuscript.replace(tag, replacement)

    # Also catch any leftover freeform figure suggestions the LLM may have added
    leftover_patterns = [
        r'\((?:Suggest|Insert|Place)\s+(?:placing\s+)?Figure\s*\d*\s*(?:here)?[:\s]*.+?\)',
        r'\[(?:Insert|Place)\s+Figure\s*\d*\s*(?:here)?[:\s]*.+?\]',
        r'\(Placement Suggestion:\s*.+?\)',
    ]
    for pat in leftover_patterns:
        manuscript = re.sub(pat, '', manuscript, flags=re.IGNORECASE)

    # Extract abstract
    abstract = ""
    abs_match = re.search(r'(?:^|\n)##?\s*Abstract\s*\n+(.*?)(?=\n##?\s|\Z)', manuscript, re.DOTALL | re.IGNORECASE)
    if abs_match:
        abstract = abs_match.group(1).strip()

    # If manuscript doesn't include References section, append one from papers
    refs_section = ""
    if papers and not re.search(r'##?\s*References', manuscript, re.IGNORECASE):
        refs_section = "\n\n## References\n\n"
        for i, p in enumerate(papers, 1):
            auth = ", ".join(p.get("authors", [])[:6])
            if len(p.get("authors", [])) > 6:
                auth += ", et al"
            title = p.get("title", "")
            year = p.get("year", "n.d.")
            journal = p.get("journal", "")
            doi = p.get("doi", "")
            ref = f"{i}. {auth}. {title}. {journal}. {year}."
            if doi:
                ref += f" doi:{doi}"
            refs_section += ref + "\n"
        manuscript += refs_section

    return {"ok": True, "manuscript": manuscript, "abstract": abstract,
            "references": refs_section,
            "tokens": result.get("usage")}


def compile_references(args):
    """Phase 4.2 — Reference Compilation: Generate bibliography from papers.
    Args: {papers: [{title, authors, year, doi, journal}], format: "vancouver"|"ama"|"bibtex"}
    Returns: {ok, bibliography, bibtex}"""
    papers = args.get("papers", [])
    fmt = args.get("format", "vancouver")

    if not papers:
        return {"error": "No papers to compile."}

    bib_entries = []
    bibtex_entries = []

    for i, p in enumerate(papers, 1):
        title = p.get("title", "")
        authors = p.get("authors", [])
        year = str(p.get("year", ""))
        doi = p.get("doi", "")
        journal = p.get("journal", "")

        # Vancouver format
        auth_str = ", ".join(authors[:6])
        if len(authors) > 6:
            auth_str += ", et al"
        vancouver = f"{i}. {auth_str}. {title}. {journal}. {year}"
        if doi:
            vancouver += f". doi:{doi}"
        bib_entries.append(vancouver)

        # BibTeX
        key = re.sub(r'[^a-zA-Z0-9]', '', (authors[0].split()[0] if authors else "unknown")) + year
        bib = f"@article{{{key}{i},\n"
        bib += f"  title = {{{title}}},\n"
        if authors:
            bib += f"  author = {{{' and '.join(authors[:5])}}},\n"
        bib += f"  year = {{{year}}},\n"
        if journal:
            bib += f"  journal = {{{journal}}},\n"
        if doi:
            bib += f"  doi = {{{doi}}},\n"
        bib += "}\n"
        bibtex_entries.append(bib)

    return {
        "ok": True,
        "bibliography": "\n".join(bib_entries),
        "bibtex": "\n".join(bibtex_entries),
        "count": len(papers),
    }


def generate_blueprint(args):
    """Phase 3.1 — Blueprint Agent: Generate paper structure based on study design.
    Args: {query, study_design, papers_summary, api_key, provider, model, step_config}
    Returns: {ok, sections: [{section, requirements, subsections, needs_table, needs_figure, word_target}]}"""
    query = args.get("query", "")
    study_design = args.get("study_design", "systematic_review")
    papers_summary = args.get("papers_context", args.get("papers_summary", ""))
    
    guidelines_map = {
        "systematic_review": "PRISMA 2020",
        "meta_analysis": "PRISMA-MA + MOOSE",
        "narrative_review": "SANRA",
        "scoping_review": "PRISMA-ScR",
        "comparative": "Comparative analysis",
        "exploratory": "Exploratory research",
    }
    guidelines = args.get("guidelines", guidelines_map.get(study_design, "academic"))
    
    extracted_texts = args.get("extracted_texts", "")
    if extracted_texts:
        papers_summary += "\n\n--- Full-text excerpts ---\n" + extracted_texts[:6000]
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "strong", "max_tokens": 8192, "temperature": 0.3, "use_structured_output": True, "use_thinking": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}

    system = _build_system_prompt(args, sc, 
        "You are a research manuscript architect. Generate a detailed section-by-section blueprint "
        "for a {study_design} paper following {guidelines} guidelines.\n\n"
        "For each section, specify: what content to cover, required tables/figures, "
        "citation density expectations, and word count targets. "
        "Adapt the blueprint to the specific research question and available literature."
    )

    thinking = {"budget": sc.get("thinking_budget", 8192)} if sc["use_thinking"] and provider == "google" else None
    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": system},
                        {"role": "user", "content": f"Research question: {query}\nStudy design: {study_design}\n\nAvailable literature summary:\n{papers_summary[:4000]}"}],
                       temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                       response_schema=_SCHEMA_BLUEPRINT if sc["use_structured_output"] else None,
                       thinking_config=thinking)
    if "error" in result:
        return result

    parsed = result.get("structured")
    if not parsed:
        try:
            parsed = json.loads(re.search(r'\[.*\]', result["text"], re.DOTALL).group())
        except Exception:
            # Fallback: default section plan
            parsed = [
                {"section": "Introduction", "requirements": ["Background context", "Research gap", "Objectives"], "subsections": [], "needs_table": False, "needs_figure": False, "word_target": 800},
                {"section": "Methods", "requirements": ["Search strategy", "Inclusion/exclusion criteria", "Data extraction", "Quality assessment"], "subsections": ["Search Strategy", "Selection Criteria", "Data Extraction"], "needs_table": True, "needs_figure": True, "word_target": 1200},
                {"section": "Results", "requirements": ["Study selection flow", "Study characteristics", "Main findings", "Risk of bias"], "subsections": ["Study Selection", "Study Characteristics", "Synthesis of Results"], "needs_table": True, "needs_figure": True, "word_target": 1500},
                {"section": "Discussion", "requirements": ["Summary of findings", "Comparison with prior work", "Strengths and limitations", "Implications"], "subsections": ["Principal Findings", "Comparison with Literature", "Limitations", "Implications"], "needs_table": False, "needs_figure": False, "word_target": 1200},
            ]

    sections_raw = parsed if isinstance(parsed, list) else [parsed]
    # Normalize sections for frontend: {id, title, description, requirements}
    sections_out = []
    figure_plan = []
    table_plan = []
    for s in sections_raw:
        sec_title = s.get("section", "Section")
        sec_id = sec_title.lower().replace(" ", "_").replace("/", "_")
        reqs = s.get("requirements", [])
        desc = ", ".join(reqs) if isinstance(reqs, list) else str(reqs)
        reqs_str = desc
        sections_out.append({"id": sec_id, "title": sec_title, "description": desc, "requirements": reqs_str})
        if s.get("needs_figure"):
            figure_plan.append(f"Figure for {sec_title}")
        if s.get("needs_table"):
            table_plan.append(f"Table for {sec_title}")

    return {"ok": True, "sections": sections_out, "figure_plan": figure_plan,
            "table_plan": table_plan, "guidelines_map": {study_design: guidelines},
            "guidelines": guidelines, "tokens": result.get("usage")}


def citation_verifier_swarm(args):
    """Phase 3.3a — Citation Verifier: Cross-reference citations against paper list.
    Args: {sections: [{type, text}] OR section_text, papers, api_key, provider, model, step_config}
    Returns: {ok, verified, hallucinated, issues, pass, tokens}"""
    sections = args.get("sections", [])
    section_text = args.get("section_text", "")
    if sections and not section_text:
        section_text = "\n\n".join(f"## {s.get('type', 'Section').title()}\n{s.get('text', '')}" for s in sections)
    papers = args.get("papers", [])
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "fast", "max_tokens": 8192, "temperature": 0.0, "use_structured_output": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}
    if not section_text:
        return {"ok": True, "pass": True, "verified": [], "hallucinated": [], "issues": []}

    # Build paper reference list
    refs = ""
    for i, p in enumerate(papers, 1):
        auth = ", ".join(p.get("authors", [])[:3])
        refs += f"[{i}] {auth}. \"{p.get('title', '')}\" ({p.get('year', 'n.d.')})\n"

    system = _build_system_prompt(args, sc, 
        "You are a citation integrity auditor. Cross-reference every numbered citation [N] "
        "in the drafted text against the provided paper list. Verify that: "
        "(1) each citation number maps to a real paper, "
        "(2) the cited claim accurately reflects the source paper's abstract/findings, "
        "(3) there are no hallucinated or fabricated references. "
        "Flag any discrepancies with specific line references and suggestions for correction."
    )

    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": system},
                        {"role": "user", "content": f"Paper reference list:\n{refs}\n\nSection text to verify:\n{section_text[:8000]}"}],
                       temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                       response_schema=_SCHEMA_CITATION_VERIFY if sc["use_structured_output"] else None)
    if "error" in result:
        return {"ok": True, "pass": True, "verified": [], "hallucinated": [], "issues": ["Verification call failed"], "tokens": result.get("usage")}

    parsed = result.get("structured")
    if not parsed:
        try:
            parsed = json.loads(re.search(r'\{.*\}', result["text"], re.DOTALL).group())
        except Exception:
            parsed = {"verified": [], "hallucinated": [], "issues": [], "pass": True}

    # Normalize issues to [{section, issue, severity}] for frontend
    issues = []
    for h in parsed.get("hallucinated", []):
        if isinstance(h, str):
            issues.append({"section": "unknown", "issue": f"Hallucinated citation: {h}", "severity": "critical"})
        else:
            issues.append({"section": h.get("section", "unknown"), "issue": h.get("issue", str(h)), "severity": "critical"})
    for i in parsed.get("issues", []):
        if isinstance(i, str):
            issues.append({"section": "unknown", "issue": i, "severity": "warning"})
        else:
            issues.append({"section": i.get("section", "unknown"), "issue": i.get("issue", str(i)), "severity": i.get("severity", "warning")})

    return {"ok": True, "issues": issues, "verified": parsed.get("verified", []),
            "hallucinated": parsed.get("hallucinated", []),
            "pass": len(parsed.get("hallucinated", [])) == 0,
            "tokens": result.get("usage")}


def guidelines_compliance_check(args):
    """Phase 3.3b — Guidelines Compliance: Check manuscript against reporting guidelines.
    Args: {sections: [{type, text}] OR section_text, study_design, guidelines_map, api_key, provider, model, step_config}
    Returns: {ok, checklist: [{item, status, fix}], tokens}"""
    sections = args.get("sections", [])
    section_text = args.get("section_text", "")
    if sections and not section_text:
        section_text = "\n\n".join(f"## {s.get('type', 'Section').title()}\n{s.get('text', '')}" for s in sections)
    section_type = args.get("section_type", args.get("study_design", ""))
    
    study_design = args.get("study_design", "systematic_review")
    guidelines_map = args.get("guidelines_map", {
        "systematic_review": "PRISMA 2020",
        "meta_analysis": "PRISMA-MA + MOOSE",
        "narrative_review": "SANRA",
        "scoping_review": "PRISMA-ScR",
        "comparative": "Comparative analysis",
        "exploratory": "Exploratory research",
    })
    guidelines = args.get("guidelines", guidelines_map.get(study_design, "PRISMA"))

    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "strong", "max_tokens": 8192, "temperature": 0.1, "use_structured_output": True})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}
    if not section_text:
        return {"ok": True, "checklist": []}

    system = _build_system_prompt(args, sc, 
        "You are a research methodology compliance checker. Evaluate this '{section_type}' section "
        "against {guidelines} reporting guidelines. Check for: completeness of required elements, "
        "methodological rigor, proper statistical reporting, bias assessment, and ethical considerations.",
        section_type=section_type, guidelines=guidelines
    )

    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": system},
                        {"role": "user", "content": f"Section type: {section_type}\nGuidelines: {guidelines}\n\nSection text:\n{section_text[:8000]}"}],
                       temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                       response_schema=_SCHEMA_GUIDELINES if sc["use_structured_output"] else None)
    if "error" in result:
        return {"ok": True, "checklist": [], "tokens": result.get("usage")}

    parsed = result.get("structured")
    if not parsed:
        try:
            parsed = json.loads(re.search(r'\{.*\}', result["text"], re.DOTALL).group())
        except Exception:
            parsed = {}

    # Normalize to checklist format [{item, status, fix}]
    checklist = parsed.get("checklist", [])
    if not checklist:
        # Convert from compliant/violations format
        for c in parsed.get("compliant", []):
            item_text = c if isinstance(c, str) else c.get("item", str(c))
            checklist.append({"item": item_text, "status": "met", "fix": ""})
        for v in parsed.get("violations", []):
            if isinstance(v, str):
                checklist.append({"item": v, "status": "not_met", "fix": ""})
            else:
                checklist.append({"item": v.get("item", str(v)), "status": "not_met", "fix": v.get("fix", v.get("suggestion", ""))})

    return {"ok": True, "checklist": checklist, "tokens": result.get("usage")}


# ══════════════════════════════════════════════════════════════════════════════
# ██  FREE DATA SOURCE TOOLS (v6.1 atomic agents)
# ══════════════════════════════════════════════════════════════════════════════

def _search_europe_pmc(query: str, max_results: int = 50) -> list:
    """Search Europe PMC REST API (free, no key required)."""
    try:
        url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
        params = {
            "query": query, "resultType": "core", "format": "json",
            "pageSize": min(max_results, 100), "sort": "CITED desc",
        }
        resp = requests.get(url, params=params, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        papers = []
        for r in data.get("resultList", {}).get("result", []):
            authors = []
            auth_list = r.get("authorList", {}).get("author", [])
            if isinstance(auth_list, list):
                authors = [f"{a.get('lastName', '')} {a.get('initials', '')}".strip() for a in auth_list[:5]]
            papers.append({
                "title": r.get("title", "").rstrip("."),
                "authors": authors,
                "year": str(r.get("pubYear", "")),
                "abstract": r.get("abstractText", ""),
                "doi": r.get("doi", ""),
                "pmid": r.get("pmid", ""),
                "citations": r.get("citedByCount", 0),
                "url": f"https://europepmc.org/article/{r.get('source','MED')}/{r.get('id','')}",
                "source": "Europe PMC",
                "journal": r.get("journalTitle", ""),
            })
        return papers
    except Exception as e:
        return []


def _search_clinical_trials(condition: str, intervention: str = "", status: str = "COMPLETED", max_results: int = 50) -> list:
    """Search ClinicalTrials.gov v2 API (free)."""
    try:
        url = "https://clinicaltrials.gov/api/v2/studies"
        params = {
            "query.cond": condition,
            "query.intr": intervention,
            "filter.overallStatus": status,
            "pageSize": min(max_results, 100),
            "format": "json",
            "fields": "NCTId,BriefTitle,OfficialTitle,BriefSummary,OverallStatus,Phase,EnrollmentCount,StartDate,CompletionDate,Condition,InterventionName,PrimaryOutcome,StudyType",
        }
        params = {k: v for k, v in params.items() if v}
        resp = requests.get(url, params=params, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        trials = []
        for s in data.get("studies", []):
            pm = s.get("protocolSection", {})
            id_m = pm.get("identificationModule", {})
            desc_m = pm.get("descriptionModule", {})
            status_m = pm.get("statusModule", {})
            design_m = pm.get("designModule", {})
            arms_m = pm.get("armsInterventionsModule", {})
            outcomes_m = pm.get("outcomesModule", {})
            interventions = [i.get("interventionName", "") for i in arms_m.get("interventions", [])]
            primary_outcomes = [o.get("measure", "") for o in outcomes_m.get("primaryOutcomes", [])]
            trials.append({
                "nct_id": id_m.get("nctId", ""),
                "title": id_m.get("briefTitle", id_m.get("officialTitle", "")),
                "summary": desc_m.get("briefSummary", ""),
                "status": status_m.get("overallStatus", ""),
                "phase": design_m.get("phases", []),
                "enrollment": design_m.get("enrollmentInfo", {}).get("count", 0),
                "start_date": status_m.get("startDateStruct", {}).get("date", ""),
                "completion_date": status_m.get("completionDateStruct", {}).get("date", ""),
                "conditions": pm.get("conditionsModule", {}).get("conditions", []),
                "interventions": interventions,
                "primary_outcomes": primary_outcomes,
                "url": f"https://clinicaltrials.gov/study/{id_m.get('nctId', '')}",
                "source": "ClinicalTrials.gov",
            })
        return trials
    except Exception as e:
        return []


def _query_openfda_adverse_events(drug_name: str, limit: int = 100) -> dict:
    """Query OpenFDA drug adverse event reports (free, no key required)."""
    try:
        url = "https://api.fda.gov/drug/event.json"
        params = {
            "search": f'patient.drug.openfda.generic_name:"{drug_name}" OR patient.drug.openfda.brand_name:"{drug_name}"',
            "count": "patient.reaction.reactionmeddrapt.exact",
            "limit": min(limit, 1000),
        }
        resp = requests.get(url, params=params, timeout=15)
        if resp.status_code == 404:
            return {"ok": True, "drug": drug_name, "adverse_events": [], "total": 0}
        resp.raise_for_status()
        data = resp.json()
        events = [{"reaction": r.get("term", ""), "count": r.get("count", 0)}
                  for r in data.get("results", [])]
        return {"ok": True, "drug": drug_name, "adverse_events": events, "total": len(events)}
    except Exception as e:
        return {"ok": False, "error": str(e), "drug": drug_name, "adverse_events": []}


def _query_openfda_drug_labels(drug_name: str) -> dict:
    """Query OpenFDA drug label database for prescribing information (free)."""
    try:
        url = "https://api.fda.gov/drug/label.json"
        params = {
            "search": f'openfda.generic_name:"{drug_name}" OR openfda.brand_name:"{drug_name}"',
            "limit": 3,
        }
        resp = requests.get(url, params=params, timeout=15)
        if resp.status_code == 404:
            return {"ok": True, "drug": drug_name, "labels": []}
        resp.raise_for_status()
        data = resp.json()
        labels = []
        for r in data.get("results", []):
            openfda = r.get("openfda", {})
            labels.append({
                "brand_name": openfda.get("brand_name", []),
                "generic_name": openfda.get("generic_name", []),
                "manufacturer": openfda.get("manufacturer_name", []),
                "route": openfda.get("route", []),
                "indications": (r.get("indications_and_usage") or [""])[0][:500],
                "warnings": (r.get("warnings") or [""])[0][:500],
                "dosage": (r.get("dosage_and_administration") or [""])[0][:500],
                "contraindications": (r.get("contraindications") or [""])[0][:300],
            })
        return {"ok": True, "drug": drug_name, "labels": labels}
    except Exception as e:
        return {"ok": False, "error": str(e), "drug": drug_name, "labels": []}


def _lookup_mesh_terms(term: str) -> dict:
    """Look up MeSH terms using NLM E-Utilities (free)."""
    try:
        # Step 1: search MeSH for the term
        search_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
        params = {"db": "mesh", "term": term, "retmode": "json", "retmax": 10}
        resp = requests.get(search_url, params=params, timeout=10)
        resp.raise_for_status()
        ids = resp.json().get("esearchresult", {}).get("idlist", [])
        if not ids:
            return {"ok": True, "term": term, "mesh_terms": [], "synonyms": []}

        # Step 2: fetch summaries for first few IDs
        summary_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
        params = {"db": "mesh", "id": ",".join(ids[:5]), "retmode": "json"}
        time.sleep(0.4)
        resp = requests.get(summary_url, params=params, timeout=10)
        resp.raise_for_status()
        doc = resp.json().get("result", {})

        mesh_terms = []
        synonyms = []
        for uid in ids[:5]:
            item = doc.get(uid, {})
            name = item.get("ds_meshterms", [])
            if isinstance(name, list):
                mesh_terms.extend(name)
            scope_note = item.get("ds_scopenote", "")
            entry_terms = item.get("ds_termsyn", [])
            if isinstance(entry_terms, list):
                synonyms.extend(entry_terms)

        return {"ok": True, "term": term, "mesh_terms": list(set(mesh_terms))[:10],
                "synonyms": list(set(synonyms))[:20]}
    except Exception as e:
        return {"ok": False, "error": str(e), "term": term, "mesh_terms": [], "synonyms": []}


def _lookup_rxnorm(drug_name: str) -> dict:
    """Look up drug information from NLM RxNorm API (free)."""
    try:
        # Get RxCUI
        url = f"https://rxnav.nlm.nih.gov/REST/rxcui.json"
        params = {"name": drug_name, "search": 1}
        resp = requests.get(url, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        rxcui_group = data.get("idGroup", {})
        rxcuis = rxcui_group.get("rxnormId", [])
        if not rxcuis:
            return {"ok": True, "drug": drug_name, "rxcui": None, "brand_names": [], "drug_classes": []}

        rxcui = rxcuis[0]
        time.sleep(0.3)

        # Get related terms (brand names, drug classes)
        rel_url = f"https://rxnav.nlm.nih.gov/REST/rxcui/{rxcui}/allrelated.json"
        resp = requests.get(rel_url, timeout=10)
        resp.raise_for_status()
        all_related = resp.json().get("allRelatedGroup", {}).get("conceptGroup", [])

        brand_names = []
        drug_classes = []
        ingredients = []
        for group in all_related:
            tty = group.get("tty", "")
            concepts = group.get("conceptProperties", [])
            if isinstance(concepts, list):
                names = [c.get("name", "") for c in concepts if c.get("name")]
                if tty == "BN":
                    brand_names.extend(names)
                elif tty == "IN":
                    ingredients.extend(names)

        # Get drug classes
        class_url = f"https://rxnav.nlm.nih.gov/REST/rxclass/class/byRxcui.json"
        params = {"rxcui": rxcui, "relaSource": "MESHPA"}
        try:
            time.sleep(0.3)
            resp = requests.get(class_url, params=params, timeout=10)
            for entry in resp.json().get("rxclassDrugInfoList", {}).get("rxclassDrugInfo", []):
                cls_name = entry.get("rxclassMinConceptItem", {}).get("className", "")
                if cls_name:
                    drug_classes.append(cls_name)
        except Exception:
            pass

        return {
            "ok": True, "drug": drug_name, "rxcui": rxcui,
            "brand_names": list(set(brand_names))[:10],
            "ingredients": list(set(ingredients))[:5],
            "drug_classes": list(set(drug_classes))[:10],
        }
    except Exception as e:
        return {"ok": False, "error": str(e), "drug": drug_name, "rxcui": None}


def _check_retraction_watch(doi: str) -> dict:
    """Check if a paper has been retracted via CrossRef + Retraction Watch CSV."""
    result = {"doi": doi, "retracted": False, "reason": None, "source": None, "date": None}
    try:
        # 1. CrossRef check
        if doi:
            url = f"https://api.crossref.org/works/{requests.utils.quote(doi, safe='')}"
            resp = requests.get(url, timeout=10, headers={"User-Agent": "ZenithResearch/6.1 (mailto:research@zenith.app)"})
            if resp.ok:
                data = resp.json().get("message", {})
                update_to = data.get("update-to", [])
                for update in update_to:
                    if update.get("type", "").lower() in ("retraction", "withdrawal"):
                        result["retracted"] = True
                        result["source"] = "CrossRef"
                        result["date"] = update.get("updated", {}).get("date-parts", [[""]])[0]
                        break

        # 2. Retraction Watch CSV (cached locally or fetched)
        if not result["retracted"]:
            rw_path = os.path.join(os.path.dirname(__file__), "retraction_watch_cache.csv")
            if os.path.exists(rw_path) and doi:
                import csv
                with open(rw_path, "r", encoding="utf-8", errors="ignore") as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        row_doi = row.get("OriginalPaperDOI", "").strip().lower()
                        if row_doi and row_doi == doi.lower():
                            result["retracted"] = True
                            result["reason"] = row.get("Reason", "")
                            result["source"] = "Retraction Watch"
                            result["date"] = row.get("RetractionDate", "")
                            break
    except Exception as e:
        result["error"] = str(e)
    return result


def _check_predatory_journal(journal_name: str) -> dict:
    """Check if a journal is on the Beall's list (cached local list)."""
    try:
        beall_path = os.path.join(os.path.dirname(__file__), "bealls_list_cache.txt")
        if not os.path.exists(beall_path):
            # Try to fetch a simple public mirror
            try:
                resp = requests.get(
                    "https://raw.githubusercontent.com/scholarly-comms-product-team/predatory-journals/main/journals.txt",
                    timeout=10
                )
                if resp.ok:
                    with open(beall_path, "w", encoding="utf-8") as f:
                        f.write(resp.text)
            except Exception:
                return {"ok": True, "journal": journal_name, "is_predatory": None, "note": "Cache not available"}

        if os.path.exists(beall_path):
            name_lower = journal_name.lower()
            with open(beall_path, "r", encoding="utf-8") as f:
                for line in f:
                    if name_lower in line.lower().strip():
                        return {"ok": True, "journal": journal_name, "is_predatory": True, "matched_entry": line.strip()}
            return {"ok": True, "journal": journal_name, "is_predatory": False}

        return {"ok": True, "journal": journal_name, "is_predatory": None, "note": "List unavailable"}
    except Exception as e:
        return {"ok": False, "error": str(e), "journal": journal_name, "is_predatory": None}


# ══════════════════════════════════════════════════════════════════════════════
# ██  STATISTICAL ANALYSIS TOOLS (v6.1 meta-analysis & visualization)
# ══════════════════════════════════════════════════════════════════════════════

def run_meta_analysis(args: dict) -> dict:
    """Run fixed-effects or random-effects meta-analysis using scipy/statsmodels.
    Args: {studies: [{study_id, effect_size, variance|se|ci_lower&ci_upper, n_total}], method: 'fixed'|'random'}
    Returns: {pooled_effect, ci_lower, ci_upper, i_squared, q_stat, p_q, tau_squared, k}"""
    try:
        import numpy as np
        studies = args.get("studies", [])
        method = args.get("method", "random").lower()

        if len(studies) < 2:
            return {"ok": False, "error": "Need at least 2 studies for meta-analysis"}

        effects = []
        variances = []
        for s in studies:
            es = float(s.get("effect_size", 0))
            if "variance" in s:
                var = float(s["variance"])
            elif "se" in s:
                var = float(s["se"]) ** 2
            elif "ci_lower" in s and "ci_upper" in s:
                se = (float(s["ci_upper"]) - float(s["ci_lower"])) / (2 * 1.96)
                var = se ** 2
            else:
                var = 0.01  # fallback
            effects.append(es)
            variances.append(var)

        effects = np.array(effects)
        variances = np.array(variances)
        weights_fixed = 1.0 / variances

        # Fixed-effect pooled estimate
        pooled_fe = np.sum(weights_fixed * effects) / np.sum(weights_fixed)
        var_fe = 1.0 / np.sum(weights_fixed)

        # Q statistic and I²
        q = np.sum(weights_fixed * (effects - pooled_fe) ** 2)
        k = len(effects)
        df = k - 1
        p_q = 1.0 - float(__import__("scipy.stats", fromlist=["chi2"]).chi2.cdf(q, df))
        i_squared = max(0.0, (q - df) / q * 100) if q > 0 else 0.0

        # DerSimonian-Laird tau² for random effects
        c = np.sum(weights_fixed) - np.sum(weights_fixed ** 2) / np.sum(weights_fixed)
        tau2 = max(0.0, (q - df) / c) if c > 0 else 0.0

        if method == "random":
            weights = 1.0 / (variances + tau2)
        else:
            weights = weights_fixed

        pooled = np.sum(weights * effects) / np.sum(weights)
        var_pooled = 1.0 / np.sum(weights)
        se_pooled = float(np.sqrt(var_pooled))
        z = pooled / se_pooled if se_pooled > 0 else 0.0
        from scipy.stats import norm
        p_value = 2 * (1 - norm.cdf(abs(z)))

        return {
            "ok": True,
            "method": method,
            "k": k,
            "pooled_effect": round(float(pooled), 4),
            "se": round(se_pooled, 4),
            "ci_lower": round(float(pooled - 1.96 * se_pooled), 4),
            "ci_upper": round(float(pooled + 1.96 * se_pooled), 4),
            "z": round(float(z), 3),
            "p_value": round(float(p_value), 4),
            "q_stat": round(float(q), 3),
            "p_q": round(float(p_q), 4),
            "i_squared": round(float(i_squared), 1),
            "tau_squared": round(float(tau2), 4),
        }
    except ImportError:
        return {"ok": False, "error": "scipy not installed. Run: pip install scipy"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generate_forest_plot(args: dict) -> dict:
    """Generate a forest plot from meta-analysis data.
    Args: {studies: [{study_id, effect_size, ci_lower, ci_upper, weight}], pooled_effect, ci_lower, ci_upper, xlabel, title}
    Returns: {image_base64, path}"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import numpy as np, io, base64

        studies = args.get("studies", [])
        pooled = args.get("pooled_effect", 0)
        pooled_lo = args.get("ci_lower", pooled - 0.1)
        pooled_hi = args.get("ci_upper", pooled + 0.1)
        xlabel = args.get("xlabel", "Effect Size (SMD)")
        title = args.get("title", "Forest Plot")

        n = len(studies)
        if n == 0:
            return {"ok": False, "error": "No studies provided"}

        fig, ax = plt.subplots(figsize=(10, max(4, n * 0.5 + 2)))
        fig.patch.set_facecolor("#0f1520")
        ax.set_facecolor("#0f1520")

        y_positions = list(range(n, 0, -1))
        max_weight = max((s.get("weight", 1) for s in studies), default=1)

        for i, (s, y) in enumerate(zip(studies, y_positions)):
            es = s.get("effect_size", 0)
            lo = s.get("ci_lower", es - 0.2)
            hi = s.get("ci_upper", es + 0.2)
            w = s.get("weight", 1) / max_weight
            color = "#22d3ee"

            ax.plot([lo, hi], [y, y], color=color, linewidth=1, alpha=0.7)
            ax.scatter([es], [y], s=max(20, w * 120), color=color, zorder=5, edgecolors="white", linewidths=0.3)
            ax.text(-0.02, y, s.get("study_id", f"Study {i+1}"),
                    ha="right", va="center", fontsize=8, color="#94a3b8")
            ax.text(hi + 0.02, y,
                    f"{es:.2f} [{lo:.2f}, {hi:.2f}]",
                    ha="left", va="center", fontsize=7, color="#64748b",
                    fontfamily="monospace")

        # Pooled diamond
        diamond_y = 0
        diamond_x = [pooled_lo, pooled, pooled_hi, pooled, pooled_lo]
        diamond_yy = [diamond_y, diamond_y + 0.35, diamond_y, diamond_y - 0.35, diamond_y]
        ax.fill(diamond_x, diamond_yy, color="#10b981", alpha=0.8, zorder=6)
        ax.text(-0.02, diamond_y, "Pooled",
                ha="right", va="center", fontsize=8, color="#10b981", fontweight="bold")
        ax.text(pooled_hi + 0.02, diamond_y,
                f"{pooled:.2f} [{pooled_lo:.2f}, {pooled_hi:.2f}]",
                ha="left", va="center", fontsize=7, color="#10b981",
                fontfamily="monospace")

        # Null effect line
        ax.axvline(x=0, color="#475569", linestyle="--", linewidth=0.8, alpha=0.6)

        ax.set_yticks([])
        ax.set_xlabel(xlabel, color="#94a3b8", fontsize=9)
        ax.set_title(title, color="#e2e8f0", fontsize=11, pad=12)
        ax.tick_params(colors="#64748b", labelsize=8)
        for spine in ax.spines.values():
            spine.set_edgecolor("#1c2536")

        plt.tight_layout(pad=1.5)

        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    facecolor=fig.get_facecolor())
        plt.close()
        buf.seek(0)
        image_b64 = base64.b64encode(buf.read()).decode()

        return {"ok": True, "image_base64": image_b64, "chart_type": "forest_plot"}
    except ImportError:
        return {"ok": False, "error": "matplotlib not installed. Run: pip install matplotlib"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generate_funnel_plot(args: dict) -> dict:
    """Generate a funnel plot to assess publication bias.
    Args: {studies: [{study_id, effect_size, se|variance}], pooled_effect}
    Returns: {image_base64, egger_p, begg_p}"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np, io, base64
        from scipy import stats

        studies = args.get("studies", [])
        pooled = args.get("pooled_effect", 0)
        n = len(studies)
        if n < 3:
            return {"ok": False, "error": "Need at least 3 studies for funnel plot"}

        effects = []
        ses = []
        for s in studies:
            es = float(s.get("effect_size", 0))
            if "se" in s:
                se = float(s["se"])
            elif "variance" in s:
                se = float(s["variance"]) ** 0.5
            elif "ci_lower" in s and "ci_upper" in s:
                se = (float(s["ci_upper"]) - float(s["ci_lower"])) / (2 * 1.96)
            else:
                se = 0.1
            effects.append(es)
            ses.append(se)

        effects = np.array(effects)
        ses = np.array(ses)

        # Egger's test (regress ES/SE on 1/SE)
        precision = 1.0 / ses
        slope, intercept, r, p_egger, _ = stats.linregress(precision, effects / ses)
        begg_tau, p_begg = stats.kendalltau(effects, ses)

        fig, ax = plt.subplots(figsize=(7, 6))
        fig.patch.set_facecolor("#0f1520")
        ax.set_facecolor("#0f1520")

        ax.scatter(effects, ses, color="#22d3ee", alpha=0.8, s=60, edgecolors="white", linewidths=0.5, zorder=5)
        for i, s in enumerate(studies):
            ax.text(effects[i], ses[i] * 1.02, s.get("study_id", ""), fontsize=6,
                    ha="center", color="#64748b")

        # Funnel borders (95% CI lines)
        se_range = np.linspace(0, max(ses) * 1.1, 100)
        ax.plot([pooled - 1.96 * se_range, pooled + 1.96 * se_range],
                [se_range, se_range], color="#475569", linestyle="--", linewidth=0.8, alpha=0.5)

        ax.axvline(x=pooled, color="#10b981", linestyle="-", linewidth=1, alpha=0.7)
        ax.invert_yaxis()
        ax.set_xlabel("Effect Size", color="#94a3b8", fontsize=9)
        ax.set_ylabel("Standard Error", color="#94a3b8", fontsize=9)
        ax.set_title(f"Funnel Plot  |  Egger p={p_egger:.3f}  |  Begg p={p_begg:.3f}",
                     color="#e2e8f0", fontsize=10, pad=10)
        ax.tick_params(colors="#64748b", labelsize=8)
        for spine in ax.spines.values():
            spine.set_edgecolor("#1c2536")

        plt.tight_layout(pad=1.5)
        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
        plt.close()
        buf.seek(0)
        image_b64 = base64.b64encode(buf.read()).decode()

        return {
            "ok": True, "image_base64": image_b64, "chart_type": "funnel_plot",
            "egger_p": round(float(p_egger), 4),
            "begg_p": round(float(p_begg), 4),
            "asymmetric": p_egger < 0.05,
        }
    except ImportError:
        return {"ok": False, "error": "matplotlib/scipy not installed"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generate_prisma_flowchart(args: dict) -> dict:
    """Generate a PRISMA 2020 flowchart.
    Args: {records_identified, records_removed_duplicates, records_screened, records_excluded,
           reports_sought, reports_not_retrieved, reports_assessed, reports_excluded_reasons: [{reason, count}],
           studies_included}
    Returns: {image_base64}"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import io, base64

        ri = args.get("records_identified", 0)
        rd = args.get("records_removed_duplicates", 0)
        rs = args.get("records_screened", ri - rd)
        re_ = args.get("records_excluded", 0)
        rso = args.get("reports_sought", rs - re_)
        rnr = args.get("reports_not_retrieved", 0)
        ra = args.get("reports_assessed", rso - rnr)
        ex_reasons = args.get("reports_excluded_reasons", [])
        si = args.get("studies_included", 0)

        fig, ax = plt.subplots(figsize=(10, 12))
        fig.patch.set_facecolor("#06080d")
        ax.set_facecolor("#06080d")
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 12)
        ax.axis("off")

        bg_id = "#0f1520"
        bg_ex = "#1a1030"
        cy = "#22d3ee"
        gr = "#94a3b8"

        def box(x, y, w, h, text, color=bg_id, text_color=gr, fontsize=8):
            rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                                            linewidth=1, edgecolor="#1c2536", facecolor=color)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h/2, text, ha="center", va="center",
                    fontsize=fontsize, color=text_color, wrap=True,
                    multialignment="center")

        def arrow(x1, y1, x2, y2):
            ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                        arrowprops=dict(arrowstyle="->", color="#475569", lw=0.8))

        # Column headers
        ax.text(3, 11.7, "Identification", ha="center", fontsize=9, color=cy, fontweight="bold")
        ax.text(3, 9.2, "Screening", ha="center", fontsize=9, color=cy, fontweight="bold")
        ax.text(3, 6.2, "Eligibility", ha="center", fontsize=9, color=cy, fontweight="bold")
        ax.text(3, 3.2, "Included", ha="center", fontsize=9, color=cy, fontweight="bold")

        # Boxes - left column (identification flow)
        box(1.5, 10.8, 3, 0.7, f"Records identified\n(n = {ri})", bg_id, gr)
        arrow(3, 10.8, 3, 9.6)
        box(1.5, 8.8, 3, 0.7, f"Records screened\n(n = {rs})", bg_id, gr)
        arrow(3, 8.8, 3, 7.6)
        box(1.5, 6.8, 3, 0.7, f"Reports sought\n(n = {rso})", bg_id, gr)
        arrow(3, 6.8, 3, 5.6)
        box(1.5, 4.8, 3, 0.7, f"Reports assessed\n(n = {ra})", bg_id, gr)
        arrow(3, 4.8, 3, 3.6)
        box(1.5, 2.8, 3, 0.7, f"Studies included\n(n = {si})", "#0a1e0e", "#10b981", 9)

        # Boxes - right column (exclusions)
        box(5.5, 10.3, 3.5, 0.7, f"Duplicates removed\n(n = {rd})", bg_ex, "#64748b")
        ax.annotate("", xy=(5.5, 10.65), xytext=(4.5, 10.65),
                    arrowprops=dict(arrowstyle="->", color="#475569", lw=0.8))

        box(5.5, 8.3, 3.5, 0.7, f"Records excluded\n(n = {re_})", bg_ex, "#64748b")
        ax.annotate("", xy=(5.5, 8.65), xytext=(4.5, 8.65),
                    arrowprops=dict(arrowstyle="->", color="#475569", lw=0.8))

        box(5.5, 6.3, 3.5, 0.7, f"Not retrieved\n(n = {rnr})", bg_ex, "#64748b")
        ax.annotate("", xy=(5.5, 6.65), xytext=(4.5, 6.65),
                    arrowprops=dict(arrowstyle="->", color="#475569", lw=0.8))

        ex_text = "\n".join(f"{r['reason']}: {r['count']}" for r in ex_reasons[:4]) or f"Excluded (n = {ra - si})"
        box(5.5, 4.3, 3.5, 0.7, ex_text, bg_ex, "#64748b", fontsize=6)
        ax.annotate("", xy=(5.5, 4.65), xytext=(4.5, 4.65),
                    arrowprops=dict(arrowstyle="->", color="#475569", lw=0.8))

        ax.set_title("PRISMA 2020 Flow Diagram", color="#e2e8f0", fontsize=12, pad=8)

        plt.tight_layout(pad=1.0)
        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
        plt.close()
        buf.seek(0)
        image_b64 = base64.b64encode(buf.read()).decode()
        return {"ok": True, "image_base64": image_b64, "chart_type": "prisma_flowchart"}
    except ImportError:
        return {"ok": False, "error": "matplotlib not installed. Run: pip install matplotlib"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generate_rob_plot(args: dict) -> dict:
    """Generate a Risk of Bias traffic light + summary plot.
    Args: {studies: [{study_id, domains: {domain_name: 'low'|'some'|'high'|'unclear'}}]}
    Returns: {image_base64}"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import numpy as np, io, base64

        studies = args.get("studies", [])
        if not studies:
            return {"ok": False, "error": "No studies provided"}

        all_domains = []
        for s in studies:
            for d in s.get("domains", {}).keys():
                if d not in all_domains:
                    all_domains.append(d)

        COLOR_MAP = {
            "low": "#10b981", "some": "#f59e0b", "some concerns": "#f59e0b",
            "high": "#ef4444", "unclear": "#64748b", "": "#1c2536"
        }
        LABEL_MAP = {"low": "L", "some": "SC", "some concerns": "SC", "high": "H", "unclear": "?", "": "—"}

        n_studies = len(studies)
        n_domains = len(all_domains)
        fig_w = max(8, n_domains * 1.2 + 3)
        fig_h = max(4, n_studies * 0.5 + 2)

        fig, ax = plt.subplots(figsize=(fig_w, fig_h))
        fig.patch.set_facecolor("#0f1520")
        ax.set_facecolor("#0f1520")

        for si, study in enumerate(studies):
            for di, domain in enumerate(all_domains):
                rating = study.get("domains", {}).get(domain, "unclear").lower()
                color = COLOR_MAP.get(rating, "#64748b")
                label = LABEL_MAP.get(rating, "?")

                rect = mpatches.FancyBboxPatch(
                    (di + 0.05, si + 0.05), 0.9, 0.9,
                    boxstyle="round,pad=0.05",
                    linewidth=0.5, edgecolor="#0f1520", facecolor=color + "cc"
                )
                ax.add_patch(rect)
                ax.text(di + 0.5, si + 0.5, label,
                        ha="center", va="center", fontsize=7,
                        color="white", fontweight="bold")

        ax.set_xlim(0, n_domains)
        ax.set_ylim(0, n_studies)
        ax.set_xticks(np.arange(n_domains) + 0.5)
        ax.set_xticklabels([d[:20] for d in all_domains], rotation=35, ha="right",
                            fontsize=7, color="#94a3b8")
        ax.set_yticks(np.arange(n_studies) + 0.5)
        ax.set_yticklabels([s.get("study_id", f"Study {i+1}")[:25]
                             for i, s in enumerate(studies)], fontsize=7, color="#94a3b8")

        ax.set_title("Risk of Bias Assessment", color="#e2e8f0", fontsize=11, pad=12)

        legend_patches = [
            mpatches.Patch(color=COLOR_MAP["low"], label="Low Risk"),
            mpatches.Patch(color=COLOR_MAP["some"], label="Some Concerns"),
            mpatches.Patch(color=COLOR_MAP["high"], label="High Risk"),
            mpatches.Patch(color=COLOR_MAP["unclear"], label="Unclear"),
        ]
        ax.legend(handles=legend_patches, loc="upper right", fontsize=7,
                  framealpha=0.3, facecolor="#1c2536", edgecolor="#475569",
                  labelcolor="#94a3b8")

        plt.tight_layout(pad=1.5)
        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
        plt.close()
        buf.seek(0)
        image_b64 = base64.b64encode(buf.read()).decode()
        return {"ok": True, "image_base64": image_b64, "chart_type": "rob_plot"}
    except ImportError:
        return {"ok": False, "error": "matplotlib not installed"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def calculate_grade(args: dict) -> dict:
    """Calculate GRADE certainty of evidence for an outcome.
    Args: {
        study_design: 'rct'|'observational',
        risk_of_bias: 0|1|2,          # 0=not serious, 1=serious, 2=very serious
        inconsistency: 0|1|2,
        indirectness: 0|1|2,
        imprecision: 0|1|2,
        publication_bias: 0|1|2,
        large_effect: 0|1|2,          # upgrade: 0=no, 1=large, 2=very large
        dose_response: bool,
        confounding_direction: 0|1    # upgrade: 0=no, 1=confounders reduce effect
        outcome_label: str
    }
    Returns: {certainty: 'high'|'moderate'|'low'|'very_low', score, rationale[]}"""
    study_design = args.get("study_design", "rct").lower()
    # Starting score: RCTs start at 4 (high), observational at 2 (low)
    score = 4 if "rct" in study_design or "randomis" in study_design else 2

    rationale = []
    # Downgrade domains
    for domain, key in [
        ("Risk of bias", "risk_of_bias"),
        ("Inconsistency", "inconsistency"),
        ("Indirectness", "indirectness"),
        ("Imprecision", "imprecision"),
        ("Publication bias", "publication_bias"),
    ]:
        val = int(args.get(key, 0))
        if val == 1:
            score -= 1
            rationale.append(f"↓ {domain}: serious")
        elif val >= 2:
            score -= 2
            rationale.append(f"↓↓ {domain}: very serious")

    # Upgrade domains (observational studies only)
    if "rct" not in study_design:
        large = int(args.get("large_effect", 0))
        if large == 1:
            score += 1; rationale.append("↑ Large effect (RR > 2)")
        elif large >= 2:
            score += 2; rationale.append("↑↑ Very large effect (RR > 5)")
        if args.get("dose_response"):
            score += 1; rationale.append("↑ Dose-response gradient")
        if int(args.get("confounding_direction", 0)) == 1:
            score += 1; rationale.append("↑ Confounders reduce effect")

    score = max(1, min(4, score))
    certainty_map = {4: "high", 3: "moderate", 2: "low", 1: "very_low"}
    certainty = certainty_map[score]

    return {
        "ok": True,
        "outcome": args.get("outcome_label", ""),
        "certainty": certainty,
        "score": score,
        "rationale": rationale,
        "grade_symbol": {"high": "⊕⊕⊕⊕", "moderate": "⊕⊕⊕◯", "low": "⊕⊕◯◯", "very_low": "⊕◯◯◯"}[certainty],
    }


def extract_pico_structured(args: dict) -> dict:
    """Extract structured PICO elements from a paper using LLM.
    Args: {text, title, api_key, provider, model}
    Returns: {population, intervention, comparator, outcome, sample_size, effect_size, ci, p_value, study_design}"""
    text = args.get("text", "")[:6000]
    title = args.get("title", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "")
    model = args.get("model", "")

    if not text:
        return {"ok": False, "error": "No text provided"}

    system = (
        "You are a biomedical data extraction specialist. Extract PICO elements from the research paper text. "
        "Return ONLY valid JSON with keys: population, intervention, comparator, outcome, sample_size, "
        "effect_size, ci_95, p_value, study_design. Use null for missing fields. Keep values concise (< 80 chars each)."
    )
    prompt = f"Paper title: {title}\n\nText excerpt:\n{text}\n\nExtract PICO elements as JSON:"

    result = _llm_chat(provider, api_key, model,
                       [{"role": "user", "content": prompt}],
                       temperature=0.1, max_tokens=1024, system_prompt=system,
                       response_format="json_object")
    if result.get("error"):
        return {"ok": False, "error": result["error"]}

    try:
        parsed = json.loads(result.get("content", "{}"))
        return {
            "ok": True,
            "population": parsed.get("population", ""),
            "intervention": parsed.get("intervention", ""),
            "comparator": parsed.get("comparator", ""),
            "outcome": parsed.get("outcome", ""),
            "sample_size": str(parsed.get("sample_size", "")),
            "effect_size": str(parsed.get("effect_size", "")),
            "ci": str(parsed.get("ci_95", "")),
            "p_value": str(parsed.get("p_value", "")),
            "study_design": parsed.get("study_design", ""),
            "tokens": result.get("usage"),
        }
    except json.JSONDecodeError:
        return {"ok": False, "error": "LLM returned invalid JSON"}


def run_pipeline_phase(args):
    """Run a specific phase of the v5.6 research pipeline.
    Args: {phase, query, study_design, papers, api_key, provider, model, step_config, ...}
    Returns: phase-specific results"""
    phase = args.get("phase", "")
    query = args.get("query", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    model = args.get("model", "")
    tavily_api_key = args.get("tavily_api_key", "")
    brave_api_key = args.get("brave_api_key", "")
    firecrawl_api_key = args.get("firecrawl_api_key", "")

    # Phases that need LLM access require an API key
    LLM_PHASES = {"validate", "generate_queries", "triage", "draft", "smooth",
                  "verify_citations", "novelty_check", "blueprint",
                  "citation_verify_swarm", "guidelines_check"}
    if phase in LLM_PHASES and not api_key:
        return {"error": "API key required."}

    if phase == "validate":
        _clear_prompt_logs()  # Clear at pipeline start
        return validate_research_query(args)

    elif phase == "generate_queries":
        return generate_search_queries(args)

    elif phase == "harvest":
        # Run all search queries across multiple databases
        search_queries = args.get("search_queries", [])
        if not search_queries:
            search_queries = [{"db": "pubmed", "query_string": query},
                              {"db": "semantic_scholar", "query_string": query},
                              {"db": "openalex", "query_string": query}]

        all_papers = []
        sources_used = []

        for sq in search_queries:
            db = sq.get("db", "").lower()
            q = sq.get("query_string", query)
            max_r = sq.get("max_results", 15)

            if db == "pubmed":
                papers = _search_pubmed(q, max_r)
                all_papers.extend(papers)
                if papers:
                    sources_used.append("PubMed")
                time.sleep(0.5)
            elif db in ("semantic_scholar", "s2"):
                papers = _search_semantic_scholar(q, max_r)
                all_papers.extend(papers)
                if papers:
                    sources_used.append("Semantic Scholar")
                time.sleep(1.0)
            elif db == "openalex":
                papers = _search_openalex(q, max_r)
                all_papers.extend(papers)
                if papers:
                    sources_used.append("OpenAlex")
                time.sleep(0.5)
            elif db == "arxiv":
                papers = _search_arxiv(q, max_r)
                all_papers.extend(papers)
                if papers:
                    sources_used.append("arXiv")
                time.sleep(1.0)
            elif db == "web":
                # Also search web for grey literature
                web_results = []
                if brave_api_key:
                    br = _web_search_brave(q, brave_api_key, 5)
                    web_results.extend(br.get("results", []))
                if tavily_api_key:
                    tv = _web_search_tavily(q, tavily_api_key, 5)
                    web_results.extend(tv.get("results", []))
                if not web_results:
                    ddg = _web_search_duckduckgo(q, 5)
                    web_results.extend(ddg.get("results", []))
                # Convert web results to paper-like format
                for wr in web_results:
                    all_papers.append({
                        "title": wr.get("title", ""), "authors": [],
                        "year": "", "abstract": wr.get("snippet", ""),
                        "doi": "", "citations": 0,
                        "url": wr.get("url", ""), "source": f"Web ({wr.get('source', 'search')})",
                    })
                if web_results:
                    sources_used.append("Web Search")

        # Deduplicate by title
        seen = set()
        unique = []
        for p in all_papers:
            key = re.sub(r'[^a-z0-9]+', ' ', p.get("title", "").lower()).strip()[:80]
            if key and key not in seen:
                seen.add(key)
                unique.append(p)

        # Enrich top papers with CrossRef citation counts
        for p in unique[:30]:
            if p.get("doi") and p.get("citations", 0) == 0:
                enriched = _enrich_crossref(p["doi"])
                if enriched.get("citations"):
                    p["citations"] = enriched["citations"]
                if enriched.get("journal"):
                    p["journal"] = enriched["journal"]
                time.sleep(0.2)

        unique.sort(key=lambda x: x.get("citations", 0), reverse=True)

        return {"ok": True, "papers": unique, "total": len(unique),
                "sources": sources_used}

    elif phase == "triage":
        return triage_papers(args)

    elif phase == "acquire":
        return acquire_papers(args)

    elif phase == "extract":
        # Extract text from downloaded PDFs
        paths = args.get("paths", [])
        results = []
        for path in paths[:30]:
            ext_result = extract_pdf_text({"path": path})
            results.append({
                "path": path,
                "ok": "error" not in ext_result,
                "text": ext_result.get("text", "")[:5000],
                "pages": ext_result.get("pages", 0),
                "error": ext_result.get("error"),
            })
        return {"ok": True, "results": results,
                "extracted_count": sum(1 for r in results if r["ok"])}

    elif phase == "draft":
        return draft_research_section(args)

    elif phase == "verify_citations":
        return verify_citations(args)

    elif phase == "novelty_check":
        return check_novelty(args)

    elif phase == "smooth":
        return smooth_manuscript(args)

    elif phase == "ingest_vectordb":
        return ingest_into_vectordb(args)

    elif phase == "query_vectordb":
        return query_vectordb(args)

    elif phase == "blueprint":
        return generate_blueprint(args)

    elif phase == "citation_verify_swarm":
        return citation_verifier_swarm(args)

    elif phase == "guidelines_check":
        return guidelines_compliance_check(args)

    elif phase == "generate_figures":
        return generate_pipeline_figures(args)

    elif phase == "illustrate":
        return scientific_illustrator_agent(args)

    elif phase == "generate_chart":
        return generate_chart(args)

    elif phase == "generate_table":
        return generate_table(args)

    elif phase == "compile_refs":
        return compile_references(args)

    elif phase == "export_snapshot":
        return export_research_snapshot(args)

    elif phase == "auto_rename":
        return auto_rename_thread(args)

    elif phase == "get_prompt_logs":
        return {"ok": True, "logs": _get_prompt_logs()}

    # ── v6.1 Free Data Source Phases ──────────────────────────────────────────

    elif phase == "europe_pmc_search":
        results = _search_europe_pmc(query, args.get("max_results", 50))
        return {"ok": True, "papers": results, "total": len(results), "source": "Europe PMC"}

    elif phase == "clinical_trials_search":
        results = _search_clinical_trials(
            args.get("condition", query),
            args.get("intervention", ""),
            args.get("status", "COMPLETED"),
            args.get("max_results", 50),
        )
        return {"ok": True, "trials": results, "total": len(results), "source": "ClinicalTrials.gov"}

    elif phase == "openfda_adverse_events":
        return _query_openfda_adverse_events(
            args.get("drug_name", query),
            args.get("limit", 100),
        )

    elif phase == "openfda_drug_labels":
        return _query_openfda_drug_labels(args.get("drug_name", query))

    elif phase == "mesh_lookup":
        return _lookup_mesh_terms(args.get("term", query))

    elif phase == "rxnorm_lookup":
        return _lookup_rxnorm(args.get("drug_name", query))

    elif phase == "retraction_check":
        dois = args.get("dois", [])
        if not dois and args.get("doi"):
            dois = [args["doi"]]
        results = [_check_retraction_watch(doi) for doi in dois[:50]]
        retracted = [r for r in results if r.get("retracted")]
        return {"ok": True, "checked": len(results), "retracted_count": len(retracted), "results": results}

    elif phase == "predatory_journal_check":
        journals = args.get("journals", [])
        if not journals and args.get("journal"):
            journals = [args["journal"]]
        results = [_check_predatory_journal(j) for j in journals[:20]]
        flagged = [r for r in results if r.get("is_predatory")]
        return {"ok": True, "checked": len(results), "flagged_count": len(flagged), "results": results}

    # ── v6.1 Statistical Analysis Phases ─────────────────────────────────────

    elif phase == "meta_analysis":
        return run_meta_analysis(args)

    elif phase == "forest_plot":
        return generate_forest_plot(args)

    elif phase == "funnel_plot":
        return generate_funnel_plot(args)

    elif phase == "prisma_flowchart":
        return generate_prisma_flowchart(args)

    elif phase == "rob_plot":
        return generate_rob_plot(args)

    elif phase == "grade_assess":
        return calculate_grade(args)

    elif phase == "pico_extract":
        return extract_pico_structured(args)

    # ── Utility Phases ────────────────────────────────────────────────────────

    elif phase == "test_connection":
        if not api_key:
            return {"ok": False, "error": "No API key provided"}
        result = _llm_chat(provider, api_key, model,
                           [{"role": "user", "content": "Reply with exactly: OK"}],
                           temperature=0.0, max_tokens=10)
        if result.get("error"):
            return {"ok": False, "error": result["error"]}
        return {"ok": True, "reply": result.get("content", ""), "model": model, "provider": provider}

    else:
        return {"error": f"Unknown pipeline phase: {phase}"}


# ══════════════════════════════════════════════════════════════════════════════
# ██  PUBLIC ACTIONS (registered in process_files.py ACTIONS dict)
# ══════════════════════════════════════════════════════════════════════════════

def research_chat(args):
    """Multi-turn research chat with optional tool dispatch.
    Args: {messages, api_key, provider, model, temperature, max_tokens,
           system_prompt, enabled_tools}
    Returns: {reply, tokens, type, data, tool_used, tool_results}
    """
    messages_in = args.get("messages", [])
    api_key = args.get("api_key", "")
    provider = args.get("provider", "openai")
    model = args.get("model", "")
    temperature = args.get("temperature", 0.7)
    max_tokens = args.get("max_tokens", 4096)
    system_prompt = args.get("system_prompt", "")
    enabled_tools = args.get("enabled_tools", [])
    tavily_api_key = args.get("tavily_api_key", "")
    brave_api_key = args.get("brave_api_key", "")
    firecrawl_api_key = args.get("firecrawl_api_key", "")

    if not api_key:
        return {"error": "API key required. Set it in Settings > API Keys."}
    if not messages_in:
        return {"error": "No messages provided."}

    # Build tool descriptions for the system prompt
    tool_desc = ""
    if enabled_tools:
        tools = []
        # ── Primary tools (v5.6 pipeline) ──
        if "pubmed" in enabled_tools:
            tools.append("- PUBMED_SEARCH: Search PubMed/MEDLINE for biomedical literature with MeSH terms")
        if "literature" in enabled_tools:
            tools.append("- LITERATURE_SEARCH: Search academic papers on arXiv, Semantic Scholar, and OpenAlex")
        if "web_search" in enabled_tools:
            tools.append("- WEB_SEARCH: Search the web for information (uses Brave, Tavily, Firecrawl, DuckDuckGo — aggregated & deduplicated)")
        if "scihub" in enabled_tools:
            tools.append("- SCIHUB_FETCH: Download a paper PDF via Sci-Hub/Unpaywall given a DOI")
        if "validate_query" in enabled_tools:
            tools.append("- VALIDATE_QUERY: Validate if a research question is suitable for systematic review")
        if "mesh_queries" in enabled_tools:
            tools.append("- MESH_QUERIES: Generate optimized MeSH/Boolean search strings for a research question")
        if "triage" in enabled_tools:
            tools.append("- TRIAGE: Screen papers for relevance to the research question")
        if "draft_section" in enabled_tools:
            tools.append("- DRAFT_SECTION: Draft a specific section of a research paper with citations")
        # ── Auxiliary tools ──
        if "pdf_extract" in enabled_tools:
            tools.append("- PDF_EXTRACT: Extract text from PDF files")
        if "novelty" in enabled_tools:
            tools.append("- NOVELTY_CHECK: Assess how novel a research idea is")
        if "citation_verify" in enabled_tools:
            tools.append("- CITATION_VERIFY: Verify if citations/references are accurate")
        if "experiment" in enabled_tools:
            tools.append("- EXPERIMENT: Run Python code in a sandboxed environment")
        if "generate_chart" in enabled_tools:
            tools.append("- GENERATE_CHART: Generate a chart (bar, line, pie, scatter, heatmap) from data")
        if "generate_table" in enabled_tools:
            tools.append("- GENERATE_TABLE: Generate a formatted data table")
        # ── V6.1 Free Data Source Tools ──
        tools.append("- EUROPE_PMC_SEARCH: Search Europe PMC for biomedical literature (free, no key)")
        tools.append("- CLINICAL_TRIALS_SEARCH: Search ClinicalTrials.gov v2 API for registered trials")
        tools.append("- OPENFDA_ADVERSE_EVENTS: Query OpenFDA for drug adverse event reports by drug name")
        tools.append("- OPENFDA_DRUG_LABELS: Fetch FDA drug label/prescribing information")
        tools.append("- MESH_LOOKUP: Look up MeSH controlled vocabulary terms and synonyms via NLM")
        tools.append("- RXNORM_LOOKUP: Look up drug info, brand names, and drug classes via NLM RxNorm")
        tools.append("- RETRACTION_CHECK: Check if papers (by DOI) have been retracted")
        # ── V6.1 Statistical Tools ──
        tools.append("- META_ANALYSIS: Run fixed/random-effects meta-analysis (provide studies with effect sizes)")
        tools.append("- FOREST_PLOT: Generate a forest plot image from meta-analysis data")
        tools.append("- FUNNEL_PLOT: Generate a funnel plot to assess publication bias")
        tools.append("- PRISMA_FLOWCHART: Generate a PRISMA 2020 flow diagram for systematic reviews")
        tools.append("- ROB_PLOT: Generate a risk of bias traffic light / summary plot")
        tools.append("- GRADE_ASSESS: Calculate GRADE certainty of evidence for an outcome")
        tools.append("- PICO_EXTRACT: Extract structured PICO elements from a paper's text using AI")
        if tools:
            tool_desc = (
                "\n\nYou have access to these research tools:\n" +
                "\n".join(tools) +
                "\n\nWhen you need to use a tool, include a tool call tag in your response like: "
                "[TOOL:TOOL_NAME]{\"param\": \"value\"}[/TOOL]\n"
                "For PUBMED_SEARCH: {\"query\": \"...\", \"max_results\": 20}\n"
                "For LITERATURE_SEARCH: {\"query\": \"...\", \"max_results\": 5}\n"
                "For WEB_SEARCH: {\"query\": \"...\"}\n"
                "For SCIHUB_FETCH: {\"doi\": \"10.1234/example\"}\n"
                "For VALIDATE_QUERY: {\"query\": \"your research question\"}\n"
                "For MESH_QUERIES: {\"query\": \"research question\", \"domain\": \"biomedical\"}\n"
                "For TRIAGE: {\"papers\": [{\"title\":\"...\",\"abstract\":\"...\"}], \"query\": \"...\"}\n"
                "For DRAFT_SECTION: {\"section_type\": \"introduction\", \"query\": \"...\", \"papers_context\": \"...\"}\n"
                "For PDF_EXTRACT: {\"path\": \"...\"}\n"
                "For NOVELTY_CHECK: {\"idea\": \"...\"}\n"
                "For CITATION_VERIFY: {\"citations\": [\"title1\", \"title2\"]}\n"
                "For EXPERIMENT: {\"code\": \"print('hello')\", \"timeout_sec\": 30}\n"
                "For GENERATE_CHART: {\"chart_type\": \"bar\", \"data\": [10,20,30], \"labels\": [\"A\",\"B\",\"C\"], \"title\": \"My Chart\"}\n"
                "For GENERATE_TABLE: {\"headers\": [\"Col1\",\"Col2\"], \"rows\": [[\"a\",\"b\"]], \"title\": \"My Table\"}\n"
                "For EUROPE_PMC_SEARCH: {\"query\": \"...\", \"max_results\": 20}\n"
                "For CLINICAL_TRIALS_SEARCH: {\"condition\": \"diabetes\", \"intervention\": \"metformin\", \"status\": \"COMPLETED\"}\n"
                "For OPENFDA_ADVERSE_EVENTS: {\"drug_name\": \"aspirin\", \"limit\": 50}\n"
                "For OPENFDA_DRUG_LABELS: {\"drug_name\": \"metformin\"}\n"
                "For MESH_LOOKUP: {\"term\": \"myocardial infarction\"}\n"
                "For RXNORM_LOOKUP: {\"drug_name\": \"ibuprofen\"}\n"
                "For RETRACTION_CHECK: {\"dois\": [\"10.1000/xyz123\"]}\n"
                "For META_ANALYSIS: {\"studies\": [{\"study_id\": \"Smith 2020\", \"effect_size\": 0.5, \"se\": 0.1}], \"method\": \"random\"}\n"
                "For FOREST_PLOT: {\"studies\": [...], \"pooled_effect\": 0.5, \"ci_lower\": 0.3, \"ci_upper\": 0.7}\n"
                "For FUNNEL_PLOT: {\"studies\": [...], \"pooled_effect\": 0.5}\n"
                "For PRISMA_FLOWCHART: {\"records_identified\": 500, \"records_removed_duplicates\": 80, \"records_excluded\": 200, \"studies_included\": 12}\n"
                "For ROB_PLOT: {\"studies\": [{\"study_id\": \"Smith 2020\", \"domains\": {\"Selection bias\": \"low\", \"Performance bias\": \"high\"}}]}\n"
                "For GRADE_ASSESS: {\"study_design\": \"rct\", \"risk_of_bias\": 0, \"inconsistency\": 1, \"indirectness\": 0, \"imprecision\": 1, \"publication_bias\": 0, \"outcome_label\": \"All-cause mortality\"}\n"
                "For PICO_EXTRACT: {\"text\": \"...\", \"title\": \"...\"}\n"
                "You can use multiple tools in a single response. Always explain what you found after using a tool."
            )

    full_system = system_prompt + tool_desc

    # Build messages for LLM
    chat_msgs = []
    if full_system:
        chat_msgs.append({"role": "system", "content": full_system})
    for m in messages_in:
        chat_msgs.append({"role": m.get("role", "user"), "content": m.get("content", "")})

    # First LLM call
    google_search = "web_search" in enabled_tools and provider == "google"
    result = _llm_chat(provider, api_key, model, chat_msgs, temperature, max_tokens, google_search=google_search)
    if "error" in result:
        return result

    reply_text = result["text"]
    usage = result.get("usage", {})
    tool_results = []

    # Check for tool calls in the response
    tool_pattern = r'\[TOOL:(\w+)\](.*?)\[/TOOL\]'
    tool_matches = re.findall(tool_pattern, reply_text, re.DOTALL)

    if tool_matches:
        for tool_name, tool_args_str in tool_matches:
            try:
                tool_args = json.loads(tool_args_str.strip())
            except json.JSONDecodeError:
                tool_args = {"query": tool_args_str.strip()}

            tool_result = _execute_tool(tool_name, tool_args, api_key, provider, model, tavily_api_key, brave_api_key, firecrawl_api_key)
            tool_results.append(tool_result)

        # If we got tool results, do a follow-up LLM call to synthesize
        if tool_results:
            tool_context = "\n\n".join([
                f"[Tool: {tr['tool_name']}]\n{tr.get('summary', json.dumps(tr.get('data', {}), indent=2)[:2000])}"
                for tr in tool_results
            ])

            # Clean tool tags from original reply
            clean_reply = re.sub(tool_pattern, '', reply_text).strip()

            followup_msgs = chat_msgs + [
                {"role": "assistant", "content": clean_reply},
                {"role": "user", "content": f"Here are the tool results:\n{tool_context}\n\nPlease synthesize these results into a clear, comprehensive response."}
            ]

            followup = _llm_chat(provider, api_key, model, followup_msgs, temperature, max_tokens, google_search=google_search)
            if "error" not in followup:
                reply_text = followup["text"]
                fu = followup.get("usage", {})
                usage["input_tokens"] = usage.get("input_tokens", 0) + fu.get("input_tokens", 0)
                usage["output_tokens"] = usage.get("output_tokens", 0) + fu.get("output_tokens", 0)

    # Determine response type
    resp_type = "text"
    resp_data = None
    if tool_results:
        for tr in tool_results:
            if tr.get("type") == "papers" and tr.get("data"):
                resp_type = "papers"
                resp_data = tr["data"]
                break

    return {
        "reply": reply_text,
        "tokens": {"input": usage.get("input_tokens", 0), "output": usage.get("output_tokens", 0)},
        "type": resp_type,
        "data": resp_data,
        "tool_results": tool_results if tool_results else None,
    }


def _execute_tool(tool_name, tool_args, api_key, provider, model, tavily_api_key="", brave_api_key="", firecrawl_api_key=""):
    """Dispatch a tool call and return structured result."""
    tool_name = tool_name.upper().strip()

    if tool_name == "LITERATURE_SEARCH":
        query = tool_args.get("query", "")
        max_results = tool_args.get("max_results", 5)
        year_min = tool_args.get("year_min", None)

        # Source order matches ARC: OpenAlex (10K/day) → S2 (1K/5min) → arXiv (1/3s)
        all_papers = []
        all_papers.extend(_search_openalex(query, max_results))
        time.sleep(0.5)
        all_papers.extend(_search_semantic_scholar(query, max_results, year_min))
        time.sleep(1.0)
        all_papers.extend(_search_arxiv(query, max_results))

        # Deduplicate by title similarity
        seen = set()
        unique = []
        for p in all_papers:
            key = p["title"].lower().strip()[:80]
            if key not in seen:
                seen.add(key)
                unique.append(p)

        # Sort by citations desc
        unique.sort(key=lambda x: x.get("citations", 0), reverse=True)
        unique = unique[:max_results * 2]  # keep top results

        summary = f"Found {len(unique)} papers for '{query}'."
        if unique:
            top3 = unique[:3]
            summary += " Top results: " + "; ".join([f"\"{p['title']}\" ({p.get('year','?')}, {p.get('citations',0)} cites)" for p in top3])

        return {"tool_name": "LITERATURE_SEARCH", "type": "papers", "data": unique, "summary": summary}

    elif tool_name == "WEB_SEARCH":
        query = tool_args.get("query", "")
        max_res = tool_args.get("max_results", 8)
        # ── Aggregate results from all available search providers ──
        all_results = []
        answer = ""
        sources_used = []

        # 1. Brave Search (highest quality, if key available)
        brave_key = brave_api_key or tool_args.get("brave_api_key", "")
        if brave_key:
            br = _web_search_brave(query, brave_key, max_res)
            if br.get("results"):
                all_results.extend(br["results"])
                sources_used.append("Brave")
            if br.get("answer"):
                answer = br["answer"]

        # 2. Tavily (AI-powered, if key available)
        tavily_key = tavily_api_key or tool_args.get("tavily_api_key", "")
        if tavily_key:
            tv = _web_search_tavily(query, tavily_key, max_res)
            if tv.get("results"):
                all_results.extend(tv["results"])
                sources_used.append("Tavily")
            if tv.get("answer") and not answer:
                answer = tv["answer"]

        # 3. Firecrawl Search (deep content, if key available)
        fc_key = firecrawl_api_key or tool_args.get("firecrawl_api_key", "")
        if fc_key:
            fc = _firecrawl_search(query, fc_key, min(max_res, 5))
            if fc.get("results"):
                all_results.extend(fc["results"])
                sources_used.append("Firecrawl")

        # 4. DuckDuckGo fallback (always available, no key needed)
        ddg = _web_search_duckduckgo(query, max_res)
        if ddg.get("results"):
            for r in ddg["results"]:
                r["source"] = "duckduckgo"
            all_results.extend(ddg["results"])
            if not sources_used:
                sources_used.append("DuckDuckGo")

        # ── Smart deduplication by domain+path ──
        seen_urls = set()
        unique = []
        for r in all_results:
            url = r.get("url", "")
            # Normalize URL for dedup: strip protocol, trailing slash, query params
            norm = re.sub(r'^https?://(www\.)?', '', url).rstrip('/').split('?')[0].split('#')[0].lower()
            if norm and norm not in seen_urls:
                seen_urls.add(norm)
                unique.append(r)
        # Also dedup by title similarity
        seen_titles = set()
        deduped = []
        for r in unique:
            title_key = re.sub(r'[^a-z0-9]+', ' ', r.get("title", "").lower()).strip()[:60]
            if title_key and title_key not in seen_titles:
                seen_titles.add(title_key)
                deduped.append(r)
            elif not title_key:
                deduped.append(r)

        results = deduped[:max_res * 2]
        src_str = " + ".join(sources_used) if sources_used else "DuckDuckGo"
        summary = f"Found {len(results)} web results for '{query}' via {src_str}."
        if results:
            summary += " Top: " + "; ".join([f"{r['title']}" for r in results[:3]])
        if answer:
            summary = answer + "\n\n" + summary
        return {"tool_name": "WEB_SEARCH", "type": "text", "data": results, "summary": summary}

    elif tool_name == "NOVELTY_CHECK":
        idea = tool_args.get("idea", "")
        # ARC pattern: keyword extraction + Jaccard similarity + multi-source search
        _STOP = frozenset({"a","an","the","and","or","but","in","on","of","for","to","with","by","at","from","as","is","are","was","were","be","been","have","has","had","do","does","did","will","would","could","should","may","might","not","using","based","via","new","novel","approach","method","study","research","paper","proposed","results"})
        idea_keywords = [t for t in re.findall(r'[a-zA-Z][a-zA-Z0-9_-]+', idea.lower()) if t not in _STOP and len(t) >= 3]
        idea_kw_set = set(idea_keywords)

        # Search multiple sources (ARC: multi-query)
        all_papers = []
        all_papers.extend(_search_openalex(idea, 10))
        time.sleep(0.5)
        all_papers.extend(_search_semantic_scholar(idea, 10))
        time.sleep(0.5)
        # Build keyword query from top keywords
        if len(idea_keywords) >= 3:
            kw_query = " ".join(idea_keywords[:5])
            all_papers.extend(_search_semantic_scholar(kw_query, 5))

        # Deduplicate
        seen = set()
        unique = []
        for p in all_papers:
            key = p["title"].lower().strip()[:80]
            if key not in seen:
                seen.add(key)
                unique.append(p)

        # Compute similarity scores (ARC Jaccard pattern)
        similar = []
        for p in unique:
            p_kw = set(re.findall(r'[a-zA-Z][a-zA-Z0-9_-]+', (p.get("title","") + " " + p.get("abstract","")).lower())) - _STOP
            if not p_kw or not idea_kw_set:
                sim = 0.0
            else:
                sim = len(idea_kw_set & p_kw) / max(len(idea_kw_set | p_kw), 1)
            if sim >= 0.15:  # ARC threshold
                similar.append({**p, "similarity": round(sim, 3)})
        similar.sort(key=lambda x: x["similarity"], reverse=True)
        similar = similar[:15]

        # Compute novelty score (ARC pattern)
        if not similar:
            novelty_score, assessment = 1.0, "high"
        else:
            top = similar[:5]
            max_sim = max(p["similarity"] for p in top)
            high_cite = sum(1 for p in top if p["similarity"] >= 0.3 and p.get("citations", 0) >= 50)
            raw = 1.0 - max_sim
            if high_cite >= 2:
                raw *= 0.7
            novelty_score = round(max(0.0, min(1.0, raw)), 3)
            if novelty_score >= 0.7: assessment = "high"
            elif novelty_score >= 0.45: assessment = "moderate"
            elif novelty_score >= 0.25: assessment = "low"
            else: assessment = "critical"

        recommendation = "proceed" if assessment == "high" else ("differentiate" if assessment in ("moderate","low") else "abort")
        summary = f"Novelty score: {novelty_score} ({assessment}) — Recommendation: {recommendation}\n"
        summary += f"Found {len(similar)} potentially overlapping papers (of {len(unique)} searched).\n"
        for p in similar[:5]:
            summary += f"  - \"{p['title']}\" ({p.get('year','?')}, sim={p['similarity']}, {p.get('citations',0)} cites)\n"

        return {"tool_name": "NOVELTY_CHECK", "type": "text",
                "data": {"novelty_score": novelty_score, "assessment": assessment, "recommendation": recommendation, "similar_papers": similar[:10]},
                "summary": summary}

    elif tool_name == "CITATION_VERIFY":
        citations = tool_args.get("citations", [])
        if isinstance(citations, str):
            citations = [c.strip() for c in citations.split(";") if c.strip()]
        # Use the full 3-layer verification system
        vr = verify_citations({"citations": citations})
        if "error" in vr:
            return {"tool_name": "CITATION_VERIFY", "type": "text", "data": None, "summary": vr["error"]}
        s = vr.get("summary", {})
        summary = (f"Citation verification: {s.get('verified',0)} verified, {s.get('suspicious',0)} suspicious, "
                   f"{s.get('hallucinated',0)} hallucinated (integrity: {s.get('integrity_score',0):.0%})")
        return {"tool_name": "CITATION_VERIFY", "type": "text", "data": vr.get("results", []), "summary": summary}

    elif tool_name == "PDF_EXTRACT":
        path = tool_args.get("path", "")
        if not path:
            return {"tool_name": "PDF_EXTRACT", "type": "text", "data": None, "summary": "No PDF path provided."}
        pdf_result = extract_pdf_text({"path": path})
        if "error" in pdf_result:
            return {"tool_name": "PDF_EXTRACT", "type": "text", "data": None, "summary": pdf_result["error"]}
        text = pdf_result.get("text", "")[:3000]
        summary = f"Extracted {pdf_result.get('pages', 0)} pages ({len(text)} chars) from PDF."
        return {"tool_name": "PDF_EXTRACT", "type": "text", "data": {"text": text, "pages": pdf_result.get("pages", 0)}, "summary": summary}

    elif tool_name == "EXPERIMENT":
        code = tool_args.get("code", "")
        if not code:
            return {"tool_name": "EXPERIMENT", "type": "code", "data": None, "summary": "No code provided."}
        exp_result = run_experiment_action({"code": code, "timeout_sec": tool_args.get("timeout_sec", 30)})
        if exp_result.get("ok"):
            summary = f"Experiment completed (exit code {exp_result.get('exit_code', -1)}).\nstdout: {exp_result.get('stdout', '')[:500]}"
            if exp_result.get("stderr"):
                summary += f"\nstderr: {exp_result['stderr'][:300]}"
        else:
            summary = f"Experiment failed: {exp_result.get('stderr', exp_result.get('error', 'Unknown error'))}"
        return {"tool_name": "EXPERIMENT", "type": "code", "data": exp_result, "summary": summary}

    # ── V5.6 Pipeline Tools ──

    elif tool_name == "PUBMED_SEARCH":
        query = tool_args.get("query", "")
        max_results = tool_args.get("max_results", 20)
        papers = _search_pubmed(query, max_results)
        summary = f"PubMed: Found {len(papers)} papers for '{query}'."
        if papers:
            summary += " Top: " + "; ".join([f"\"{p['title'][:60]}\" ({p.get('year','?')})" for p in papers[:3]])
        return {"tool_name": "PUBMED_SEARCH", "type": "papers", "data": papers, "summary": summary}

    elif tool_name == "SCIHUB_FETCH":
        doi = tool_args.get("doi", "")
        if not doi:
            return {"tool_name": "SCIHUB_FETCH", "type": "text", "data": None, "summary": "No DOI provided."}
        # Try Unpaywall first (legal OA), then Sci-Hub
        result = _fetch_unpaywall(doi)
        if not result.get("ok"):
            result = _fetch_scihub(doi)
        if result.get("ok"):
            summary = f"Downloaded paper (DOI: {doi}) → {result['path']} ({result.get('size', 0) // 1024}KB)"
        else:
            summary = f"Could not download DOI {doi}: {result.get('error', 'unknown')}"
        return {"tool_name": "SCIHUB_FETCH", "type": "text", "data": result, "summary": summary}

    elif tool_name == "VALIDATE_QUERY":
        vq = validate_research_query({**tool_args, "api_key": api_key, "provider": provider, "model": model})
        if vq.get("ok"):
            valid = "VALID" if vq.get("is_valid") else "INVALID"
            summary = f"Query validation: {valid}. {vq.get('reason', '')[:200]}"
            if vq.get("keywords"):
                summary += f"\nKeywords: {', '.join(vq['keywords'][:8])}"
        else:
            summary = vq.get("error", "Validation failed")
        return {"tool_name": "VALIDATE_QUERY", "type": "text", "data": vq, "summary": summary}

    elif tool_name == "MESH_QUERIES":
        mq = generate_search_queries({**tool_args, "api_key": api_key, "provider": provider, "model": model})
        if mq.get("ok"):
            queries = mq.get("queries", [])
            summary = f"Generated {len(queries)} search queries:\n"
            for q in queries[:5]:
                summary += f"  [{q.get('db','')}] {q.get('query_string','')[:80]}\n"
        else:
            summary = mq.get("error", "Query generation failed")
        return {"tool_name": "MESH_QUERIES", "type": "text", "data": mq, "summary": summary}

    elif tool_name == "TRIAGE":
        tr = triage_papers({**tool_args, "api_key": api_key, "provider": provider, "model": model})
        if tr.get("ok"):
            summary = f"Screened {tr.get('total_screened', 0)} papers: {tr.get('relevant_count', 0)} relevant."
        else:
            summary = tr.get("error", "Triage failed")
        return {"tool_name": "TRIAGE", "type": "text", "data": tr, "summary": summary}

    elif tool_name == "DRAFT_SECTION":
        ds = draft_research_section({**tool_args, "api_key": api_key, "provider": provider, "model": model})
        if ds.get("ok"):
            summary = f"Drafted '{ds.get('section_type', 'section')}' ({len(ds.get('text', ''))} chars, {ds.get('citations_used', 0)} citations)"
        else:
            summary = ds.get("error", "Drafting failed")
        return {"tool_name": "DRAFT_SECTION", "type": "text", "data": ds, "summary": summary}

    elif tool_name == "GENERATE_CHART":
        cr = generate_chart(tool_args)
        if cr.get("ok"):
            summary = f"Generated {cr.get('chart_type', 'chart')} → {cr.get('path', '')}"
        else:
            summary = cr.get("error", "Chart generation failed")
        return {"tool_name": "GENERATE_CHART", "type": "text", "data": cr, "summary": summary}

    elif tool_name == "GENERATE_TABLE":
        tr = generate_table(tool_args)
        if tr.get("ok"):
            summary = f"Table generated: {tr.get('title', 'Table')}\n{tr.get('markdown', '')}"
        else:
            summary = tr.get("error", "Table generation failed")
        return {"tool_name": "GENERATE_TABLE", "type": "table", "data": tr, "summary": summary}

    # ── V6.1 Free Data Source Tools ──────────────────────────────────────────

    elif tool_name == "EUROPE_PMC_SEARCH":
        papers = _search_europe_pmc(tool_args.get("query", ""), tool_args.get("max_results", 20))
        summary = f"Europe PMC: Found {len(papers)} papers."
        if papers:
            summary += " Top: " + "; ".join(f"\"{p['title'][:50]}\" ({p.get('year','?')})" for p in papers[:3])
        return {"tool_name": "EUROPE_PMC_SEARCH", "type": "papers", "data": papers, "summary": summary}

    elif tool_name == "CLINICAL_TRIALS_SEARCH":
        trials = _search_clinical_trials(
            tool_args.get("condition", tool_args.get("query", "")),
            tool_args.get("intervention", ""),
            tool_args.get("status", "COMPLETED"),
            tool_args.get("max_results", 20),
        )
        summary = f"ClinicalTrials.gov: Found {len(trials)} trials."
        if trials:
            summary += " Top: " + "; ".join(f"\"{t['title'][:50]}\" ({t.get('phase',[])})" for t in trials[:3])
        return {"tool_name": "CLINICAL_TRIALS_SEARCH", "type": "text", "data": trials, "summary": summary}

    elif tool_name == "OPENFDA_ADVERSE_EVENTS":
        result = _query_openfda_adverse_events(tool_args.get("drug_name", ""), tool_args.get("limit", 50))
        top_ae = result.get("adverse_events", [])[:5]
        summary = f"OpenFDA AE for '{result.get('drug','')}': {result.get('total',0)} event types."
        if top_ae:
            summary += " Top AEs: " + ", ".join(f"{ae['reaction']} ({ae['count']})" for ae in top_ae)
        return {"tool_name": "OPENFDA_ADVERSE_EVENTS", "type": "text", "data": result, "summary": summary}

    elif tool_name == "OPENFDA_DRUG_LABELS":
        result = _query_openfda_drug_labels(tool_args.get("drug_name", ""))
        labels = result.get("labels", [])
        summary = f"OpenFDA labels for '{result.get('drug','')}': {len(labels)} label(s) found."
        if labels:
            summary += f" Indications: {labels[0].get('indications','')[:200]}"
        return {"tool_name": "OPENFDA_DRUG_LABELS", "type": "text", "data": result, "summary": summary}

    elif tool_name == "MESH_LOOKUP":
        result = _lookup_mesh_terms(tool_args.get("term", ""))
        summary = f"MeSH lookup: {len(result.get('mesh_terms',[]))} terms, {len(result.get('synonyms',[]))} synonyms for '{result.get('term','')}'"
        if result.get("mesh_terms"):
            summary += "\nTerms: " + ", ".join(result["mesh_terms"][:8])
        return {"tool_name": "MESH_LOOKUP", "type": "text", "data": result, "summary": summary}

    elif tool_name == "RXNORM_LOOKUP":
        result = _lookup_rxnorm(tool_args.get("drug_name", ""))
        summary = f"RxNorm: '{result.get('drug','')}' → RxCUI {result.get('rxcui','N/A')}, {len(result.get('brand_names',[]))} brand names"
        if result.get("drug_classes"):
            summary += f"\nClasses: {', '.join(result['drug_classes'][:4])}"
        return {"tool_name": "RXNORM_LOOKUP", "type": "text", "data": result, "summary": summary}

    elif tool_name == "RETRACTION_CHECK":
        dois = tool_args.get("dois", [])
        if not dois and tool_args.get("doi"):
            dois = [tool_args["doi"]]
        results = [_check_retraction_watch(doi) for doi in dois[:20]]
        retracted = [r for r in results if r.get("retracted")]
        summary = f"Retraction check: {len(dois)} DOIs checked, {len(retracted)} retracted."
        if retracted:
            summary += " RETRACTED: " + "; ".join(r["doi"] for r in retracted)
        return {"tool_name": "RETRACTION_CHECK", "type": "text", "data": results, "summary": summary}

    # ── V6.1 Statistical Analysis Tools ──────────────────────────────────────

    elif tool_name == "META_ANALYSIS":
        result = run_meta_analysis(tool_args)
        if result.get("ok"):
            summary = (f"Meta-analysis ({result.get('method','?')}): k={result['k']}, "
                       f"pooled={result['pooled_effect']} [{result['ci_lower']}, {result['ci_upper']}], "
                       f"I²={result['i_squared']}%, p={result['p_value']}")
        else:
            summary = result.get("error", "Meta-analysis failed")
        return {"tool_name": "META_ANALYSIS", "type": "text", "data": result, "summary": summary}

    elif tool_name == "FOREST_PLOT":
        result = generate_forest_plot(tool_args)
        if result.get("ok"):
            summary = f"Forest plot generated ({len(tool_args.get('studies', []))} studies)"
        else:
            summary = result.get("error", "Forest plot failed")
        return {"tool_name": "FOREST_PLOT", "type": "text", "data": result, "summary": summary}

    elif tool_name == "FUNNEL_PLOT":
        result = generate_funnel_plot(tool_args)
        if result.get("ok"):
            summary = f"Funnel plot generated. Egger's p={result.get('egger_p','N/A')} ({'asymmetric' if result.get('asymmetric') else 'symmetric'})"
        else:
            summary = result.get("error", "Funnel plot failed")
        return {"tool_name": "FUNNEL_PLOT", "type": "text", "data": result, "summary": summary}

    elif tool_name == "PRISMA_FLOWCHART":
        result = generate_prisma_flowchart(tool_args)
        if result.get("ok"):
            summary = f"PRISMA 2020 flowchart generated ({tool_args.get('studies_included','?')} included studies)"
        else:
            summary = result.get("error", "PRISMA flowchart failed")
        return {"tool_name": "PRISMA_FLOWCHART", "type": "text", "data": result, "summary": summary}

    elif tool_name == "ROB_PLOT":
        result = generate_rob_plot(tool_args)
        if result.get("ok"):
            summary = f"Risk of Bias plot generated ({len(tool_args.get('studies',[]))} studies)"
        else:
            summary = result.get("error", "RoB plot failed")
        return {"tool_name": "ROB_PLOT", "type": "text", "data": result, "summary": summary}

    elif tool_name == "GRADE_ASSESS":
        result = calculate_grade(tool_args)
        if result.get("ok"):
            summary = f"GRADE: {result['outcome']} → {result['certainty'].upper()} {result['grade_symbol']}"
            if result.get("rationale"):
                summary += "\n" + "\n".join(result["rationale"])
        else:
            summary = result.get("error", "GRADE assessment failed")
        return {"tool_name": "GRADE_ASSESS", "type": "text", "data": result, "summary": summary}

    elif tool_name == "PICO_EXTRACT":
        result = extract_pico_structured({**tool_args, "api_key": api_key, "provider": provider, "model": model})
        if result.get("ok"):
            summary = f"PICO extracted: P={result.get('population','?')[:50]}, I={result.get('intervention','?')[:50]}"
        else:
            summary = result.get("error", "PICO extraction failed")
        return {"tool_name": "PICO_EXTRACT", "type": "text", "data": result, "summary": summary}

    else:
        return {"tool_name": tool_name, "type": "text", "data": None, "summary": f"Unknown tool: {tool_name}"}


def search_papers(args):
    """Search academic papers across multiple sources.
    Args: {query, max_results, sources, year_min}
    Returns: {ok, papers: [...]}
    """
    query = args.get("query", "")
    max_results = args.get("max_results", 10)
    sources = args.get("sources", ["arxiv", "semantic_scholar", "openalex"])
    year_min = args.get("year_min", None)

    if not query:
        return {"error": "Query is required."}

    # Source order matches ARC: OpenAlex (10K/day) → S2 (1K/5min) → arXiv (1/3s)
    all_papers = []
    if "openalex" in sources:
        all_papers.extend(_search_openalex(query, max_results))
        time.sleep(0.5)
    if "semantic_scholar" in sources:
        all_papers.extend(_search_semantic_scholar(query, max_results, year_min))
        time.sleep(1.0)
    if "arxiv" in sources:
        all_papers.extend(_search_arxiv(query, max_results))

    # Deduplicate
    seen = set()
    unique = []
    for p in all_papers:
        key = p["title"].lower().strip()[:80]
        if key not in seen:
            seen.add(key)
            unique.append(p)

    unique.sort(key=lambda x: x.get("citations", 0), reverse=True)
    return {"ok": True, "papers": unique[:max_results * 2], "total": len(unique)}


def web_search_action(args):
    """Web search action.
    Args: {query, max_results}
    Returns: {ok, results: [...]}
    """
    query = args.get("query", "")
    max_results = args.get("max_results", 5)

    if not query:
        return {"error": "Query is required."}

    result = _web_search_duckduckgo(query, max_results)
    return {"ok": True, "results": result.get("results", [])}


def extract_pdf_text(args):
    """Extract text from a PDF file.
    Args: {path}
    Returns: {ok, text, pages, metadata}
    """
    path = args.get("path", "")
    if not path or not os.path.isfile(path):
        return {"error": f"File not found: {path}"}

    text = ""
    pages = 0
    metadata = {}

    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = len(pdf.pages)
            metadata = pdf.metadata or {}
            text_parts = []
            for page in pdf.pages[:50]:  # limit to 50 pages
                t = page.extract_text() or ""
                text_parts.append(t)
            text = "\n\n".join(text_parts)
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            pages = len(reader.pages)
            metadata = dict(reader.metadata) if reader.metadata else {}
            text_parts = []
            for page in reader.pages[:50]:
                t = page.extract_text() or ""
                text_parts.append(t)
            text = "\n\n".join(text_parts)
        except ImportError:
            return {"error": "PDF extraction requires pdfplumber or PyPDF2. Install via: pip install pdfplumber"}
        except Exception as e:
            return {"error": f"PDF read error: {e}"}
    except Exception as e:
        return {"error": f"PDF read error: {e}"}

    # Sanitize metadata for JSON
    safe_meta = {}
    for k, v in metadata.items():
        try:
            json.dumps(v)
            safe_meta[k] = v
        except (TypeError, ValueError):
            safe_meta[k] = str(v)

    return {"ok": True, "text": text[:100000], "pages": pages, "metadata": safe_meta}


def check_novelty(args):
    """Assess novelty of a research idea.
    Args: {idea, papers (optional), api_key, provider, model}
    Returns: {ok, score, assessment, similar_papers, gaps}
    """
    idea = args.get("idea", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "openai")
    model = args.get("model", "")

    if not idea:
        return {"error": "Research idea is required."}
    if not api_key:
        return {"error": "API key required."}

    # Search for related papers
    related = _search_semantic_scholar(idea, 10)
    related.extend(_search_arxiv(idea, 5))

    # Deduplicate
    seen = set()
    unique = []
    for p in related:
        key = p["title"].lower().strip()[:80]
        if key not in seen:
            seen.add(key)
            unique.append(p)
    unique.sort(key=lambda x: x.get("citations", 0), reverse=True)
    top_papers = unique[:10]

    # Ask LLM to assess novelty
    papers_context = "\n".join([
        f"- \"{p['title']}\" ({p.get('year','?')}, {p.get('citations',0)} citations)"
        for p in top_papers
    ])

    prompt = (
        f"Research idea: {idea}\n\n"
        f"Related existing papers:\n{papers_context}\n\n"
        "Assess the novelty of this research idea on a scale of 1-10. Provide:\n"
        "1. Novelty score (1-10)\n"
        "2. Brief assessment of what's novel vs what exists\n"
        "3. Identified gaps this idea could fill\n"
        "4. Suggestions to increase novelty\n"
        "Format as JSON: {\"score\": N, \"assessment\": \"...\", \"gaps\": [\"...\"], \"suggestions\": [\"...\"]}"
    )

    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": "You are a research novelty assessor. Be honest and constructive."},
                        {"role": "user", "content": prompt}],
                       temperature=0.3, max_tokens=2048)

    if "error" in result:
        return result

    # Try to parse JSON from response
    text = result["text"]
    try:
        # Extract JSON from response
        json_match = re.search(r'\{[^{}]*"score"[^{}]*\}', text, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            return {
                "ok": True,
                "score": parsed.get("score", 5),
                "assessment": parsed.get("assessment", text),
                "gaps": parsed.get("gaps", []),
                "suggestions": parsed.get("suggestions", []),
                "similar_papers": top_papers[:5],
                "tokens": result.get("usage"),
            }
    except (json.JSONDecodeError, AttributeError):
        pass

    return {
        "ok": True, "score": 5, "assessment": text,
        "gaps": [], "suggestions": [],
        "similar_papers": top_papers[:5],
        "tokens": result.get("usage"),
    }


# ── Title similarity (mirrors ARC verify.py) ──────────────────────────────

def _title_similarity(a, b):
    """Word-overlap Jaccard-ish similarity between two titles (0.0-1.0)."""
    def _words(t):
        return set(re.sub(r'[^a-z0-9\s]', '', t.lower()).split()) - {''}
    wa, wb = _words(a), _words(b)
    if not wa or not wb:
        return 0.0
    return len(wa & wb) / max(len(wa), len(wb))


def _verify_by_doi(doi, expected_title):
    """L2: Verify DOI via CrossRef API, with DataCite fallback for arXiv DOIs (ARC pattern)."""
    encoded_doi = urllib.parse.quote(doi, safe='')
    url = f"https://api.crossref.org/works/{encoded_doi}"
    try:
        req = urllib.request.Request(url, headers={
            "User-Agent": "Zenith/1.0 (mailto:zenith@example.com)", "Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=20) as resp:
            body = json.loads(resp.read().decode())
        titles = body.get("message", {}).get("title", [])
        found_title = titles[0] if titles else ""
        if not found_title:
            return {"status": "verified", "confidence": 0.85, "method": "doi", "details": f"DOI {doi} resolves via CrossRef"}
        sim = _title_similarity(expected_title, found_title)
        if sim >= 0.80:
            return {"status": "verified", "confidence": sim, "method": "doi", "details": f"Confirmed via CrossRef: '{found_title}'"}
        elif sim >= 0.50:
            return {"status": "suspicious", "confidence": sim, "method": "doi", "details": f"DOI resolves but title differs (sim={sim:.2f})"}
        else:
            return {"status": "suspicious", "confidence": sim, "method": "doi", "details": f"DOI resolves but title mismatch (sim={sim:.2f})"}
    except urllib.error.HTTPError as e:
        if e.code == 404:
            # Try DataCite for arXiv DOIs
            if doi.startswith("10.48550/") or doi.startswith("10.5281/"):
                return _verify_by_datacite(doi, expected_title)
            return {"status": "hallucinated", "confidence": 0.9, "method": "doi", "details": f"DOI {doi} not found (HTTP 404)"}
        return None
    except Exception:
        return None


def _verify_by_datacite(doi, expected_title):
    """DataCite fallback for arXiv/Zenodo DOIs (ARC pattern)."""
    try:
        encoded = urllib.parse.quote(doi, safe='')
        url = f"https://api.datacite.org/dois/{encoded}"
        req = urllib.request.Request(url, headers={"User-Agent": "Zenith/1.0", "Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            body = json.loads(resp.read().decode())
        dc_titles = body.get("data", {}).get("attributes", {}).get("titles", [])
        found_title = dc_titles[0].get("title", "") if dc_titles else ""
        if not found_title:
            return {"status": "verified", "confidence": 0.85, "method": "doi", "details": f"DOI {doi} resolves via DataCite"}
        sim = _title_similarity(expected_title, found_title)
        if sim >= 0.80:
            return {"status": "verified", "confidence": sim, "method": "doi", "details": f"Confirmed via DataCite: '{found_title}'"}
        return {"status": "suspicious", "confidence": sim, "method": "doi", "details": f"DataCite title differs (sim={sim:.2f})"}
    except Exception:
        return None


def _verify_by_openalex(title):
    """L3a: Verify via OpenAlex title search (10K+/day, ARC pattern)."""
    try:
        params = urllib.parse.urlencode({"filter": "title.search:" + title.replace(",", " ").replace(":", " "),
                                         "per_page": "5", "mailto": "zenith@example.com"})
        url = f"https://api.openalex.org/works?{params}"
        req = urllib.request.Request(url, headers={"User-Agent": "Zenith/1.0 (mailto:zenith@example.com)", "Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            body = json.loads(resp.read().decode())
        results = body.get("results", [])
        if not results:
            return {"status": "hallucinated", "confidence": 0.7, "method": "openalex", "details": "No results found via OpenAlex"}
        best_sim, best_title = 0.0, ""
        for r in results:
            ft = r.get("title", "")
            if ft:
                sim = _title_similarity(title, ft)
                if sim > best_sim:
                    best_sim, best_title = sim, ft
        if best_sim >= 0.80:
            return {"status": "verified", "confidence": best_sim, "method": "openalex", "details": f"Confirmed via OpenAlex: '{best_title}'"}
        elif best_sim >= 0.50:
            return {"status": "suspicious", "confidence": best_sim, "method": "openalex", "details": f"Partial match via OpenAlex (sim={best_sim:.2f})"}
        return {"status": "hallucinated", "confidence": 0.7, "method": "openalex", "details": "No close match found via OpenAlex"}
    except Exception:
        return None


def verify_citations(args):
    """Verify citations using ARC's 3-layer system: DOI/CrossRef → OpenAlex → S2+arXiv title search.
    Args: {citations: [str] or [{title, doi, arxiv_id}]}
    Returns: {ok, results: [{ref, status, confidence, method, details, doi, url}]}
    """
    citations = args.get("citations", [])
    if isinstance(citations, str):
        citations = [c.strip() for c in citations.split("\n") if c.strip()]

    if not citations:
        return {"error": "No citations to verify."}

    # Adaptive delays (ARC pattern)
    _DELAY_CROSSREF = 0.3
    _DELAY_OPENALEX = 0.2
    _DELAY_S2 = 1.5
    api_calls = 0
    _start = time.time()
    _TIMEOUT = 300  # 5 min global timeout (ARC BUG-22 fix)

    results = []
    for i, cite in enumerate(citations[:20]):
        # Global timeout check
        if time.time() - _start > _TIMEOUT:
            for remaining in citations[i:]:
                ref = remaining if isinstance(remaining, str) else remaining.get("title", str(remaining))
                results.append({"ref": ref, "status": "skipped", "confidence": 0, "method": "timeout", "details": "Verification timeout exceeded"})
            break

        # Normalize input
        if isinstance(cite, dict):
            title = cite.get("title", "")
            doi = cite.get("doi", "")
            arxiv_id = cite.get("arxiv_id", "")
            ref_str = title or str(cite)
        else:
            title = cite
            doi, arxiv_id = "", ""
            ref_str = cite

        if not title:
            results.append({"ref": ref_str, "status": "skipped", "confidence": 0, "method": "skipped", "details": "No title provided"})
            continue

        result = None

        # L2: DOI via CrossRef (fast, generous limits)
        if result is None and doi:
            if api_calls > 0:
                time.sleep(_DELAY_CROSSREF)
            result = _verify_by_doi(doi, title)
            api_calls += 1

        # L3a: OpenAlex title search (10K/day)
        if result is None:
            if api_calls > 0:
                time.sleep(_DELAY_OPENALEX)
            result = _verify_by_openalex(title)
            api_calls += 1

        # L3b: S2 title search (last resort)
        if result is None:
            if api_calls > 0:
                time.sleep(_DELAY_S2)
            found = _search_semantic_scholar(title, 3)
            api_calls += 1
            if found:
                best_sim, best = 0.0, None
                for p in found:
                    sim = _title_similarity(title, p.get("title", ""))
                    if sim > best_sim:
                        best_sim, best = sim, p
                if best_sim >= 0.80:
                    result = {"status": "verified", "confidence": best_sim, "method": "title_search",
                              "details": f"Found via S2: '{best['title']}'"}
                elif best_sim >= 0.50:
                    result = {"status": "suspicious", "confidence": best_sim, "method": "title_search",
                              "details": f"Partial match via S2 (sim={best_sim:.2f})"}
                else:
                    result = {"status": "hallucinated", "confidence": 1.0 - best_sim, "method": "title_search",
                              "details": "No close match found"}
            else:
                result = {"status": "hallucinated", "confidence": 0.7, "method": "title_search",
                          "details": "No results found via S2 + arXiv"}

        # Fallback: all layers failed
        if result is None:
            result = {"status": "skipped", "confidence": 0, "method": "skipped", "details": "All verification methods failed"}

        results.append({"ref": ref_str, "doi": doi, **result})

    verified = sum(1 for r in results if r["status"] == "verified")
    suspicious = sum(1 for r in results if r["status"] == "suspicious")
    hallucinated = sum(1 for r in results if r["status"] == "hallucinated")
    integrity = verified / max(1, len(results) - sum(1 for r in results if r["status"] == "skipped")) if results else 1.0

    return {
        "ok": True, "results": results,
        "summary": {"total": len(results), "verified": verified, "suspicious": suspicious,
                     "hallucinated": hallucinated, "integrity_score": round(integrity, 3)},
    }


def run_experiment_action(args):
    """Run Python code in a sandboxed subprocess.
    Args: {code, timeout_sec, packages}
    Returns: {ok, stdout, stderr, exit_code}
    """
    code = args.get("code", "")
    timeout = min(args.get("timeout_sec", 30), 60)  # max 60s

    if not code:
        return {"error": "No code provided."}

    # Write code to temp file
    import uuid
    exp_id = uuid.uuid4().hex[:8]
    exp_dir = os.path.join(EXPERIMENTS_DIR, f"exp_{exp_id}")
    os.makedirs(exp_dir, exist_ok=True)
    script_path = os.path.join(exp_dir, "run.py")

    with open(script_path, "w", encoding="utf-8") as f:
        f.write(code)

    try:
        import subprocess
        result = subprocess.run(
            [sys.executable, script_path],
            capture_output=True, text=True, timeout=timeout,
            cwd=exp_dir, env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"}
        )
        return {
            "ok": True, "stdout": result.stdout[:10000],
            "stderr": result.stderr[:5000], "exit_code": result.returncode,
            "experiment_dir": exp_dir,
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "stdout": "", "stderr": f"Experiment timed out after {timeout}s", "exit_code": -1}
    except Exception as e:
        return {"ok": False, "stdout": "", "stderr": str(e), "exit_code": -1}


def export_chat(args):
    """Export a research conversation to file.
    Args: {messages, format, thread_title}
    Returns: {ok, path, size}
    """
    messages = args.get("messages", [])
    fmt = args.get("format", "markdown")
    title = args.get("thread_title", "Research Export")

    if not messages:
        return {"error": "No messages to export."}

    timestamp = time.strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:50]

    if fmt == "markdown":
        content = f"# {title}\n\n*Exported from Zenith Research — {time.strftime('%Y-%m-%d %H:%M')}*\n\n---\n\n"
        for m in messages:
            role = m.get("role", "user").capitalize()
            text = m.get("content", "")
            if role == "User":
                content += f"## You\n\n{text}\n\n"
            elif role == "Assistant":
                content += f"## Assistant\n\n{text}\n\n"
            elif role == "Tool":
                content += f"### Tool Result\n\n{text}\n\n"
            content += "---\n\n"
        ext = ".md"

    elif fmt == "json":
        content = json.dumps({"title": title, "exported_at": time.strftime('%Y-%m-%dT%H:%M:%S'), "messages": messages}, indent=2, ensure_ascii=False)
        ext = ".json"

    elif fmt == "latex":
        content = "\\documentclass{article}\n\\usepackage[utf8]{inputenc}\n\\usepackage{hyperref}\n"
        content += f"\\title{{{title}}}\n\\date{{\\today}}\n\\begin{{document}}\n\\maketitle\n\n"
        for m in messages:
            role = m.get("role", "user")
            text = m.get("content", "").replace("&", "\\&").replace("%", "\\%").replace("#", "\\#").replace("_", "\\_")
            if role == "user":
                content += f"\\subsection*{{User}}\n{text}\n\n"
            elif role == "assistant":
                content += f"\\subsection*{{Assistant}}\n{text}\n\n"
        content += "\\end{document}\n"
        ext = ".tex"

    elif fmt == "bibtex":
        # Extract paper references from messages
        content = f"% BibTeX export from Zenith Research — {title}\n% Generated {time.strftime('%Y-%m-%d')}\n\n"
        entry_count = 0
        for m in messages:
            data = m.get("data")
            if data and isinstance(data, list):
                for paper in data:
                    if isinstance(paper, dict) and paper.get("title"):
                        entry_count += 1
                        key = re.sub(r'[^a-zA-Z0-9]', '', (paper.get("authors", ["unknown"])[0] if paper.get("authors") else "unknown"))
                        key += str(paper.get("year", ""))
                        content += f"@article{{{key}{entry_count},\n"
                        content += f"  title = {{{paper.get('title', '')}}},\n"
                        if paper.get("authors"):
                            content += f"  author = {{{' and '.join(paper['authors'][:5])}}},\n"
                        if paper.get("year"):
                            content += f"  year = {{{paper['year']}}},\n"
                        if paper.get("doi"):
                            content += f"  doi = {{{paper['doi']}}},\n"
                        if paper.get("url"):
                            content += f"  url = {{{paper['url']}}},\n"
                        content += "}\n\n"
        if entry_count == 0:
            content += "% No paper references found in this conversation.\n"
        ext = ".bib"

    elif fmt == "pdf":
        # Full markdown-to-PDF via reportlab with proper rendering
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Preformatted,
                Table, TableStyle, HRFlowable,
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm, mm
            from reportlab.lib.colors import HexColor
            from reportlab.lib.enums import TA_LEFT

            out_path = os.path.join(EXPORTS_DIR, f"{safe_title}_{timestamp}.pdf")
            doc = SimpleDocTemplate(
                out_path, pagesize=A4,
                leftMargin=2 * cm, rightMargin=2 * cm,
                topMargin=2 * cm, bottomMargin=2 * cm,
            )
            styles = getSampleStyleSheet()

            # ── Custom styles ──
            styles.add(ParagraphStyle(
                "MDH1", parent=styles["Heading1"], fontSize=18, spaceAfter=10, spaceBefore=14,
                textColor=HexColor("#1a1a2e"),
            ))
            styles.add(ParagraphStyle(
                "MDH2", parent=styles["Heading2"], fontSize=15, spaceAfter=8, spaceBefore=12,
                textColor=HexColor("#1a1a2e"),
            ))
            styles.add(ParagraphStyle(
                "MDH3", parent=styles["Heading3"], fontSize=13, spaceAfter=6, spaceBefore=10,
                textColor=HexColor("#2a2a4e"),
            ))
            styles.add(ParagraphStyle(
                "MDH4", parent=styles["Heading4"], fontSize=11, spaceAfter=4, spaceBefore=8,
                textColor=HexColor("#2a2a4e"),
            ))
            styles.add(ParagraphStyle(
                "CodeBlock", fontName="Courier", fontSize=8, leading=10,
                backColor=HexColor("#f5f5f5"), textColor=HexColor("#333333"),
                leftIndent=12, rightIndent=12, spaceBefore=6, spaceAfter=6,
                borderPadding=(6, 6, 6, 6),
            ))
            styles.add(ParagraphStyle(
                "MDBody", parent=styles["Normal"], fontSize=10, leading=14,
                spaceAfter=4, alignment=TA_LEFT,
            ))
            styles.add(ParagraphStyle(
                "Blockquote", parent=styles["Normal"], fontSize=10, leading=13,
                leftIndent=20, textColor=HexColor("#555555"), fontName="Helvetica-Oblique",
                spaceBefore=4, spaceAfter=4,
            ))
            styles.add(ParagraphStyle(
                "BulletItem", parent=styles["Normal"], fontSize=10, leading=13,
                leftIndent=20, bulletIndent=10, spaceBefore=1, spaceAfter=1,
            ))
            styles.add(ParagraphStyle(
                "RoleLabel", parent=styles["Heading3"], fontSize=11, spaceAfter=4,
                spaceBefore=8, textColor=HexColor("#0891b2"),
            ))

            def _esc(t):
                """Escape HTML entities for reportlab Paragraph markup."""
                return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            def _inline(t):
                """Convert inline markdown to reportlab XML: bold, italic, code, links."""
                t = _esc(t)
                # Bold: **text**
                t = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', t)
                # Italic: *text*
                t = re.sub(r'\*(.+?)\*', r'<i>\1</i>', t)
                # Inline code: `text`
                t = re.sub(r'`(.+?)`', r'<font face="Courier" size="8" color="#c0392b">\1</font>', t)
                # Links: [text](url)
                t = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2" color="blue">\1</a>', t)
                # Bare URLs
                t = re.sub(r'(https?://[^\s<>]+)', r'<a href="\1" color="blue">\1</a>', t)
                return t

            def _md_to_flowables(md_text, _img_dirs=None):
                """Convert markdown text to a list of reportlab flowables.
                _img_dirs: list of directories to search for images."""
                from reportlab.platypus import Image as RLImage
                flowables = []
                lines = md_text.split("\n")
                i = 0
                while i < len(lines):
                    line = lines[i]

                    # Markdown image: ![alt](path)
                    img_m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
                    if img_m:
                        alt_text = img_m.group(1)
                        img_ref = img_m.group(2)
                        img_fname = os.path.basename(img_ref)
                        img_path = None
                        for d in (_img_dirs or []):
                            candidate = os.path.join(d, img_fname)
                            if os.path.isfile(candidate):
                                img_path = candidate
                                break
                        if not img_path:
                            # Try the reference as-is
                            if os.path.isfile(img_ref):
                                img_path = img_ref
                        if img_path:
                            try:
                                img_w = min(A4[0] - 4 * cm, 14 * cm)
                                flowables.append(Spacer(1, 3 * mm))
                                flowables.append(RLImage(img_path, width=img_w, kind="proportional"))
                                if alt_text:
                                    flowables.append(Paragraph(f'<i>{_esc(alt_text)}</i>', styles["MDBody"]))
                                flowables.append(Spacer(1, 3 * mm))
                            except Exception:
                                flowables.append(Paragraph(f'[Image: {_esc(alt_text)}]', styles["MDBody"]))
                        else:
                            flowables.append(Paragraph(f'[Image: {_esc(alt_text)}]', styles["MDBody"]))
                        i += 1
                        continue

                    # Code block
                    if line.strip().startswith("```"):
                        code_lines = []
                        i += 1
                        while i < len(lines) and not lines[i].strip().startswith("```"):
                            code_lines.append(lines[i])
                            i += 1
                        i += 1  # skip closing ```
                        code_text = _esc("\n".join(code_lines))
                        flowables.append(Preformatted(code_text, styles["CodeBlock"]))
                        continue

                    # Headers
                    hm = re.match(r'^(#{1,4})\s+(.+)$', line)
                    if hm:
                        level = len(hm.group(1))
                        style_name = f"MDH{level}"
                        flowables.append(Paragraph(_inline(hm.group(2)), styles[style_name]))
                        i += 1
                        continue

                    # Horizontal rule
                    if re.match(r'^-{3,}$', line.strip()) or re.match(r'^\*{3,}$', line.strip()):
                        flowables.append(HRFlowable(
                            width="100%", thickness=0.5,
                            color=HexColor("#cccccc"), spaceBefore=6, spaceAfter=6,
                        ))
                        i += 1
                        continue

                    # Blockquote
                    bq = re.match(r'^>\s?(.*)$', line)
                    if bq:
                        flowables.append(Paragraph(_inline(bq.group(1)), styles["Blockquote"]))
                        i += 1
                        continue

                    # Table: detect | header | header | lines
                    if line.strip().startswith("|") and "|" in line[1:]:
                        table_rows = []
                        while i < len(lines) and lines[i].strip().startswith("|"):
                            row_line = lines[i].strip()
                            # Skip separator rows like |:---|:---|
                            if re.match(r'^\|[\s:|-]+\|$', row_line):
                                i += 1
                                continue
                            cells = [c.strip() for c in row_line.split("|")[1:-1]]
                            table_rows.append(cells)
                            i += 1
                        if table_rows:
                            # Normalize column count
                            max_cols = max(len(r) for r in table_rows)
                            for r in table_rows:
                                while len(r) < max_cols:
                                    r.append("")
                            # Convert cells to Paragraphs for wrapping
                            pdf_data = []
                            for ri, row in enumerate(table_rows):
                                pdf_row = []
                                for cell in row:
                                    st = styles["MDBody"] if ri > 0 else styles["MDH4"]
                                    pdf_row.append(Paragraph(_inline(cell), st))
                                pdf_data.append(pdf_row)
                            # Build table with styling
                            col_w = (A4[0] - 4 * cm) / max_cols
                            tbl = Table(pdf_data, colWidths=[col_w] * max_cols)
                            tbl.setStyle(TableStyle([
                                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#e8e8e8")),
                                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#1a1a2e")),
                                ("FONTSIZE", (0, 0), (-1, -1), 9),
                                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("TOPPADDING", (0, 0), (-1, -1), 4),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ]))
                            flowables.append(Spacer(1, 4 * mm))
                            flowables.append(tbl)
                            flowables.append(Spacer(1, 4 * mm))
                        continue

                    # Unordered list item
                    ul = re.match(r'^(\s*)[-*]\s+(.+)$', line)
                    if ul:
                        indent = len(ul.group(1)) // 2
                        bullet = "\u2022 " if indent == 0 else "  \u25E6 "
                        st = ParagraphStyle(
                            f"_bullet_{i}", parent=styles["BulletItem"],
                            leftIndent=20 + indent * 14,
                        )
                        flowables.append(Paragraph(bullet + _inline(ul.group(2)), st))
                        i += 1
                        continue

                    # Ordered list item
                    ol = re.match(r'^(\s*)(\d+)[.)]\s+(.+)$', line)
                    if ol:
                        indent = len(ol.group(1)) // 2
                        num = ol.group(2)
                        st = ParagraphStyle(
                            f"_ol_{i}", parent=styles["BulletItem"],
                            leftIndent=20 + indent * 14,
                        )
                        flowables.append(Paragraph(f"{num}. " + _inline(ol.group(3)), st))
                        i += 1
                        continue

                    # Empty line → small spacer
                    if not line.strip():
                        flowables.append(Spacer(1, 3 * mm))
                        i += 1
                        continue

                    # Regular paragraph — accumulate consecutive non-empty lines
                    para_lines = []
                    while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(("#", "```", "|", ">", "---", "***", "- ", "* ")):
                        # Also break on numbered list items
                        if re.match(r'^\s*\d+[.)]\s+', lines[i]):
                            break
                        para_lines.append(lines[i])
                        i += 1
                    if para_lines:
                        flowables.append(Paragraph(_inline(" ".join(para_lines)), styles["MDBody"]))
                    else:
                        # Single line that didn't match anything
                        flowables.append(Paragraph(_inline(line), styles["MDBody"]))
                        i += 1

                return flowables

            # ── Build the PDF ──
            story = []
            story.append(Paragraph(_esc(title), styles["Title"]))
            story.append(Paragraph(
                f'<font size="9" color="#888888">Exported from Zenith Research \u2014 {time.strftime("%Y-%m-%d %H:%M")}</font>',
                styles["Normal"],
            ))
            story.append(Spacer(1, 0.5 * cm))
            story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#0891b2"), spaceBefore=4, spaceAfter=10))

            for m in messages:
                role = m.get("role", "user").capitalize()
                text = m.get("content", "")
                mtype = m.get("type", "text")

                # Role header
                role_label = {"User": "You", "Assistant": "Assistant", "Tool": f"Tool: {m.get('tool_used', 'result')}", "System": "System"}.get(role, role)
                story.append(Paragraph(role_label, styles["RoleLabel"]))

                # Convert full message content via markdown parser — NO truncation
                story.extend(_md_to_flowables(text))

                # If the message has paper data, render it as a table
                data = m.get("data")
                if data and isinstance(data, list) and mtype == "papers":
                    paper_rows = [["Title", "Authors", "Year", "Citations"]]
                    for p in data[:20]:
                        if isinstance(p, dict) and p.get("title"):
                            authors = ", ".join(p.get("authors", [])[:3])
                            if len(p.get("authors", [])) > 3:
                                authors += " et al."
                            paper_rows.append([
                                p.get("title", "")[:80],
                                authors[:40],
                                str(p.get("year", "")),
                                str(p.get("citations", "")),
                            ])
                    if len(paper_rows) > 1:
                        col_widths = [200, 120, 40, 50]
                        ptbl = Table(
                            [[Paragraph(_esc(c), styles["MDBody"]) for c in row] for row in paper_rows],
                            colWidths=col_widths,
                        )
                        ptbl.setStyle(TableStyle([
                            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#0891b2")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                            ("FONTSIZE", (0, 0), (-1, -1), 8),
                            ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#cccccc")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]))
                        story.append(Spacer(1, 3 * mm))
                        story.append(ptbl)

                story.append(Spacer(1, 3 * mm))
                story.append(HRFlowable(width="100%", thickness=0.3, color=HexColor("#dddddd"), spaceBefore=2, spaceAfter=6))

            doc.build(story)
            size = os.path.getsize(out_path)
            return {"ok": True, "path": out_path, "size": size, "format": "pdf"}
        except ImportError:
            # Fallback: save as markdown with .pdf note
            fmt = "markdown"
            ext = ".md"
            content = f"# {title}\n\n*PDF export requires reportlab. Exported as Markdown instead.*\n\n"
            for m in messages:
                role = m.get("role", "user").capitalize()
                content += f"## {role}\n\n{m.get('content', '')}\n\n---\n\n"
    else:
        return {"error": f"Unknown export format: {fmt}"}

    # Write file
    out_path = os.path.join(EXPORTS_DIR, f"{safe_title}_{timestamp}{ext}")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)

    size = os.path.getsize(out_path)
    return {"ok": True, "path": out_path, "size": size, "format": fmt}


def auto_rename_thread(args):
    """Use a cheap LLM call to generate a short, descriptive thread title.
    Args: {content, api_key, provider, model}
    Returns: {ok, title}"""
    content = args.get("content", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    model = args.get("model", "")

    if not api_key or not content:
        # Fallback: extract first meaningful words
        words = content.strip().split()[:8]
        return {"ok": True, "title": " ".join(words)[:50] if words else "Research"}

    # Use cheapest model for each provider
    cheap_models = {
        "google": "gemini-3.1-flash-lite-preview",
        "openai": "gpt-4.1-nano",
        "anthropic": "claude-haiku-4-5-20250514",
        "deepseek": "deepseek-chat",
        "groq": "llama-3.1-8b-instant",
    }
    rename_model = cheap_models.get(provider, model)

    result = _llm_chat(provider, api_key, rename_model,
                       [{"role": "system", "content": "Generate a short, descriptive title (3-8 words) for this research conversation. Return ONLY the title text, no quotes, no explanation."},
                        {"role": "user", "content": content[:500]}],
                       temperature=0.3, max_tokens=64)
    if "error" in result:
        words = content.strip().split()[:8]
        return {"ok": True, "title": " ".join(words)[:50]}

    title = result["text"].strip().strip('"').strip("'")[:60]
    return {"ok": True, "title": title, "tokens": result.get("usage")}


def export_research_snapshot(args):
    """Export a complete research snapshot: folder with manuscript, refs, papers list, logs, etc.
    Args: {manuscript, papers, bibliography, query, study_design, logs, thread_title,
           draft_sections, extracted_texts, messages, format}
    Returns: {ok, folder, files: [{name, path, size}]}"""
    manuscript = args.get("manuscript", "")
    papers = args.get("papers", [])
    bibliography = args.get("bibliography", "")
    query = args.get("query", "")
    study_design = args.get("study_design", "systematic_review")
    logs = args.get("logs", [])
    title = args.get("thread_title", "Research Export")
    draft_sections = args.get("draft_sections", [])
    messages = args.get("messages", [])
    fmt = args.get("format", "markdown")

    timestamp = time.strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:40]
    export_folder = os.path.join(EXPORTS_DIR, f"{safe_title}_{timestamp}")
    os.makedirs(export_folder, exist_ok=True)

    files = []

    # 1. Manuscript (markdown)
    if manuscript:
        md_path = os.path.join(export_folder, "manuscript.md")
        header = f"# {title}\n\n"
        header += f"**Research Question:** {query}\n\n"
        header += f"**Study Design:** {study_design.replace('_', ' ').title()}\n\n"
        header += f"**Generated:** {time.strftime('%Y-%m-%d %H:%M')}\n\n---\n\n"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(header + manuscript)
        files.append({"name": "manuscript.md", "path": md_path, "size": os.path.getsize(md_path)})

    # 2. Bibliography (BibTeX)
    if bibliography:
        bib_path = os.path.join(export_folder, "bibliography.bib")
        with open(bib_path, "w", encoding="utf-8") as f:
            f.write(f"% Bibliography for: {title}\n% Generated {time.strftime('%Y-%m-%d')}\n\n{bibliography}")
        files.append({"name": "bibliography.bib", "path": bib_path, "size": os.path.getsize(bib_path)})

    # 3. Papers list (JSON)
    if papers:
        papers_path = os.path.join(export_folder, "papers.json")
        with open(papers_path, "w", encoding="utf-8") as f:
            json.dump({"query": query, "total": len(papers), "papers": papers}, f, indent=2, ensure_ascii=False)
        files.append({"name": "papers.json", "path": papers_path, "size": os.path.getsize(papers_path)})

    # 4. Draft sections (individual markdown files)
    if draft_sections:
        sections_dir = os.path.join(export_folder, "sections")
        os.makedirs(sections_dir, exist_ok=True)
        for s in draft_sections:
            stype = s.get("type", "section")
            sec_path = os.path.join(sections_dir, f"{stype}.md")
            with open(sec_path, "w", encoding="utf-8") as f:
                f.write(f"## {stype.title()}\n\n{s.get('text', '')}")
            files.append({"name": f"sections/{stype}.md", "path": sec_path, "size": os.path.getsize(sec_path)})

    # 5. Pipeline log
    if logs:
        log_path = os.path.join(export_folder, "pipeline_log.txt")
        with open(log_path, "w", encoding="utf-8") as f:
            for entry in logs:
                f.write(f"[{entry.get('time', '')}] [{entry.get('phase', '')}] {entry.get('message', '')}\n")
        files.append({"name": "pipeline_log.txt", "path": log_path, "size": os.path.getsize(log_path)})

    # 6. Chat export (markdown)
    if messages:
        chat_path = os.path.join(export_folder, "chat_history.md")
        content = f"# Chat History — {title}\n\n"
        for m in messages:
            role = m.get("role", "user").capitalize()
            if role == "User":
                content += f"## You\n\n{m.get('content', '')}\n\n---\n\n"
            elif role == "Assistant":
                content += f"## Assistant\n\n{m.get('content', '')}\n\n---\n\n"
            elif role == "Tool":
                content += f"### Tool: {m.get('tool_used', 'result')}\n\n{m.get('content', '')}\n\n---\n\n"
        with open(chat_path, "w", encoding="utf-8") as f:
            f.write(content)
        files.append({"name": "chat_history.md", "path": chat_path, "size": os.path.getsize(chat_path)})

    # 7. Copy generated figures/tables into assets/ folder
    generated_figures = args.get("generated_figures", [])
    generated_tables = args.get("generated_tables", [])
    if generated_figures or generated_tables:
        import shutil as _shutil
        assets_dir = os.path.join(export_folder, "assets")
        os.makedirs(assets_dir, exist_ok=True)
        for fig in generated_figures:
            src = fig.get("path", "")
            if src and os.path.isfile(src):
                dest = os.path.join(assets_dir, os.path.basename(src))
                _shutil.copy2(src, dest)
                files.append({"name": f"assets/{os.path.basename(src)}", "path": dest, "size": os.path.getsize(dest)})
        for tbl in generated_tables:
            src = tbl.get("path", "")
            if src and os.path.isfile(src):
                dest = os.path.join(assets_dir, os.path.basename(src))
                _shutil.copy2(src, dest)
                files.append({"name": f"assets/{os.path.basename(src)}", "path": dest, "size": os.path.getsize(dest)})
        # Also copy any charts from the charts directory that are referenced in the manuscript
        chart_dir = os.path.join(RESEARCH_DIR, "charts")
        if os.path.isdir(chart_dir) and manuscript:
            for fname in os.listdir(chart_dir):
                if fname.endswith(".png") and fname in manuscript:
                    src = os.path.join(chart_dir, fname)
                    dest = os.path.join(assets_dir, fname)
                    if not os.path.exists(dest):
                        _shutil.copy2(src, dest)
                        files.append({"name": f"assets/{fname}", "path": dest, "size": os.path.getsize(dest)})

    # 7.5. Copy acquired PDF reference files into references/ folder
    acquired_pdfs = args.get("acquired_pdfs", [])
    if acquired_pdfs:
        import shutil as _shutil2
        refs_dir = os.path.join(export_folder, "references")
        os.makedirs(refs_dir, exist_ok=True)
        copied_count = 0
        for pdf_info in acquired_pdfs:
            src = pdf_info.get("path", "")
            if src and os.path.isfile(src):
                dest = os.path.join(refs_dir, os.path.basename(src))
                if not os.path.exists(dest):
                    _shutil2.copy2(src, dest)
                files.append({"name": f"references/{os.path.basename(src)}", "path": dest, "size": os.path.getsize(dest)})
                copied_count += 1
        # Also scan PAPERS_DIR for any PDFs not in acquired_pdfs list (belt and suspenders)
        if os.path.isdir(PAPERS_DIR):
            for fname in os.listdir(PAPERS_DIR):
                if fname.endswith(".pdf"):
                    src = os.path.join(PAPERS_DIR, fname)
                    dest = os.path.join(refs_dir, fname)
                    if not os.path.exists(dest) and os.path.getsize(src) > 1024:
                        _shutil2.copy2(src, dest)
                        files.append({"name": f"references/{fname}", "path": dest, "size": os.path.getsize(dest)})
                        copied_count += 1

    # 8. Telemetry / token usage summary
    telemetry = {
        "query": query,
        "study_design": study_design,
        "timestamp": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "papers_found": len(papers),
        "sections_drafted": len(draft_sections),
        "manuscript_length": len(manuscript),
        "figures_generated": len(generated_figures),
        "tables_generated": len(generated_tables),
        "reference_pdfs": len(acquired_pdfs),
        "log_entries": len(logs),
    }
    tel_path = os.path.join(export_folder, "telemetry.json")
    with open(tel_path, "w", encoding="utf-8") as f:
        json.dump(telemetry, f, indent=2)
    files.append({"name": "telemetry.json", "path": tel_path, "size": os.path.getsize(tel_path)})

    # 8.5. Export prompt logs — one txt file per agent with full input/output
    prompt_logs = args.get("prompt_logs", {}) or _get_prompt_logs()
    if prompt_logs:
        prompts_dir = os.path.join(export_folder, "prompts")
        os.makedirs(prompts_dir, exist_ok=True)
        for agent_name, entries in prompt_logs.items():
            safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', agent_name)
            prompt_path = os.path.join(prompts_dir, f"{safe_name}.txt")
            with open(prompt_path, "w", encoding="utf-8") as f:
                f.write(f"{'='*80}\n")
                f.write(f"AGENT: {agent_name}\n")
                f.write(f"Total calls: {len(entries)}\n")
                f.write(f"{'='*80}\n\n")
                for ci, entry in enumerate(entries, 1):
                    f.write(f"{'─'*60}\n")
                    f.write(f"CALL {ci} — {entry.get('timestamp', '')}\n")
                    f.write(f"{'─'*60}\n\n")
                    # Variables
                    v = entry.get("variables", {})
                    f.write(f"── Variables ──\n")
                    for k, val in v.items():
                        f.write(f"  {k}: {val}\n")
                    f.write(f"\n── System Prompt ──\n{entry.get('system_prompt', '(none)')}\n\n")
                    f.write(f"── User Prompt (INPUT) ──\n{entry.get('user_prompt', '(none)')}\n\n")
                    f.write(f"── Model Output (OUTPUT) ──\n{entry.get('output', '(none)')}\n\n")
                    t = entry.get("tokens", {})
                    if t:
                        f.write(f"── Tokens ──\n  Input: {t.get('input_tokens', 0)}  Output: {t.get('output_tokens', 0)}\n\n")
            files.append({"name": f"prompts/{safe_name}.txt", "path": prompt_path, "size": os.path.getsize(prompt_path)})

    # 9. Generate PDF with embedded images
    if manuscript:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Preformatted,
                Table as RLTable, TableStyle as RLTableStyle, HRFlowable,
                Image as RLImage,
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm, mm
            from reportlab.lib.colors import HexColor
            from reportlab.lib.enums import TA_LEFT

            pdf_dest = os.path.join(export_folder, "manuscript.pdf")
            pdf_doc = SimpleDocTemplate(
                pdf_dest, pagesize=A4,
                leftMargin=2 * cm, rightMargin=2 * cm,
                topMargin=2 * cm, bottomMargin=2 * cm,
            )
            _styles = getSampleStyleSheet()
            for lvl, (sz, sp) in enumerate([(18, 14), (15, 12), (13, 10), (11, 8)], 1):
                _styles.add(ParagraphStyle(f"_H{lvl}", parent=_styles[f"Heading{lvl}"],
                    fontSize=sz, spaceBefore=sp, spaceAfter=sz // 3, textColor=HexColor("#1a1a2e")))
            _styles.add(ParagraphStyle("_Body", parent=_styles["Normal"], fontSize=10,
                leading=14, spaceAfter=4, alignment=TA_LEFT))
            _styles.add(ParagraphStyle("_Code", fontName="Courier", fontSize=8, leading=10,
                backColor=HexColor("#f5f5f5"), textColor=HexColor("#333333"),
                leftIndent=12, rightIndent=12, spaceBefore=6, spaceAfter=6))
            _styles.add(ParagraphStyle("_Caption", parent=_styles["Normal"], fontSize=9,
                leading=12, alignment=1, textColor=HexColor("#555555"), fontName="Helvetica-Oblique"))
            _styles.add(ParagraphStyle("_Bullet", parent=_styles["Normal"], fontSize=10,
                leading=13, leftIndent=20, bulletIndent=10, spaceBefore=1, spaceAfter=1))

            def _pe(t):
                return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            def _pi(t):
                t = _pe(t)
                t = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', t)
                t = re.sub(r'\*(.+?)\*', r'<i>\1</i>', t)
                t = re.sub(r'`(.+?)`', r'<font face="Courier" size="8" color="#c0392b">\1</font>', t)
                t = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2" color="blue">\1</a>', t)
                return t

            # Image search dirs: export assets, charts dir, original gen paths
            _img_search = [
                os.path.join(export_folder, "assets"),
                os.path.join(RESEARCH_DIR, "charts"),
            ]
            for _fig in generated_figures:
                _d = os.path.dirname(_fig.get("path", ""))
                if _d and _d not in _img_search:
                    _img_search.append(_d)

            story = []
            story.append(Paragraph(_pe(title), _styles["Title"]))
            story.append(Paragraph(f'<font size="9" color="#888">{_pe(query)}</font>', _styles["Normal"]))
            story.append(Paragraph(f'<font size="8" color="#aaa">{study_design.replace("_", " ").title()} — {time.strftime("%Y-%m-%d")}</font>', _styles["Normal"]))
            story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#0891b2"), spaceBefore=6, spaceAfter=10))

            _mlines = manuscript.split("\n")
            _mi = 0
            while _mi < len(_mlines):
                _ln = _mlines[_mi]

                # Image
                _im = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', _ln.strip())
                if _im:
                    _alt, _ref = _im.group(1), _im.group(2)
                    _ifname = os.path.basename(_ref)
                    _ipath = None
                    for _sd in _img_search:
                        _c = os.path.join(_sd, _ifname)
                        if os.path.isfile(_c):
                            _ipath = _c
                            break
                    if _ipath:
                        try:
                            _iw = min(A4[0] - 4 * cm, 14 * cm)
                            story.append(Spacer(1, 3 * mm))
                            story.append(RLImage(_ipath, width=_iw, kind="proportional"))
                            if _alt:
                                story.append(Paragraph(f'<i>{_pe(_alt)}</i>', _styles["_Caption"]))
                            story.append(Spacer(1, 3 * mm))
                        except Exception:
                            story.append(Paragraph(f'[Image: {_pe(_alt)}]', _styles["_Body"]))
                    _mi += 1
                    continue

                # Headers
                _hm = re.match(r'^(#{1,4})\s+(.+)$', _ln)
                if _hm:
                    _hl = min(len(_hm.group(1)), 4)
                    story.append(Paragraph(_pi(_hm.group(2)), _styles[f"_H{_hl}"]))
                    _mi += 1
                    continue

                # HR
                if re.match(r'^-{3,}$', _ln.strip()):
                    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#ccc"), spaceBefore=4, spaceAfter=4))
                    _mi += 1
                    continue

                # Code block
                if _ln.strip().startswith("```"):
                    _cb = []
                    _mi += 1
                    while _mi < len(_mlines) and not _mlines[_mi].strip().startswith("```"):
                        _cb.append(_mlines[_mi])
                        _mi += 1
                    _mi += 1
                    story.append(Preformatted(_pe("\n".join(_cb)), _styles["_Code"]))
                    continue

                # Table
                if _ln.strip().startswith("|") and "|" in _ln[1:]:
                    _trows = []
                    while _mi < len(_mlines) and _mlines[_mi].strip().startswith("|"):
                        _r = _mlines[_mi].strip()
                        if not all(c in "-| :" for c in _r):
                            _trows.append([c.strip() for c in _r.split("|")[1:-1]])
                        _mi += 1
                    if _trows:
                        _nc = max(len(r) for r in _trows)
                        for r in _trows:
                            while len(r) < _nc:
                                r.append("")
                        _cw = (A4[0] - 4 * cm) / _nc
                        _td = [[Paragraph(_pi(c), _styles["_Body"]) for c in r] for r in _trows]
                        _t = RLTable(_td, colWidths=[_cw] * _nc)
                        _t.setStyle(RLTableStyle([
                            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#e8e8e8")),
                            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#ccc")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("TOPPADDING", (0, 0), (-1, -1), 4),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ]))
                        story.append(Spacer(1, 3 * mm))
                        story.append(_t)
                        story.append(Spacer(1, 3 * mm))
                    continue

                # Bullet
                _bm = re.match(r'^(\s*)[-*]\s+(.+)$', _ln)
                if _bm:
                    story.append(Paragraph("\u2022 " + _pi(_bm.group(2)), _styles["_Bullet"]))
                    _mi += 1
                    continue

                # Bold caption lines
                if _ln.strip().startswith("**Figure ") or _ln.strip().startswith("**Table "):
                    story.append(Paragraph(f'<b>{_pi(_ln.strip().replace("**", ""))}</b>', _styles["_Caption"]))
                    _mi += 1
                    continue

                # Empty
                if not _ln.strip():
                    story.append(Spacer(1, 2 * mm))
                    _mi += 1
                    continue

                # Paragraph
                _plines = []
                while _mi < len(_mlines) and _mlines[_mi].strip() and not _mlines[_mi].strip().startswith(("#", "```", "|", "---", "- ", "* ", "![")):
                    if re.match(r'^\s*\d+[.)]\s+', _mlines[_mi]):
                        break
                    if _mlines[_mi].strip().startswith("**Figure ") or _mlines[_mi].strip().startswith("**Table "):
                        break
                    _plines.append(_mlines[_mi])
                    _mi += 1
                if _plines:
                    story.append(Paragraph(_pi(" ".join(_plines)), _styles["_Body"]))
                else:
                    story.append(Paragraph(_pi(_ln), _styles["_Body"]))
                    _mi += 1

            pdf_doc.build(story)
            files.append({"name": "manuscript.pdf", "path": pdf_dest, "size": os.path.getsize(pdf_dest)})
        except ImportError:
            pass
        except Exception:
            pass

    # 10. Generate DOCX (editable manuscript) with embedded figures and tables
    if manuscript:
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = DocxDocument()
            # Title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Research Question: {query}")
            doc.add_paragraph(f"Study Design: {study_design.replace('_', ' ').title()}")
            doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("")

            # Build a map of figure filenames for embedding — search all possible dirs
            fig_paths = {}
            _docx_img_dirs = [
                os.path.join(export_folder, "assets"),
                os.path.join(RESEARCH_DIR, "charts"),
            ]
            for _gf in generated_figures:
                _gd = os.path.dirname(_gf.get("path", ""))
                if _gd and _gd not in _docx_img_dirs:
                    _docx_img_dirs.append(_gd)
            for _did in _docx_img_dirs:
                if os.path.isdir(_did):
                    for fname in os.listdir(_did):
                        if fname.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp")):
                            fig_paths[fname] = os.path.join(_did, fname)

            # Parse markdown sections
            lines = manuscript.split("\n")
            i_line = 0
            while i_line < len(lines):
                line = lines[i_line]
                if line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.strip().startswith("!["):
                    # Markdown image: ![alt](path)
                    img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
                    if img_match:
                        alt_text = img_match.group(1)
                        img_ref = img_match.group(2)
                        img_fname = os.path.basename(img_ref)
                        img_path = fig_paths.get(img_fname)
                        if not img_path:
                            # Try finding in charts dir
                            chart_path = os.path.join(RESEARCH_DIR, "charts", img_fname)
                            if os.path.isfile(chart_path):
                                img_path = chart_path
                        if img_path and os.path.isfile(img_path):
                            doc.add_picture(img_path, width=Inches(5.5))
                            caption_para = doc.add_paragraph(alt_text)
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.runs[0].italic = True if caption_para.runs else None
                        else:
                            doc.add_paragraph(f"[{alt_text}]", style="Normal")
                elif line.strip().startswith("**Figure ") or line.strip().startswith("**Table "):
                    # Bold caption line
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip().replace("**", ""))
                    run.bold = True
                    run.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.strip().startswith("|") and "|" in line[1:]:
                    # Markdown table — collect all rows and render as DOCX table
                    table_lines = []
                    while i_line < len(lines) and lines[i_line].strip().startswith("|"):
                        row_text = lines[i_line].strip()
                        # Skip separator rows (|---|---|)
                        if not re.match(r'^\|[\s:-]+\|$', row_text.replace("|", "").replace("-", "").replace(":", "").strip() and row_text or ""):
                            if not all(c in "-| :" for c in row_text):
                                cells = [c.strip() for c in row_text.strip("|").split("|")]
                                table_lines.append(cells)
                        i_line += 1
                    i_line -= 1  # will be incremented at loop end
                    if table_lines:
                        ncols = max(len(row) for row in table_lines)
                        tbl = doc.add_table(rows=len(table_lines), cols=ncols)
                        tbl.style = "Light Grid Accent 1"
                        for ri, row_cells in enumerate(table_lines):
                            for ci, cell_val in enumerate(row_cells):
                                if ci < ncols:
                                    tbl.rows[ri].cells[ci].text = str(cell_val)
                        doc.add_paragraph("")  # spacing after table
                elif line.strip():
                    doc.add_paragraph(line.strip())
                i_line += 1

            docx_path = os.path.join(export_folder, "manuscript.docx")
            doc.save(docx_path)
            files.append({"name": "manuscript.docx", "path": docx_path, "size": os.path.getsize(docx_path)})
        except ImportError:
            pass  # python-docx not installed
        except Exception:
            pass

    return {"ok": True, "folder": export_folder, "files": files, "file_count": len(files)}


# ══════════════════════════════════════════════════════════════════════════════
#  Phase 3.2b — Data Analyst Agent: Generate figures/tables from blueprint plan
# ══════════════════════════════════════════════════════════════════════════════

def generate_pipeline_figures(args):
    """Phase 3.2b — Data Analyst Agent: Generate charts/tables for each figure_plan item.
    Uses LLM to analyze papers and produce chart data, then matplotlib to render PNGs.
    Args: {figure_plan: str[], table_plan: str[], papers_context, draft_sections,
           query, api_key, provider, model, step_config}
    Returns: {ok, figures: [{description, path, chart_type, size}],
                  tables:  [{description, path, markdown, size}],
                  errors: str[]}"""
    figure_plan = args.get("figure_plan", [])
    table_plan = args.get("table_plan", [])
    papers_context = args.get("papers_context", "")
    draft_sections = args.get("draft_sections", [])
    query = args.get("query", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "fast", "max_tokens": 4096, "temperature": 0.3})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required."}

    # Collect figure placeholders from draft sections (multiple formats LLMs use)
    draft_figure_reqs = []
    for s in draft_sections:
        text = s.get("text", "")
        # Match all placeholder formats the Lead Author might use
        for pattern in [
            r'\[FIGURE:\s*(.+?)\]',
            r'\((?:Suggest|Insert|Place)\s+(?:placing\s+)?Figure\s*\d*\s*(?:here)?[:\s]*(.+?)\)',
            r'\[(?:Insert|Place)\s+Figure\s*\d*\s*(?:here)?[:\s]*(.+?)\]',
            r'\(Placement Suggestion:\s*(.+?)\)',
        ]:
            for m in re.findall(pattern, text, re.IGNORECASE):
                desc = m.strip().rstrip('.')
                if desc and desc not in figure_plan and desc not in draft_figure_reqs and len(desc) > 5:
                    draft_figure_reqs.append(desc)

    all_figures = list(figure_plan) + draft_figure_reqs
    if not all_figures and not table_plan:
        return {"ok": True, "figures": [], "tables": [], "errors": []}

    # Combine draft text for context
    draft_text = "\n\n".join(f"## {s.get('type', '')}\n{s.get('text', '')[:2000]}" for s in draft_sections)

    generated_figures = []
    generated_tables = []
    errors = []

    # ── Classify each figure: "chart" (data viz) vs "illustration" (scientific diagram) ──
    ILLUSTRATION_KW = frozenset([
        "mechanism", "pathway", "diagram", "architecture", "framework",
        "workflow", "process", "overview", "schematic", "cross-section", "structure",
        "interaction", "cycle", "cascade", "signaling", "flowchart", "illustration",
        "anatomy", "morphology", "circuit", "network diagram", "conceptual model",
    ])
    chart_figures = []
    illustration_figures = []
    for desc in all_figures:
        dl = desc.lower()
        if any(kw in dl for kw in ILLUSTRATION_KW):
            illustration_figures.append(desc)
        else:
            chart_figures.append(desc)

    # ─��� Generate charts/graphs via Gemini code execution (matplotlib/seaborn) ──
    # Uses code_execution tool so Gemini writes & runs Python, then we extract the PNG.
    charts_dir = os.path.join(RESEARCH_DIR, "charts")
    os.makedirs(charts_dir, exist_ok=True)

    for i, fig_desc in enumerate(chart_figures):
        try:
            code_exec_prompt = (
                f"Research question: {query}\n\n"
                f"Draft manuscript excerpt:\n{draft_text[:4000]}\n\n"
                f"Available literature:\n{papers_context[:3000]}\n\n"
                f"Create a publication-quality chart for: \"{fig_desc}\"\n\n"
                "INSTRUCTIONS:\n"
                "1. Write Python code using matplotlib and/or seaborn to create the chart.\n"
                "2. Use realistic data extracted from the literature context above.\n"
                "3. Style: white background, professional academic style, clear axis labels, legend if needed.\n"
                "4. IMPORTANT: Use text wrapping for long labels — never truncate with '...' or ellipsis.\n"
                "   For long x-axis labels, rotate them 45° or use textwrap.fill(label, 15).\n"
                "   For long titles, use textwrap.fill(title, 50).\n"
                "5. Use plt.tight_layout() before saving.\n"
                "6. Save the figure to '/tmp/chart_output.png' with dpi=200, facecolor='white'.\n"
                "7. After saving, print the one-sentence figure caption for the manuscript.\n"
                "8. Print 'CAPTION:' followed by the caption text on a single line.\n\n"
                "Write and execute the code now."
            )
            result = _llm_chat(
                provider, api_key, model,
                [{"role": "system", "content": "You are a data visualization expert for academic research. Write and execute Python code to create charts using matplotlib/seaborn. Always use white backgrounds, clear labels, and wrap long text instead of truncating."},
                 {"role": "user", "content": code_exec_prompt}],
                temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                code_execution=True,
            )
            if "error" in result:
                errors.append(f"Figure '{fig_desc}': code execution error — {result['error']}")
                continue

            # Extract the executed code output and any generated image
            text = result.get("text", "")
            structured = result.get("structured") or {}
            code_results = structured.get("code_execution_results", [])

            # Check if code execution produced an image (base64 in output)
            chart_saved = False
            caption = fig_desc

            # Extract caption from text
            cap_match = re.search(r'CAPTION:\s*(.+)', text)
            if cap_match:
                caption = cap_match.group(1).strip()

            # Try to find generated PNG from code execution output
            # Gemini code execution runs in sandbox — check if it produced output with image data
            exec_code = ""
            exec_output = ""
            for cr in code_results:
                if cr.get("code"):
                    exec_code = cr["code"]
                if cr.get("output"):
                    exec_output += cr["output"]

            # If code execution ran successfully, try to extract/re-execute the chart code locally
            if exec_code and "plt." in exec_code:
                chart_path = os.path.join(charts_dir, f"figure_{i + 1}_{int(time.time())}.png")
                try:
                    # Re-execute the code locally to get the actual PNG file
                    local_code = exec_code
                    # Replace the save path to our local path
                    local_code = re.sub(
                        r"""(plt\.savefig|fig\.savefig)\s*\(\s*['"][^'"]+['"]""",
                        f'plt.savefig("{chart_path.replace(chr(92), "/")}"',
                        local_code
                    )
                    # Ensure it saves to our path if no savefig found
                    if "savefig" not in local_code:
                        local_code += f'\nplt.savefig("{chart_path.replace(chr(92), "/")}", dpi=200, bbox_inches="tight", facecolor="white")\nplt.close()'
                    else:
                        local_code += "\nplt.close()"
                    # Execute in isolated namespace — redirect stdout to prevent
                    # any print() in LLM-generated code from corrupting the JSON sidecar output
                    import io as _io, sys as _sys
                    _old_stdout = _sys.stdout
                    _sys.stdout = _io.StringIO()
                    try:
                        _ns = {"__builtins__": __builtins__}
                        exec(local_code, _ns)  # noqa: S102
                    finally:
                        _sys.stdout = _old_stdout

                    if os.path.isfile(chart_path) and os.path.getsize(chart_path) > 500:
                        chart_saved = True
                        generated_figures.append({
                            "description": fig_desc,
                            "caption": caption,
                            "path": chart_path,
                            "chart_type": "code_execution",
                            "size": os.path.getsize(chart_path),
                            "index": i + 1,
                        })
                except Exception as ce:
                    errors.append(f"Figure '{fig_desc}': local exec error — {ce}")

            # Fallback: if code execution produced no file, ask the LLM for JSON chart data
            # (works for all providers, not just Gemini code-execution)
            if not chart_saved:
                try:
                    json_prompt = (
                        f"Research question: {query}\n\n"
                        f"Create a publication-quality chart for: \"{fig_desc}\"\n\n"
                        f"Context (literature):\n{papers_context[:2000]}\n\n"
                        "Return ONLY a JSON object (no markdown, no explanation) with these fields:\n"
                        '{"chart_type":"bar","title":"...","labels":["Label1","Label2","Label3"],'
                        '"data":[25,40,35],"xlabel":"X Axis","ylabel":"Y Axis","caption":"One sentence describing the figure."}'
                        "\n\nchart_type must be one of: bar, line, pie, scatter\n"
                        "Use realistic numerical data extracted from the literature above."
                    )
                    jr = _llm_chat(
                        provider, api_key, model,
                        [{"role": "system", "content": "Return ONLY valid JSON with no markdown fences, comments, or extra text."},
                         {"role": "user", "content": json_prompt}],
                        temperature=0.3, max_tokens=512,
                    )
                    if "error" not in jr:
                        jtext = jr.get("text", "").strip()
                        # Strip markdown fences if present
                        jtext = re.sub(r'^```(?:json)?\s*', '', jtext, flags=re.MULTILINE)
                        jtext = re.sub(r'\s*```\s*$', '', jtext, flags=re.MULTILINE)
                        jmatch = re.search(r'\{.*\}', jtext, re.DOTALL)
                        if jmatch:
                            chart_data = json.loads(jmatch.group())
                            chart_result = generate_chart({
                                "chart_type": chart_data.get("chart_type", "bar"),
                                "data": chart_data.get("data", []),
                                "title": chart_data.get("title", fig_desc),
                                "xlabel": chart_data.get("xlabel", ""),
                                "ylabel": chart_data.get("ylabel", ""),
                                "labels": chart_data.get("labels", []),
                            })
                            if chart_result.get("ok"):
                                generated_figures.append({
                                    "description": fig_desc,
                                    "caption": chart_data.get("caption", fig_desc),
                                    "path": chart_result["path"],
                                    "chart_type": chart_result.get("chart_type", "bar"),
                                    "size": chart_result.get("size", 0),
                                    "index": i + 1,
                                })
                                chart_saved = True
                except Exception:
                    pass

                if not chart_saved:
                    errors.append(f"Figure '{fig_desc}': could not generate chart")

        except Exception as e:
            errors.append(f"Figure '{fig_desc}': {e}")

    # ── Generate tables via Gemini code execution (matplotlib/seaborn) ──
    for i, tbl_desc in enumerate(table_plan):
        try:
            tbl_code_prompt = (
                f"Research question: {query}\n\n"
                f"Draft manuscript excerpt:\n{draft_text[:4000]}\n\n"
                f"Available literature:\n{papers_context[:3000]}\n\n"
                f"Create a publication-quality table for: \"{tbl_desc}\"\n\n"
                "INSTRUCTIONS:\n"
                "1. Write Python code using matplotlib to render a table as an image.\n"
                "2. Use realistic data from the literature context above.\n"
                "3. Style: white background, light blue header row (#e0f2fe), clear grid lines.\n"
                "4. IMPORTANT: Use text wrapping for ALL cell content — NEVER truncate with '...' or ellipsis.\n"
                "   Use textwrap.fill(text, 20) for cells and textwrap.fill(text, 15) for headers.\n"
                "5. Use auto_set_column_width and scale(1, 1.6) for readability.\n"
                "6. Save to '/tmp/table_output.png' with dpi=200, facecolor='white'.\n"
                "7. Also print the table in markdown format (| col1 | col2 | ... |) for text embedding.\n"
                "8. Print 'CAPTION:' followed by a one-sentence caption.\n\n"
                "Write and execute the code now."
            )
            result = _llm_chat(
                provider, api_key, model,
                [{"role": "system", "content": "You are a data visualization expert creating academic tables. Write and execute Python code using matplotlib. Wrap all text — never truncate."},
                 {"role": "user", "content": tbl_code_prompt}],
                temperature=sc["temperature"], max_tokens=sc["max_tokens"],
                code_execution=True,
            )
            if "error" in result:
                errors.append(f"Table '{tbl_desc}': code execution error — {result['error']}")
                continue

            text = result.get("text", "")
            structured = result.get("structured") or {}
            code_results = structured.get("code_execution_results", [])

            caption = tbl_desc
            cap_match = re.search(r'CAPTION:\s*(.+)', text)
            if cap_match:
                caption = cap_match.group(1).strip()

            # Extract markdown table from output
            md_table = ""
            md_match = re.search(r'(\|.+\|(?:\n\|.+\|)+)', text)
            if md_match:
                md_table = md_match.group(1)

            # Re-execute code locally
            exec_code = ""
            for cr in code_results:
                if cr.get("code"):
                    exec_code = cr["code"]

            tbl_saved = False
            if exec_code and ("plt." in exec_code or "matplotlib" in exec_code):
                tbl_path = os.path.join(charts_dir, f"table_{i + 1}_{int(time.time())}.png")
                try:
                    local_code = exec_code
                    local_code = re.sub(
                        r"""(plt\.savefig|fig\.savefig)\s*\(\s*['"][^'"]+['"]""",
                        f'plt.savefig("{tbl_path.replace(chr(92), "/")}"',
                        local_code
                    )
                    if "savefig" not in local_code:
                        local_code += f'\nplt.savefig("{tbl_path.replace(chr(92), "/")}", dpi=200, bbox_inches="tight", facecolor="white")\nplt.close()'
                    else:
                        local_code += "\nplt.close()"
                    # Capture stdout so any print() in LLM-generated code doesn't
                    # corrupt the JSON sidecar output channel
                    import io as _io, sys as _sys
                    _old_stdout = _sys.stdout
                    _sys.stdout = _io.StringIO()
                    try:
                        _ns = {"__builtins__": __builtins__}
                        exec(local_code, _ns)  # noqa: S102
                    finally:
                        _sys.stdout = _old_stdout

                    if os.path.isfile(tbl_path) and os.path.getsize(tbl_path) > 500:
                        tbl_saved = True
                        generated_tables.append({
                            "description": tbl_desc,
                            "caption": caption,
                            "markdown": md_table,
                            "path": tbl_path,
                            "size": os.path.getsize(tbl_path),
                            "index": i + 1,
                        })
                except Exception as ce:
                    errors.append(f"Table '{tbl_desc}': local exec error — {ce}")

            # Fallback: ask LLM for structured JSON table data, then call generate_table()
            if not tbl_saved:
                try:
                    json_prompt = (
                        f'Create a publication-quality data table for: "{tbl_desc}"\n'
                        f'Context: {query[:300]}\n\n'
                        'Return ONLY a JSON object with this exact schema '
                        '(no markdown fences, no explanation):\n'
                        '{"title":"...", "caption":"...", "headers":["Col1","Col2",...], '
                        '"rows":[["val","val",...],...]}\n\n'
                        'Requirements: meaningful headers, realistic data values, '
                        'at least 3 columns and 4-8 rows, data relevant to the description.'
                    )
                    jr = _llm_chat(
                        provider, api_key, model,
                        [{"role": "user", "content": json_prompt}],
                        temperature=0.3, max_tokens=600,
                    )
                    jr_text = (jr.get("content") or "").strip()
                    # Strip markdown fences if present
                    jr_text = re.sub(r'^```[a-z]*\n?', '', jr_text)
                    jr_text = re.sub(r'\n?```$', '', jr_text).strip()
                    tbl_data = json.loads(jr_text)
                    tbl_result = generate_table({
                        "headers": tbl_data.get("headers", []),
                        "rows": tbl_data.get("rows", []),
                        "title": tbl_data.get("title", tbl_desc),
                        "format": "image",
                    })
                    if tbl_result.get("ok"):
                        generated_tables.append({
                            "description": tbl_desc,
                            "caption": tbl_data.get("caption", tbl_desc),
                            "markdown": tbl_result.get("markdown", ""),
                            "path": tbl_result.get("path", ""),
                            "size": tbl_result.get("size", 0),
                            "index": i + 1,
                        })
                        tbl_saved = True
                except (json.JSONDecodeError, Exception) as fe:
                    errors.append(f"Table '{tbl_desc}': JSON fallback failed — {fe}")
            if not tbl_saved:
                errors.append(f"Table '{tbl_desc}': no table generated")

        except Exception as e:
            errors.append(f"Table '{tbl_desc}': {e}")

    # ── Generate scientific illustrations ONLY via Nano Banana Pro ──
    if illustration_figures and api_key:
        for desc in illustration_figures:
            try:
                ill_prompt = (
                    f"Create a clean, publication-quality scientific illustration for an academic paper. "
                    f"Subject: {desc}. Context: {query[:200]}. "
                    f"Style: professional scientific diagram with clear labels, white background, "
                    f"suitable for a peer-reviewed journal. No text watermarks. "
                    f"Include accurate anatomical/molecular/process details. Use proper scientific nomenclature."
                )
                ill_result = _generate_illustration(ill_prompt, api_key, provider)
                if ill_result.get("ok"):
                    generated_figures.append({
                        "description": desc,
                        "caption": f"Illustration: {desc}",
                        "path": ill_result["path"],
                        "chart_type": "illustration",
                        "size": ill_result.get("size", 0),
                        "index": len(generated_figures) + 1,
                    })
                else:
                    errors.append(f"Illustration '{desc}': {ill_result.get('error', 'failed')}")
            except Exception as e:
                errors.append(f"Illustration '{desc}': {e}")

    return {"ok": True, "figures": generated_figures, "tables": generated_tables, "errors": errors}


def _generate_illustration(prompt, api_key, provider="google"):
    """Generate a scientific illustration using Gemini Nano Banana image generation.
    Uses the same API as ZenithEditor's generate_image action."""
    import base64 as _b64
    model = "gemini-3-pro-image-preview"
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"

    body = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "responseModalities": ["IMAGE", "TEXT"],
            "imageConfig": {"aspectRatio": "16:9", "imageSize": "1K"},
        },
    }

    payload = json.dumps(body).encode()
    req = urllib.request.Request(url, data=payload, headers={
        "Content-Type": "application/json", "User-Agent": _USER_AGENT})

    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        body_text = e.read().decode() if e.fp else ""
        return {"ok": False, "error": f"Gemini image API error {e.code}: {body_text[:300]}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

    candidates = data.get("candidates", [])
    if not candidates:
        feedback = data.get("promptFeedback", {})
        block = feedback.get("blockReason", "")
        return {"ok": False, "error": f"Blocked: {block}" if block else "No image returned"}

    result_b64 = None
    for part in candidates[0].get("content", {}).get("parts", []):
        if "inlineData" in part:
            result_b64 = part["inlineData"]["data"]
            break

    if not result_b64:
        return {"ok": False, "error": "No image in Gemini response"}

    img_bytes = _b64.b64decode(result_b64)
    chart_dir = os.path.join(RESEARCH_DIR, "charts")
    os.makedirs(chart_dir, exist_ok=True)
    safe = re.sub(r'[^a-zA-Z0-9_-]', '_', prompt[:30])
    path = os.path.join(chart_dir, f"illustration_{safe}_{int(time.time())}.png")
    with open(path, "wb") as f:
        f.write(img_bytes)
    return {"ok": True, "path": path, "size": len(img_bytes)}


# ══════════════════════════════════════════════════════════════════════════════
#  Scientific Illustrator Agent — dedicated phase for Nano Banana Pro images
# ══════════════════════════════════════════════════════════════════════════════

def scientific_illustrator_agent(args):
    """Dedicated Scientific Illustrator Agent.
    Analyses the draft manuscript + blueprint figure plan and generates scientific
    illustrations using Nano Banana Pro (gemini-3-pro-image-preview).

    This agent:
    1. Uses the LLM to decide which figures need scientific illustrations (vs charts).
    2. Writes a detailed visual brief for each illustration.
    3. Generates them via Nano Banana Pro image generation.
    4. Returns the illustrations alongside any previously generated charts.

    Args: {draft_sections, figure_plan, generated_figures (existing charts), query,
           api_key, provider, model, step_config}
    Returns: {ok, illustrations: [{description, caption, path, chart_type, size, index}],
              errors: str[]}"""
    draft_sections = args.get("draft_sections", [])
    figure_plan = args.get("figure_plan", [])
    existing_figures = args.get("generated_figures", [])
    query = args.get("query", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "google")
    sc = _get_step_config(args, {"model_tier": "fast", "max_tokens": 4096, "temperature": 0.3})
    model = _resolve_model(provider, sc["model_tier"], args.get("model", ""))

    if not api_key:
        return {"error": "API key required for Scientific Illustrator."}

    illustrations = []
    errors = []

    # Combine draft text for context
    draft_text = "\n\n".join(
        f"## {s.get('type', '')}\n{s.get('text', '')[:2000]}" for s in draft_sections
    )

    # ── Step 1: Collect all figure descriptions that need illustration ──
    # Sources: blueprint figure_plan, placeholders in draft, and any existing figures
    # that were classified as "illustration" but failed generation.
    illustration_requests = []

    # From figure_plan (blueprint)
    for desc in figure_plan:
        illustration_requests.append(desc)

    # From draft sections — scan for figure placeholders not yet generated
    existing_descs_lower = {f.get("description", "").lower().strip() for f in existing_figures}
    for s in draft_sections:
        text = s.get("text", "")
        for pattern in [
            r'\[FIGURE:\s*(.+?)\]',
            r'\((?:Suggest|Insert|Place)\s+(?:placing\s+)?Figure\s*\d*\s*(?:here)?[:\s]*(.+?)\)',
            r'\[(?:Insert|Place)\s+Figure\s*\d*\s*(?:here)?[:\s]*(.+?)\]',
            r'\(Placement Suggestion:\s*(.+?)\)',
        ]:
            for m in re.findall(pattern, text, re.IGNORECASE):
                desc = m.strip().rstrip('.')
                if desc and desc.lower().strip() not in existing_descs_lower and len(desc) > 5:
                    illustration_requests.append(desc)

    # Deduplicate
    seen = set()
    unique_requests = []
    for desc in illustration_requests:
        key = desc.lower().strip()[:80]
        if key not in seen:
            seen.add(key)
            unique_requests.append(desc)

    if not unique_requests:
        return {"ok": True, "illustrations": [], "errors": ["No illustration requests found"]}

    # ── Step 2: Ask the LLM which figures are illustrations vs charts ──
    # Filter out anything already generated as a chart
    existing_chart_descs = {
        f.get("description", "").lower().strip()
        for f in existing_figures
        if f.get("chart_type") != "illustration"
    }

    # Use LLM to write detailed visual briefs for each illustration
    brief_prompt = (
        f"Research topic: {query}\n\n"
        f"Manuscript excerpt:\n{draft_text[:3000]}\n\n"
        f"The following figures were requested in the manuscript blueprint. "
        f"For each one, decide if it should be a SCIENTIFIC ILLUSTRATION "
        f"(diagram, mechanism, pathway, process flow, anatomy, schematic, conceptual model) "
        f"or a DATA CHART (bar chart, line graph, scatter plot, pie chart, etc.).\n\n"
        f"For each figure that IS a scientific illustration, write a detailed visual brief "
        f"describing exactly what should be drawn. Be specific about layout, elements, labels, "
        f"arrows, colors, and scientific accuracy.\n\n"
        f"Figures:\n"
    )
    for i, desc in enumerate(unique_requests, 1):
        is_existing_chart = desc.lower().strip() in existing_chart_descs
        brief_prompt += f"{i}. {desc}{' [ALREADY GENERATED AS CHART]' if is_existing_chart else ''}\n"

    brief_prompt += (
        "\n\nRespond in JSON format:\n"
        '{"illustrations": [{"index": 1, "description": "original desc", '
        '"is_illustration": true, "visual_brief": "detailed visual description..."}]}'
    )

    result = _llm_chat(
        provider, api_key, model,
        [{"role": "system", "content":
            "You are a Scientific Illustrator for academic publications. "
            "Your job is to identify which figures need to be scientific illustrations "
            "(not data charts) and write detailed visual briefs for an image generation model. "
            "Be specific about scientific accuracy, layout, labels, arrows, and colors. "
            "Respond ONLY in valid JSON."},
         {"role": "user", "content": brief_prompt}],
        temperature=sc["temperature"], max_tokens=sc["max_tokens"],
        structured_output=True,
    )

    if "error" in result:
        # Fallback: treat all as illustrations
        briefs = [{"index": i, "description": d, "is_illustration": True,
                    "visual_brief": d} for i, d in enumerate(unique_requests, 1)]
    else:
        text = result.get("text", "")
        structured = result.get("structured")
        briefs = []
        if isinstance(structured, dict):
            briefs = structured.get("illustrations", [])
        elif isinstance(structured, list):
            briefs = structured
        if not briefs:
            # Try to parse from text
            json_match = re.search(r'\{[\s\S]*"illustrations"[\s\S]*\}', text)
            if json_match:
                try:
                    briefs = json.loads(json_match.group()).get("illustrations", [])
                except (json.JSONDecodeError, Exception):
                    pass
        if not briefs:
            briefs = [{"index": i, "description": d, "is_illustration": True,
                        "visual_brief": d} for i, d in enumerate(unique_requests, 1)]

    # ── Step 3: Generate each illustration via Nano Banana Pro ──
    ill_count = 0
    for brief in briefs:
        if not brief.get("is_illustration", True):
            continue
        desc = brief.get("description", "")
        visual_brief = brief.get("visual_brief", desc)
        if not visual_brief or len(visual_brief) < 5:
            continue

        # Skip if already generated
        if desc.lower().strip() in existing_descs_lower:
            continue

        ill_count += 1
        try:
            ill_prompt = (
                f"Create a clean, publication-quality scientific illustration for an academic paper.\n\n"
                f"VISUAL BRIEF: {visual_brief}\n\n"
                f"Research context: {query[:200]}\n\n"
                f"STYLE REQUIREMENTS:\n"
                f"- Professional scientific diagram suitable for a peer-reviewed journal\n"
                f"- White/light background\n"
                f"- Clear labels with proper scientific nomenclature\n"
                f"- Accurate anatomical/molecular/process details\n"
                f"- Clean arrows and connections where needed\n"
                f"- No text watermarks or decorative elements\n"
                f"- Publication-ready quality (300 DPI equivalent)\n"
            )
            ill_result = _generate_illustration(ill_prompt, api_key, provider)
            if ill_result.get("ok"):
                illustrations.append({
                    "description": desc,
                    "caption": f"Figure: {desc}",
                    "path": ill_result["path"],
                    "chart_type": "illustration",
                    "size": ill_result.get("size", 0),
                    "index": len(existing_figures) + ill_count,
                })
            else:
                errors.append(f"Illustration '{desc}': {ill_result.get('error', 'failed')}")
        except Exception as e:
            errors.append(f"Illustration '{desc}': {e}")

        time.sleep(1)  # Rate limit for image generation API

    return {"ok": True, "illustrations": illustrations, "errors": errors}


def generate_chart(args):
    """Generate a matplotlib chart from data and save as PNG (white/print-friendly background).
    Args: {chart_type, data, title, xlabel, ylabel, labels}
    Returns: {ok, path, chart_type}"""
    chart_type = args.get("chart_type", "bar")
    data = args.get("data", [])
    title = args.get("title", "Chart")
    xlabel = args.get("xlabel", "")
    ylabel = args.get("ylabel", "")
    labels = args.get("labels", [])

    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        # Academic print-friendly style: white background, dark text
        fig, ax = plt.subplots(figsize=(10, 6))
        fig.patch.set_facecolor("white")
        ax.set_facecolor("white")
        ax.tick_params(colors="#333333", labelsize=10)
        ax.xaxis.label.set_color("#222222")
        ax.yaxis.label.set_color("#222222")
        ax.title.set_color("#111111")
        for spine in ax.spines.values():
            spine.set_color("#cccccc")

        # Academic color palette (colorblind-friendly)
        COLORS = ["#2563eb", "#059669", "#d97706", "#dc2626", "#7c3aed", "#0891b2", "#be185d", "#4f46e5"]

        if chart_type == "bar":
            x = labels or [f"Item {i+1}" for i in range(len(data))]
            ax.bar(x, data, color=[COLORS[i % len(COLORS)] for i in range(len(data))], edgecolor="white", linewidth=0.5)
            plt.xticks(rotation=45, ha="right", fontsize=9)
        elif chart_type == "line":
            ax.plot(data, color=COLORS[0], linewidth=2, marker="o", markersize=5, markerfacecolor="white", markeredgecolor=COLORS[0], markeredgewidth=1.5)
            if labels:
                ax.set_xticks(range(len(labels)))
                ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=9)
        elif chart_type == "pie":
            x = labels or [f"Slice {i+1}" for i in range(len(data))]
            wedges, texts, autotexts = ax.pie(data, labels=x,
                colors=[COLORS[i % len(COLORS)] for i in range(len(data))],
                autopct="%1.1f%%", textprops={"color": "#333333", "fontsize": 10},
                wedgeprops={"edgecolor": "white", "linewidth": 1.5})
            for at in autotexts:
                at.set_fontsize(9)
                at.set_fontweight("bold")
        elif chart_type == "scatter":
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], (list, tuple)):
                xs, ys = zip(*data)
                ax.scatter(xs, ys, color=COLORS[0], s=60, alpha=0.8, edgecolors="white", linewidths=0.5)
            else:
                ax.scatter(range(len(data)), data, color=COLORS[0], s=60, alpha=0.8, edgecolors="white", linewidths=0.5)
        elif chart_type == "heatmap":
            try:
                import numpy as np
                arr = np.array(data)
                im = ax.imshow(arr, cmap="YlOrRd", aspect="auto")
                fig.colorbar(im, ax=ax)
            except ImportError:
                return {"ok": False, "error": "numpy required for heatmap"}

        ax.set_title(title, fontsize=14, fontweight="bold", pad=15, color="#111111")
        if xlabel:
            ax.set_xlabel(xlabel, fontsize=11)
        if ylabel:
            ax.set_ylabel(ylabel, fontsize=11)
        ax.grid(True, alpha=0.3, color="#e5e7eb", linestyle="--")
        plt.tight_layout()

        chart_dir = os.path.join(RESEARCH_DIR, "charts")
        os.makedirs(chart_dir, exist_ok=True)
        safe = re.sub(r'[^a-zA-Z0-9_-]', '_', title)[:30]
        path = os.path.join(chart_dir, f"{safe}_{int(time.time())}.png")
        fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        return {"ok": True, "path": path, "chart_type": chart_type, "size": os.path.getsize(path)}
    except ImportError:
        return {"ok": False, "error": "matplotlib required. Install via: pip install matplotlib"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generate_table(args):
    """Generate a formatted data table as markdown and optionally as a print-friendly PNG image.
    Args: {headers, rows, title, format}
    Returns: {ok, markdown, path (if image)}"""
    headers = args.get("headers", [])
    rows = args.get("rows", [])
    title = args.get("title", "Table")
    out_format = args.get("format", "markdown")

    if not headers or not rows:
        return {"ok": False, "error": "Headers and rows required."}

    # Generate markdown table
    md = f"### {title}\n\n"
    md += "| " + " | ".join(str(h) for h in headers) + " |\n"
    md += "| " + " | ".join("---" for _ in headers) + " |\n"
    for row in rows:
        md += "| " + " | ".join(str(c) for c in row) + " |\n"

    result = {"ok": True, "markdown": md, "title": title}

    if out_format == "image":
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            # Calculate proper sizing — wider columns for more data
            col_width = max(1.8, min(3.0, 14.0 / max(len(headers), 1)))
            fig_w = max(8, col_width * len(headers) + 1)
            fig_h = max(2.5, len(rows) * 0.45 + 1.8)
            fig, ax = plt.subplots(figsize=(fig_w, fig_h))
            fig.patch.set_facecolor("white")
            ax.axis("off")

            # Truncate long cell values to prevent overlap
            max_cell_chars = max(12, int(60 / max(len(headers), 1)))
            display_rows = []
            for row in rows:
                display_rows.append([str(c)[:max_cell_chars] + ("..." if len(str(c)) > max_cell_chars else "") for c in row])
            display_headers = [str(h)[:max_cell_chars] + ("..." if len(str(h)) > max_cell_chars else "") for h in headers]

            table = ax.table(cellText=display_rows, colLabels=display_headers, loc="center",
                             cellLoc="center", colColours=["#e0f2fe"] * len(headers))
            table.auto_set_font_size(False)
            table.set_fontsize(9)
            table.auto_set_column_width(list(range(len(headers))))
            table.scale(1, 1.6)

            for key, cell in table.get_celld().items():
                cell.set_edgecolor("#d1d5db")
                cell.set_linewidth(0.5)
                if key[0] == 0:
                    # Header row — bold, dark bg
                    cell.set_text_props(color="#1e3a5f", fontweight="bold", fontsize=10)
                    cell.set_facecolor("#e0f2fe")
                    cell.set_height(cell.get_height() * 1.3)
                else:
                    cell.set_text_props(color="#374151", fontsize=9)
                    cell.set_facecolor("white" if key[0] % 2 == 0 else "#f9fafb")

            ax.set_title(title, color="#111827", fontsize=13, fontweight="bold", pad=12)
            plt.tight_layout(pad=0.5)
            chart_dir = os.path.join(RESEARCH_DIR, "charts")
            os.makedirs(chart_dir, exist_ok=True)
            safe = re.sub(r'[^a-zA-Z0-9_-]', '_', title)[:30]
            path = os.path.join(chart_dir, f"table_{safe}_{int(time.time())}.png")
            fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
            plt.close(fig)
            result["path"] = path
            result["size"] = os.path.getsize(path)
        except ImportError:
            pass

    return result


def generate_section(args):
    """Generate a research paper section.
    Args: {section_type, context, api_key, provider, model}
    Returns: {ok, text, tokens}
    """
    section_type = args.get("section_type", "related_work")
    context = args.get("context", "")
    api_key = args.get("api_key", "")
    provider = args.get("provider", "openai")
    model = args.get("model", "")

    if not api_key:
        return {"error": "API key required."}

    SECTION_PROMPTS = {
        "abstract": "Write a concise academic abstract (150-250 words) for the following research:",
        "introduction": "Write an introduction section for the following research, including motivation, problem statement, and contributions:",
        "related_work": "Write a related work section reviewing the literature for:",
        "methodology": "Write a methodology section describing the approach for:",
        "results": "Write a results section presenting the findings for:",
        "conclusion": "Write a conclusion section summarizing the key findings and future work for:",
        "discussion": "Write a discussion section analyzing the implications of:",
    }

    sys_prompt = (
        "You are an expert academic writer. Write in formal academic style with proper citations. "
        "Use LaTeX citation commands like \\cite{} where appropriate. "
        "Be thorough, precise, and well-structured."
    )

    section_prompt = SECTION_PROMPTS.get(section_type, f"Write a {section_type} section for:")
    full_prompt = f"{section_prompt}\n\n{context}"

    result = _llm_chat(provider, api_key, model,
                       [{"role": "system", "content": sys_prompt},
                        {"role": "user", "content": full_prompt}],
                       temperature=0.4, max_tokens=4096)

    if "error" in result:
        return result

    return {
        "ok": True, "text": result["text"],
        "tokens": result.get("usage"),
        "section_type": section_type,
    }
